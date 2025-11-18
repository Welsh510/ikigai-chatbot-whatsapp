# Enhanced WhatsApp Business API Chatbot with Full Media Support and Language Detection
# Requirements:
# pip install flask requests python-dotenv openai python-telegram-bot mysql-connector-python pandas langid PyPDF2 python-docx

import os
import json
import requests
import logging
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_file
import openai
from dotenv import load_dotenv
import time
import mysql.connector
from mysql.connector import Error
import pandas as pd
import langid  # For language detection
import tempfile
import os
from urllib.parse import urlparse
import hashlib
import mimetypes
from pathlib import Path
import PyPDF2
import docx
import io
import base64
from flask import send_from_directory
import subprocess
import threading
from queue import Queue
from concurrent.futures import ThreadPoolExecutor
import pytz
import re

# Malaysia Timezone Configuration - Add after imports, before classes
MALAYSIA_TZ = pytz.timezone('Asia/Kuala_Lumpur')  # UTC+8

def get_malaysia_time():
    """Get current time in Malaysia timezone"""
    return datetime.now(MALAYSIA_TZ)

def get_malaysia_timestamp():
    """Get Malaysia time as timestamp string for database"""
    return get_malaysia_time().strftime('%Y-%m-%d %H:%M:%S')

def convert_to_malaysia_time(dt):
    """Convert datetime to Malaysia timezone"""
    if dt is None:
        return None
    if dt.tzinfo is None:
        # Assume UTC if no timezone info
        dt = pytz.UTC.localize(dt)
    return dt.astimezone(MALAYSIA_TZ)

class MalaysiaTimeFormatter(logging.Formatter):
    def formatTime(self, record, datefmt=None):
        malaysia_time = get_malaysia_time()
        if datefmt:
            return malaysia_time.strftime(datefmt)
        else:
            return malaysia_time.strftime('%Y-%m-%d %H:%M:%S %Z')
            
class AutoGitManager:
    """Manages automatic git operations with proper synchronization"""
    
    def __init__(self, repo_path="."):
        self.repo_path = repo_path
        self.git_queue = Queue()
        self.is_running = True
        self.default_branch = None  # Will be detected
        
        # Setup git configuration
        self.setup_git()
        
        # Initialize repository state
        self.initialize_repository()
        
        # Detect the default branch properly
        self.detect_default_branch()
        
        # Start background git processor
        self.git_thread = threading.Thread(target=self._process_git_queue, daemon=True)
        self.git_thread.start()
        
        logging.info(f"AutoGitManager initialized with default branch: {self.default_branch}")
    
    def initialize_repository(self):
        """Initialize repository and sync with remote"""
        try:
            logging.info("Initializing repository state...")
            
            # First, ensure we have a clean working directory
            self.clean_working_directory()
            
            # Fetch latest from remote
            self.fetch_from_remote()
            
            logging.info("Repository initialization completed")
            
        except Exception as e:
            logging.error(f"Error initializing repository: {e}")
    
    def clean_working_directory(self):
        """Clean working directory of uncommitted changes"""
        try:
            # Check if there are any changes
            status_result = subprocess.run([
                "git", "status", "--porcelain"
            ], cwd=self.repo_path, capture_output=True, text=True)
            
            if status_result.returncode == 0 and status_result.stdout.strip():
                logging.info("Found uncommitted changes, cleaning up...")
                
                # Add all changes
                subprocess.run([
                    "git", "add", "."
                ], cwd=self.repo_path, capture_output=True, text=True)
                
                # MODIFIED: Commit changes with Malaysia time
                commit_message = f"Auto-cleanup: commit existing changes - {get_malaysia_timestamp()}"
                commit_result = subprocess.run([
                    "git", "commit", "-m", commit_message
                ], cwd=self.repo_path, capture_output=True, text=True)
                
                if commit_result.returncode == 0:
                    logging.info("Successfully committed existing changes")
                else:
                    # If commit fails, try stashing instead
                    logging.info("Commit failed, trying to stash changes...")
                    subprocess.run([
                        "git", "stash", "--include-untracked"
                    ], cwd=self.repo_path, capture_output=True, text=True)
                    logging.info("Changes stashed")
            
        except Exception as e:
            logging.error(f"Error cleaning working directory: {e}")
    
    def fetch_from_remote(self):
        """Fetch latest changes from remote"""
        try:
            fetch_result = subprocess.run([
                "git", "fetch", "origin"
            ], cwd=self.repo_path, capture_output=True, text=True, timeout=60)
            
            if fetch_result.returncode == 0:
                logging.info("Successfully fetched from remote")
                
                # Set remote HEAD
                subprocess.run([
                    "git", "remote", "set-head", "origin", "-a"
                ], cwd=self.repo_path, capture_output=True, timeout=30)
            else:
                logging.warning(f"Fetch failed: {fetch_result.stderr}")
                
        except Exception as e:
            logging.error(f"Error fetching from remote: {e}")
    
    def detect_default_branch(self):
        """Enhanced default branch detection"""
        try:
            # First, try to get the default branch from GitHub API
            github_token = os.getenv("GITHUB_TOKEN")
            if github_token:
                try:
                    api_url = "https://api.github.com/repos/Welsh510/python-chatbot"
                    headers = {
                        "Authorization": f"token {github_token}",
                        "Accept": "application/vnd.github.v3+json"
                    }
                    response = requests.get(api_url, headers=headers, timeout=10)
                    
                    if response.status_code == 200:
                        repo_data = response.json()
                        github_default = repo_data.get('default_branch')
                        if github_default:
                            self.default_branch = github_default
                            logging.info(f"GitHub API detected default branch: {self.default_branch}")
                            self.ensure_branch_sync()
                            return
                except Exception as api_error:
                    logging.warning(f"GitHub API detection failed: {api_error}")
            
            # Fallback detection methods
            self.default_branch = self.detect_from_remote_refs()
            if not self.default_branch:
                self.default_branch = "main"  # Modern default
            
            logging.info(f"Using default branch: {self.default_branch}")
            self.ensure_branch_sync()
                
        except Exception as e:
            logging.error(f"Error detecting default branch: {e}")
            self.default_branch = "main"
    
    def detect_from_remote_refs(self):
        """Detect default branch from remote refs"""
        try:
            # Try to get remote HEAD
            result = subprocess.run([
                "git", "symbolic-ref", "refs/remotes/origin/HEAD"
            ], cwd=self.repo_path, capture_output=True, text=True)
            
            if result.returncode == 0 and result.stdout.strip():
                remote_ref = result.stdout.strip()
                if "refs/remotes/origin/" in remote_ref:
                    return remote_ref.split("refs/remotes/origin/")[1]
            
            # Check remote branches
            result = subprocess.run([
                "git", "branch", "-r"
            ], cwd=self.repo_path, capture_output=True, text=True)
            
            if result.returncode == 0:
                remote_branches = result.stdout.strip()
                if "origin/main" in remote_branches:
                    return "main"
                elif "origin/master" in remote_branches:
                    return "master"
            
            return None
            
        except Exception as e:
            logging.error(f"Error detecting from remote refs: {e}")
            return None
    
    def ensure_branch_sync(self):
        """Ensure local branch is in sync with remote"""
        try:
            # Check current branch
            current_result = subprocess.run([
                "git", "branch", "--show-current"
            ], cwd=self.repo_path, capture_output=True, text=True)
            
            current_branch = current_result.stdout.strip() if current_result.returncode == 0 else ""
            
            if current_branch != self.default_branch:
                logging.info(f"Switching from '{current_branch}' to '{self.default_branch}'")
                self.switch_to_branch(self.default_branch)
            
            # Sync with remote
            self.sync_with_remote()
            
        except Exception as e:
            logging.error(f"Error ensuring branch sync: {e}")
    
    def switch_to_branch(self, branch_name):
        """Switch to specified branch with multiple fallback methods"""
        try:
            # Method 1: Simple checkout
            result = subprocess.run([
                "git", "checkout", branch_name
            ], cwd=self.repo_path, capture_output=True, text=True)
            
            if result.returncode == 0:
                logging.info(f"Successfully switched to {branch_name}")
                return True
            
            # Method 2: Checkout with tracking
            result = subprocess.run([
                "git", "checkout", "-b", branch_name, f"origin/{branch_name}"
            ], cwd=self.repo_path, capture_output=True, text=True)
            
            if result.returncode == 0:
                logging.info(f"Successfully created tracking branch {branch_name}")
                return True
            
            # Method 3: Force checkout
            subprocess.run([
                "git", "reset", "--hard", "HEAD"
            ], cwd=self.repo_path, capture_output=True, text=True)
            
            result = subprocess.run([
                "git", "checkout", "-B", branch_name
            ], cwd=self.repo_path, capture_output=True, text=True)
            
            if result.returncode == 0:
                logging.info(f"Force created branch {branch_name}")
                return True
            
            logging.error(f"All methods failed to switch to {branch_name}")
            return False
            
        except Exception as e:
            logging.error(f"Error switching to branch {branch_name}: {e}")
            return False
    
    def sync_with_remote(self):
        """Synchronize local branch with remote"""
        try:
            # Check if remote branch exists
            remote_check = subprocess.run([
                "git", "ls-remote", "--heads", "origin", self.default_branch
            ], cwd=self.repo_path, capture_output=True, text=True)
            
            if remote_check.returncode == 0 and remote_check.stdout.strip():
                # Remote branch exists, sync with it
                logging.info(f"Syncing local {self.default_branch} with remote")
                
                # Reset to remote state to avoid conflicts
                reset_result = subprocess.run([
                    "git", "reset", "--hard", f"origin/{self.default_branch}"
                ], cwd=self.repo_path, capture_output=True, text=True)
                
                if reset_result.returncode == 0:
                    logging.info(f"Successfully synced with origin/{self.default_branch}")
                else:
                    logging.warning(f"Reset failed: {reset_result.stderr}")
            else:
                logging.info(f"Remote branch origin/{self.default_branch} doesn't exist yet")
                
        except Exception as e:
            logging.error(f"Error syncing with remote: {e}")
            
    def setup_git(self):
        """Setup git configuration and authentication"""
        try:
            # Initialize git if needed
            if not os.path.exists(os.path.join(self.repo_path, ".git")):
                subprocess.run(["git", "init"], cwd=self.repo_path, check=True)
                logging.info("Git repository initialized")
            
            # Configure git user
            subprocess.run([
                "git", "config", "user.email", "chatbot@railway.app"
            ], cwd=self.repo_path, capture_output=True)
            
            subprocess.run([
                "git", "config", "user.name", "Railway ChatBot"
            ], cwd=self.repo_path, capture_output=True)
            
            # Setup GitHub authentication
            github_token = os.getenv("GITHUB_TOKEN")
            if github_token:
                repo_url = f"https://{github_token}@github.com/Welsh510/python-chatbot.git"
                
                # Check if remote exists
                result = subprocess.run([
                    "git", "remote", "get-url", "origin"
                ], cwd=self.repo_path, capture_output=True)
                
                if result.returncode == 0:
                    # Update existing remote
                    subprocess.run([
                        "git", "remote", "set-url", "origin", repo_url
                    ], cwd=self.repo_path, check=True)
                else:
                    # Add new remote
                    subprocess.run([
                        "git", "remote", "add", "origin", repo_url
                    ], cwd=self.repo_path, check=True)
                
                logging.info("Git authentication configured with GitHub token")
            else:
                logging.warning("GITHUB_TOKEN not found - git push will fail")
                
            # Configure git settings
            subprocess.run([
                "git", "config", "credential.helper", "store"
            ], cwd=self.repo_path, capture_output=True)
            
            subprocess.run([
                "git", "config", "pull.rebase", "false"
            ], cwd=self.repo_path, capture_output=True)
            
            # Set push behavior
            subprocess.run([
                "git", "config", "push.default", "simple"
            ], cwd=self.repo_path, capture_output=True)
            
        except Exception as e:
            logging.error(f"Error setting up git: {e}")
    
    def add_file_to_queue(self, file_path, commit_message):
        """Add a file to the git processing queue"""
        self.git_queue.put({
            'action': 'commit_and_push',
            'file_path': file_path,
            'commit_message': commit_message,
            'timestamp': time.time()  # This can remain as Unix timestamp for internal queue processing
        })
        logging.info(f"Added file to git queue: {file_path}")
    
    def _process_git_queue(self):
        """Background thread to process git operations"""
        while self.is_running:
            try:
                try:
                    git_task = self.git_queue.get(timeout=30)
                except:
                    continue
                
                if git_task['action'] == 'commit_and_push':
                    success = self._commit_and_push_file_safe(
                        git_task['file_path'], 
                        git_task['commit_message']
                    )
                    
                    if success:
                        logging.info(f"Successfully committed and pushed to {self.default_branch}: {git_task['file_path']}")
                    else:
                        logging.error(f"Failed to commit and push: {git_task['file_path']}")
                
                self.git_queue.task_done()
                
            except Exception as e:
                logging.error(f"Error in git background processor: {e}")
                time.sleep(5)
    
    def _commit_and_push_file_safe(self, file_path, commit_message):
        """Safely commit and push a file with proper synchronization"""
        try:
            # Ensure clean state before operations
            self.ensure_clean_state()
            
            # Add file to git
            add_result = subprocess.run([
                "git", "add", file_path
            ], cwd=self.repo_path, capture_output=True, text=True, timeout=30)
            
            if add_result.returncode != 0:
                logging.error(f"Git add failed: {add_result.stderr}")
                return False
            
            # Check if there are changes to commit
            status_result = subprocess.run([
                "git", "status", "--porcelain"
            ], cwd=self.repo_path, capture_output=True, text=True, timeout=30)
            
            if not status_result.stdout.strip():
                logging.info(f"No changes to commit for: {file_path}")
                return True
            
            # MODIFIED: Commit the file with Malaysia timestamp
            commit_message_with_time = f"{commit_message} - {get_malaysia_timestamp()}"
            commit_result = subprocess.run([
                "git", "commit", "-m", commit_message_with_time
            ], cwd=self.repo_path, capture_output=True, text=True, timeout=30)
            
            if commit_result.returncode != 0:
                logging.error(f"Git commit failed: {commit_result.stderr}")
                return False
            
            # Sync and push safely
            return self.safe_push()
                
        except Exception as e:
            logging.error(f"Error in safe commit and push: {e}")
            return False
    
    def ensure_clean_state(self):
        """Ensure repository is in clean state for operations"""
        try:
            # Pull latest changes first
            pull_result = subprocess.run([
                "git", "pull", "origin", self.default_branch, "--no-rebase"
            ], cwd=self.repo_path, capture_output=True, text=True, timeout=60)
            
            if pull_result.returncode != 0:
                # If pull fails, try to resolve conflicts
                logging.warning(f"Pull failed, attempting to resolve: {pull_result.stderr}")
                
                # Reset to known state and retry
                subprocess.run([
                    "git", "fetch", "origin"
                ], cwd=self.repo_path, capture_output=True, timeout=60)
                
                # Merge strategy
                merge_result = subprocess.run([
                    "git", "merge", f"origin/{self.default_branch}", "--no-edit"
                ], cwd=self.repo_path, capture_output=True, text=True, timeout=60)
                
                if merge_result.returncode == 0:
                    logging.info("Successfully merged remote changes")
                else:
                    logging.warning(f"Merge failed: {merge_result.stderr}")
            else:
                logging.info("Successfully pulled latest changes")
                
        except Exception as e:
            logging.error(f"Error ensuring clean state: {e}")
    
    def safe_push(self):
        """Safely push changes to remote"""
        try:
            # First attempt: normal push
            push_result = subprocess.run([
                "git", "push", "origin", self.default_branch
            ], cwd=self.repo_path, capture_output=True, text=True, timeout=120)
            
            if push_result.returncode == 0:
                logging.info(f"Successfully pushed to {self.default_branch}")
                return True
            
            # If normal push fails, try with upstream
            logging.info("Normal push failed, trying with upstream...")
            upstream_result = subprocess.run([
                "git", "push", "--set-upstream", "origin", self.default_branch
            ], cwd=self.repo_path, capture_output=True, text=True, timeout=120)
            
            if upstream_result.returncode == 0:
                logging.info(f"Successfully pushed with upstream to {self.default_branch}")
                return True
            
            # If that fails, try force push (last resort)
            logging.warning("Attempting force push as last resort...")
            force_result = subprocess.run([
                "git", "push", "--force-with-lease", "origin", self.default_branch
            ], cwd=self.repo_path, capture_output=True, text=True, timeout=120)
            
            if force_result.returncode == 0:
                logging.info(f"Successfully force pushed to {self.default_branch}")
                return True
            else:
                logging.error(f"All push attempts failed: {force_result.stderr}")
                return False
                
        except Exception as e:
            logging.error(f"Error in safe push: {e}")
            return False
    
    def get_current_branch(self):
        """Get the current git branch"""
        try:
            result = subprocess.run([
                "git", "branch", "--show-current"
            ], cwd=self.repo_path, capture_output=True, text=True, timeout=10)
            
            return result.stdout.strip() if result.returncode == 0 else self.default_branch
                
        except Exception as e:
            logging.error(f"Error getting current branch: {e}")
            return self.default_branch
    
    def stop(self):
        """Stop the background git processor"""
        self.is_running = False
        if hasattr(self, 'git_thread'):
            self.git_thread.join(timeout=10)
            
class ChatHistoryManager:
    def __init__(self):
        # Storage paths for all media types
        self.audio_storage_path = Path("media") / "audio"
        self.audio_storage_path.mkdir(parents=True, exist_ok=True)
        
        self.image_storage_path = Path("media") / "images"
        self.image_storage_path.mkdir(parents=True, exist_ok=True)
        
        self.video_storage_path = Path("media") / "videos"
        self.video_storage_path.mkdir(parents=True, exist_ok=True)
        
        self.document_storage_path = Path("media") / "documents"
        self.document_storage_path.mkdir(parents=True, exist_ok=True)

        # Add file storage path for additional content files
        self.files_storage_path = Path("media") / "files"
        self.files_storage_path.mkdir(parents=True, exist_ok=True)

        self.connection = self.create_connection()

        if self.connection and self.connection.is_connected():
            logging.info("Database connection established")
            self.initialize_database_tables()
        else:
            logging.error("Failed to establish database connection")
    
    def initialize_database_tables(self):
        """Create required database tables if they don't exist"""
        if not self.connection:
            return False
            
        try:
            cursor = self.connection.cursor()
            
            # ADD: Set session timezone to Malaysia time
            cursor.execute("SET time_zone = '+08:00'")
        
            # Create chat_sessions table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS chat_sessions (
                    session_id INT AUTO_INCREMENT PRIMARY KEY,
                    phone_number VARCHAR(20) NOT NULL,
                    session_start DATETIME DEFAULT CURRENT_TIMESTAMP,
                    session_end DATETIME NULL,
                    message_count INT DEFAULT 0,
                    new_message_count INT DEFAULT 0,
                    last_activity DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    current_opening_step INT DEFAULT 0,
                    opening_sequence_completed BOOLEAN DEFAULT FALSE,
                    human_reply TINYINT(1) DEFAULT 0,
                    manual_reply TINYINT(1) DEFAULT 0,
                    appointment_section TINYINT(1) DEFAULT 0,
                    human_section TINYINT(1) DEFAULT 0,
                    INDEX idx_phone (phone_number)
                )
            ''')
            
            # Create chat_messages table with ONLY processed column added
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS chat_messages (
                    message_id INT AUTO_INCREMENT PRIMARY KEY,
                    session_id INT,
                    phone_number VARCHAR(20) NOT NULL,
                    message TEXT,
                    role ENUM('user', 'assistant') NOT NULL,
                    message_type ENUM('text', 'audio', 'image', 'video', 'document') DEFAULT 'text',
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                    opening_message_id INT NULL,
                    processed BOOLEAN DEFAULT FALSE,
                    INDEX idx_phone_time (phone_number, timestamp),
                    INDEX idx_processed (processed)
                )
            ''')
            
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS chat_followup_keep (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    phone_number VARCHAR(20) NOT NULL UNIQUE,
                    created_datetime DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_phone_number (phone_number),
                    INDEX idx_created_datetime (created_datetime)
                )
            ''')

            # --- NEW: opening_categories (主表) ---
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS opening_categories (
                    PKKEY INT AUTO_INCREMENT PRIMARY KEY,
                    TITLE VARCHAR(255),
                    KEYWORDS TEXT NULL,
                    SEQUENCE INT DEFAULT 1,
                    STATUS TINYINT DEFAULT 1,
                    `DEFAULT` TINYINT DEFAULT 0,
                    CREATED_AT DATETIME NULL DEFAULT CURRENT_TIMESTAMP,
                    UPDATED_AT DATETIME NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                )
            ''')
            
            # NEW: Create special_reply table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS special_reply (
                    PKKEY INT AUTO_INCREMENT PRIMARY KEY,
                    KEYWORDS TEXT,
                    STATUS TINYINT DEFAULT 1,
                    COOLDOWN_HOURS INT DEFAULT 24,
                    CREATED_AT DATETIME DEFAULT CURRENT_TIMESTAMP,
                    UPDATED_AT DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    INDEX idx_status (STATUS)
                )
            ''')
            
            # --- NEW: opening_messages (子表) ---
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS opening_messages (
                    PKKEY INT AUTO_INCREMENT PRIMARY KEY,
                    CATEGORY_ID INT NOT NULL,
                    MESSAGE TEXT,
                    SEQUENCE INT DEFAULT 1,
                    STATUS TINYINT DEFAULT 1,
                    CREATED_AT DATETIME NULL DEFAULT CURRENT_TIMESTAMP,
                    UPDATED_AT DATETIME NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    FOREIGN KEY (CATEGORY_ID) REFERENCES opening_categories(PKKEY) ON DELETE CASCADE,
                    INDEX idx_category (CATEGORY_ID),
                    INDEX idx_sequence (CATEGORY_ID, SEQUENCE)
                )
            ''')
            
            # ============================================================================
            # MODIFIED: Create special_reply_usage table with user_id instead of phone_number
            # ============================================================================
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS special_reply_usage (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    user_id VARCHAR(20) NOT NULL,
                    special_reply_id INT NOT NULL,
                    last_triggered_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    trigger_count INT DEFAULT 1,
                    INDEX idx_user_special (user_id, special_reply_id),
                    INDEX idx_triggered_at (last_triggered_at)
                )
            ''')
            
            # Add COOLDOWN_HOURS column if it doesn't exist (for existing installations)
            cursor.execute('''
                SELECT COUNT(*)
                FROM information_schema.columns
                WHERE table_schema = DATABASE() 
                AND table_name = 'special_reply' 
                AND column_name = 'COOLDOWN_HOURS'
            ''')
            if cursor.fetchone()[0] == 0:
                cursor.execute('''
                    ALTER TABLE special_reply
                    ADD COLUMN COOLDOWN_HOURS INT DEFAULT 24
                    AFTER STATUS
                ''')
                logging.info("Added COOLDOWN_HOURS column to special_reply table")
                
            # Create master_opening table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS master_opening (
                    PKKEY INT AUTO_INCREMENT PRIMARY KEY,
                    MESSAGE TEXT,
                    SEQUENCE INT DEFAULT 1,
                    STATUS TINYINT DEFAULT 1,
                    CREATED_AT DATETIME NULL DEFAULT CURRENT_TIMESTAMP,
                    UPDATED_AT DATETIME NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                )
            ''')
            
            # Add session_id column to chat_messages if it doesn't exist
            cursor.execute('''
                SELECT COUNT(*)
                FROM information_schema.columns
                WHERE table_schema = DATABASE() 
                AND table_name = 'chat_messages' 
                AND column_name = 'session_id'
            ''')
            if cursor.fetchone()[0] == 0:
                cursor.execute('''
                    ALTER TABLE chat_messages
                    ADD COLUMN session_id INT AFTER message_id,
                    ADD FOREIGN KEY (session_id) REFERENCES chat_sessions(session_id) ON DELETE CASCADE
                ''')
                logging.info("Added session_id column to chat_messages table")
            
            # Add opening_message_id column if it doesn't exist
            cursor.execute('''
                SELECT COUNT(*)
                FROM information_schema.columns
                WHERE table_schema = DATABASE() 
                AND table_name = 'chat_messages' 
                AND column_name = 'opening_message_id'
            ''')
            if cursor.fetchone()[0] == 0:
                cursor.execute('''
                    ALTER TABLE chat_messages
                    ADD COLUMN opening_message_id INT NULL AFTER timestamp
                ''')
                logging.info("Added opening_message_id column to chat_messages table")
            
            # Add processed column if it doesn't exist
            cursor.execute('''
                SELECT COUNT(*)
                FROM information_schema.columns
                WHERE table_schema = DATABASE() 
                AND table_name = 'chat_messages' 
                AND column_name = 'processed'
            ''')
            if cursor.fetchone()[0] == 0:
                cursor.execute('''
                    ALTER TABLE chat_messages
                    ADD COLUMN processed BOOLEAN DEFAULT FALSE AFTER opening_message_id,
                    ADD INDEX idx_processed (processed)
                ''')
                logging.info("Added processed column to chat_messages table")
            
            self.connection.commit()
            cursor.close()
            logging.info("Database tables initialized successfully")
            return True
            
        except Exception as e:
            logging.error(f"Error initializing database tables: {e}")
            return False
        
    def create_connection(self):
        """Create connection to MySQL database - Railway Fixed"""
        try:
        # Get MySQL credentials from environment
            mysql_host = os.getenv("MYSQL_HOST")
            mysql_user = os.getenv("MYSQL_USER")
            mysql_password = os.getenv("MYSQL_PASSWORD")
            mysql_database = os.getenv("MYSQL_DATABASE")
            mysql_port = os.getenv("MYSQL_PORT", "3306")
            
            # Check if we have the required credentials
            if not all([mysql_host, mysql_user, mysql_password, mysql_database]):
                logging.warning("Missing MySQL environment variables, trying Railway URL format...")
                
                # Try to get MySQL URL (Railway format)
                mysql_url = os.getenv("MYSQL_URL")
                if mysql_url:
                    # Parse MySQL URL format: mysql://user:password@host:port/database
                    import re
                    match = re.match(r'mysql://([^:]+):([^@]+)@([^:]+):(\d+)/(.+)', mysql_url)
                    if match:
                        mysql_user, mysql_password, mysql_host, mysql_port, mysql_database = match.groups()
                        logging.info("Successfully parsed MySQL URL")
                    else:
                        logging.error("Could not parse MySQL URL format")
                        return None
                else:
                    logging.error("No MySQL credentials or URL found")
                    return None
            
            # Log connection attempt (without password)
            logging.info(f"Attempting MySQL connection to {mysql_host}:{mysql_port} as {mysql_user}")
            
            connection = mysql.connector.connect(
                host=mysql_host,
                user=mysql_user,
                password=mysql_password,
                database=mysql_database,
                port=int(mysql_port),
                autocommit=True,
                connection_timeout=30,
                charset='utf8mb4',
                collation='utf8mb4_unicode_ci',
                time_zone='+08:00'  # ADD this line
            )
            
            if connection.is_connected():
                cursor = connection.cursor()
                cursor.execute("SET time_zone = '+08:00'")
                cursor.close()
                
                logging.info("MySQL connection successful with Malaysia timezone!")
                return connection
            else:
                logging.error("MySQL connection failed")
                return None
                
        except Error as e:
            logging.error(f"MySQL connection error: {e}")
            return None
        except Exception as e:
            logging.error(f"Unexpected error connecting to MySQL: {e}")
            return None

    def is_first_time_user(self, phone_number):
        """Check if this is a first-time user with improved logic"""
        if not self.connection:
            return True  # Assume first time if no database connection
            
        try:
            cursor = self.connection.cursor()
            
            # Check if user has any chat sessions at all
            cursor.execute('''
                SELECT COUNT(*) FROM chat_sessions
                WHERE phone_number = %s
            ''', (phone_number,))
            
            session_count = cursor.fetchone()[0]
            
            # If no sessions exist, definitely first time
            if session_count == 0:
                cursor.close()
                return True
            
            # Check if user has completed opening sequence
            cursor.execute('''
                SELECT opening_sequence_completed 
                FROM chat_sessions 
                WHERE phone_number = %s 
                ORDER BY session_start DESC 
                LIMIT 1
            ''', (phone_number,))
            
            result = cursor.fetchone()
            cursor.close()
            
            if result:
                # If opening sequence is not completed, treat as first time
                return not result[0]
            
            # If we can't determine, assume not first time
            return False
            
        except Error as e:
            logging.error(f"Error checking first-time user: {e}")
            return True  # Assume first time on error for safety
    
    def get_opening_categories(self, default_only=None):
        """读取 opening_categories；default_only 可为 0 / 1 / None"""
        if not self.connection:
            return []
        try:
            cursor = self.connection.cursor(dictionary=True)
            sql = "SELECT PKKEY, TITLE, KEYWORDS, SEQUENCE, STATUS, `DEFAULT` FROM opening_categories WHERE STATUS=1"
            params = []
            if default_only is not None:
                sql += " AND `DEFAULT`=%s"
                params.append(int(default_only))
            sql += " ORDER BY SEQUENCE ASC, PKKEY ASC"
            cursor.execute(sql, params)
            return cursor.fetchall() or []
        except Exception as e:
            logging.error(f"Error get_opening_categories: {e}")
            return []

    def get_default_opening_category(self):
        """返回 default=1 的分类（如果有多条，取 SEQUENCE 最前的一条）"""
        cats = self.get_opening_categories(default_only=1)
        return cats[0] if cats else None

    def get_opening_messages_by_category(self, category_id):
        """按分类取 master_opening 的开场消息（按 sequence）"""
        if not self.connection:
            return []
        try:
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute("""
                SELECT PKKEY, CATEGORY_ID, TYPE, TITLE, TEXT_CONTENT, MEDIANAME, SEQUENCE, STATUS
                FROM master_opening
                WHERE STATUS=1 AND CATEGORY_ID=%s
                ORDER BY SEQUENCE ASC, PKKEY ASC
            """, (category_id,))
            return cursor.fetchall() or []
        except Exception as e:
            logging.error(f"Error get_opening_messages_by_category({category_id}): {e}")
            return []
            
    def get_opening_messages(self):
        """Get opening messages ordered by sequence"""
        if not self.connection:
            return []
            
        try:
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute('''
                SELECT PKKEY, TYPE, TITLE, TEXT_CONTENT, MEDIANAME, SEQUENCE
                FROM master_opening 
                WHERE STATUS = 1 
                ORDER BY SEQUENCE ASC
            ''')
            
            opening_messages = cursor.fetchall()
            cursor.close()
            
            logging.info(f"Retrieved {len(opening_messages)} opening messages")
            return opening_messages
            
        except Error as e:
            logging.error(f"Error getting opening messages: {e}")
            return []

    def get_user_opening_progress(self, phone_number):
        """Get user's opening sequence progress"""
        if not self.connection:
            return 0, False
            
        try:
            cursor = self.connection.cursor()
            cursor.execute('''
                SELECT current_opening_step, opening_sequence_completed
                FROM chat_sessions 
                WHERE phone_number = %s 
                ORDER BY session_start DESC 
                LIMIT 1
            ''', (phone_number,))
            
            result = cursor.fetchone()
            cursor.close()
            
            if result:
                return result[0], result[1]
            else:
                return 0, False
                
        except Error as e:
            logging.error(f"Error getting opening progress: {e}")
            return 0, False

    def update_opening_progress(self, phone_number, step):
        """Update user's opening sequence progress"""
        if not self.connection:
            return False
            
        try:
            cursor = self.connection.cursor()
            
            # Get opening messages count to check if sequence is completed
            opening_messages = self.get_opening_messages()
            total_steps = len(opening_messages)
            
            is_completed = step >= total_steps
            
            cursor.execute('''
                UPDATE chat_sessions 
                SET current_opening_step = %s, 
                    opening_sequence_completed = %s,
                    last_activity = CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00')
                WHERE phone_number = %s
            ''', (step, is_completed, phone_number))
            
            self.connection.commit()
            cursor.close()
            
            logging.info(f"Updated opening progress for {phone_number}: step {step}, completed: {is_completed}")
            return True
            
        except Error as e:
            logging.error(f"Error updating opening progress: {e}")
            return False
    
    def get_special_replies(self):
        """Get all active special replies from database"""
        if not self.connection: 
            return []
        try:
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute("""
                SELECT PKKEY, KEYWORDS, STATUS
                FROM special_reply
                WHERE STATUS=1
                ORDER BY PKKEY ASC
            """)
            results = cursor.fetchall()
            cursor.close()
            return results or []
        except Exception as e:
            logging.error(f"Error get_special_replies: {e}")
            return []

    def get_special_reply_subs(self, special_id: int):
        """Get special reply sub-items by special_id"""
        if not self.connection: 
            return []
        try:
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute("""
                SELECT PKKEY, CATEGORY_ID, TYPE, TEXT_CONTENT, MEDIANAME, SEQUENCE
                FROM special_reply_sub
                WHERE CATEGORY_ID=%s
                ORDER BY SEQUENCE ASC, PKKEY ASC
            """, (special_id,))
            results = cursor.fetchall()
            cursor.close()
            return results or []
        except Exception as e:
            logging.error(f"Error get_special_reply_subs({special_id}): {e}")
            return []
    
    def is_special_reply_in_cooldown_for_user(self, phone_number, special_reply_id):
        """
        Check if a specific special reply is in cooldown for a specific user
        Returns: (is_in_cooldown: bool, remaining_hours: float)
        """
        if not self.connection:
            return False, 0
            
        try:
            cursor = self.connection.cursor()
            
            # Get cooldown hours for this special reply
            cursor.execute('''
                SELECT COOLDOWN_HOURS FROM special_reply 
                WHERE PKKEY = %s AND STATUS = 1
            ''', (special_reply_id,))
            
            result = cursor.fetchone()
            if not result:
                cursor.close()
                return False, 0
            
            cooldown_hours = result[0] or 24  # Default to 24 hours if NULL
            
            # MODIFIED: Check last trigger time using user_id instead of phone_number
            cursor.execute('''
                SELECT last_triggered_at 
                FROM special_reply_usage 
                WHERE user_id = %s AND special_reply_id = %s
                ORDER BY last_triggered_at DESC
                LIMIT 1
            ''', (phone_number, special_reply_id))
            
            last_trigger_result = cursor.fetchone()
            cursor.close()
            
            if not last_trigger_result:
                # Never triggered by this user
                logging.info(f"Special reply {special_reply_id} never triggered by user {phone_number}")
                return False, 0
            
            last_triggered_at = last_trigger_result[0]
            
            # Calculate time difference
            from datetime import datetime, timedelta
            current_time = get_malaysia_time().replace(tzinfo=None)
            time_diff = current_time - last_triggered_at
            hours_since_trigger = time_diff.total_seconds() / 3600
            
            if hours_since_trigger < cooldown_hours:
                remaining_hours = cooldown_hours - hours_since_trigger
                logging.info(f"Special reply {special_reply_id} is in cooldown for user {phone_number}. "
                            f"Triggered {hours_since_trigger:.1f}h ago, {remaining_hours:.1f}h remaining")
                return True, remaining_hours
            else:
                logging.info(f"Special reply {special_reply_id} cooldown expired for user {phone_number}. "
                            f"Last triggered {hours_since_trigger:.1f}h ago")
                return False, 0
            
        except Exception as e:
            logging.error(f"Error checking per-user special reply cooldown: {e}")
            return False, 0  # Allow on error to prevent blocking

    def record_special_reply_usage(self, phone_number, special_reply_id):
        """
        Record that a user has triggered a specific special reply
        Updates existing record or creates new one
        """
        if not self.connection:
            return False
            
        try:
            cursor = self.connection.cursor()
            
            # MODIFIED: Check if record exists using user_id instead of phone_number
            cursor.execute('''
                SELECT id, trigger_count FROM special_reply_usage 
                WHERE user_id = %s AND special_reply_id = %s
            ''', (phone_number, special_reply_id))
            
            existing_record = cursor.fetchone()
            
            if existing_record:
                # Update existing record
                record_id, current_count = existing_record
                cursor.execute('''
                    UPDATE special_reply_usage 
                    SET last_triggered_at = CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00'),
                        trigger_count = trigger_count + 1
                    WHERE id = %s
                ''', (record_id,))
                new_count = current_count + 1
                logging.info(f"Updated special reply usage: user {phone_number}, special_reply {special_reply_id}, count: {new_count}")
            else:
                # MODIFIED: Create new record with user_id instead of phone_number
                cursor.execute('''
                    INSERT INTO special_reply_usage (user_id, special_reply_id, last_triggered_at, trigger_count)
                    VALUES (%s, %s, CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00'), 1)
                ''', (phone_number, special_reply_id))
                logging.info(f"Created new special reply usage record: user {phone_number}, special_reply {special_reply_id}")
            
            self.connection.commit()
            cursor.close()
            return True
            
        except Exception as e:
            logging.error(f"Error recording special reply usage: {e}")
            return False

    def get_user_special_reply_stats(self, phone_number):
        """Get special reply usage statistics for a specific user"""
        if not self.connection:
            return []
            
        try:
            cursor = self.connection.cursor(dictionary=True)
            
            # MODIFIED: Query using user_id instead of phone_number
            cursor.execute('''
                SELECT 
                    sru.special_reply_id,
                    sr.KEYWORDS,
                    sr.COOLDOWN_HOURS,
                    sru.last_triggered_at,
                    sru.trigger_count,
                    CASE 
                        WHEN TIMESTAMPDIFF(HOUR, sru.last_triggered_at, CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00')) < sr.COOLDOWN_HOURS 
                        THEN 1 
                        ELSE 0 
                    END as is_in_cooldown,
                    GREATEST(0, sr.COOLDOWN_HOURS - TIMESTAMPDIFF(HOUR, sru.last_triggered_at, CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00'))) as remaining_hours
                FROM special_reply_usage sru
                JOIN special_reply sr ON sru.special_reply_id = sr.PKKEY
                WHERE sru.user_id = %s
                ORDER BY sru.last_triggered_at DESC
            ''', (phone_number,))
            
            stats = cursor.fetchall()
            cursor.close()
            return stats
            
        except Exception as e:
            logging.error(f"Error getting user special reply stats: {e}")
            return []
        
    def save_message(self, phone_number, message, role, message_type='text', opening_message_id=None):
        """Save a message to the database with duplicate prevention"""
        if not self.connection or not self.connection.is_connected():
            logging.warning("Database connection lost, attempting to reconnect...")
            self.connection = self.create_connection()
            
        if not self.connection:
            logging.error("No database connection available")
            return False
            
        try:
            cursor = self.connection.cursor()
            
            # DUPLICATE PREVENTION: Check for recent duplicate messages
            cursor.execute('''
                SELECT message_id, timestamp FROM chat_messages 
                WHERE phone_number = %s AND message = %s AND role = %s 
                AND timestamp > DATE_SUB(CONVERT_TZ(NOW(), @@session.time_zone, '+08:00'), INTERVAL 10 MINUTE)
                ORDER BY timestamp DESC
                LIMIT 1
            ''', (phone_number, message, role))
            
            existing = cursor.fetchone()
            if existing:
                # Check if the existing message is very recent (within 2 minutes)
                existing_id, existing_timestamp = existing
                # MODIFY: Use Malaysia time for comparison
                malaysia_now = get_malaysia_time().replace(tzinfo=None)
                time_diff = malaysia_now - existing_timestamp
                
                if time_diff.total_seconds() < 120:  # 2 minutes
                    logging.warning(f"Duplicate message found within 2 minutes, returning existing ID: {existing_id}")
                    cursor.close()
                    return existing_id
                else:
                    logging.info(f"Similar message found but it's {time_diff.total_seconds()} seconds old, allowing new save")


            # Get or create session
            cursor.execute('''
                SELECT session_id FROM chat_sessions 
                WHERE phone_number = %s AND session_end IS NULL 
                ORDER BY session_start DESC LIMIT 1
            ''', (phone_number,))
            session = cursor.fetchone()
            
            if not session:
                cursor.execute('''
                    INSERT INTO chat_sessions (phone_number) 
                    VALUES (%s)
                ''', (phone_number,))
                session_id = cursor.lastrowid
                logging.info(f"Created new session {session_id} for {phone_number}")
            else:
                session_id = session[0]
            
            # INSERT with only essential columns
            cursor.execute('''
                INSERT INTO chat_messages 
                (session_id, phone_number, message, role, message_type, opening_message_id, processed, timestamp)
                VALUES (%s, %s, %s, %s, %s, %s, %s, CONVERT_TZ(NOW(), @@session.time_zone, '+08:00'))
            ''', (session_id, phone_number, message, role, message_type, opening_message_id, False))
            
            message_id = cursor.lastrowid
            
            # Update session activity
            if role == 'user':
                # Increment both total message count and new message count for user messages
                cursor.execute('''
                    UPDATE chat_sessions 
                    SET message_count = message_count + 1, 
                        new_message_count = new_message_count + 1,
                        last_activity = CONVERT_TZ(NOW(), @@session.time_zone, '+08:00')
                    WHERE session_id = %s
                ''', (session_id,))
            else:
                # Only increment total message count for assistant messages
                cursor.execute('''
                    UPDATE chat_sessions 
                    SET message_count = message_count + 1, 
                        last_activity = CONVERT_TZ(NOW(), @@session.time_zone, '+08:00')
                    WHERE session_id = %s
                ''', (session_id,))
            
            logging.info(f"Successfully saved message with ID: {message_id}")
            
            self.connection.commit()
            cursor.close()
            return message_id
            
        except Error as e:
            logging.error(f"Error saving message: {e}")
            return False
    
    def mark_message_processed(self, message_id):
        """Mark a message as processed after AI responds (including errors)"""
        if not self.connection or not message_id:
            return False
            
        try:
            cursor = self.connection.cursor()
            cursor.execute('''
                UPDATE chat_messages 
                SET processed = TRUE 
                WHERE message_id = %s
            ''', (message_id,))
            
            self.connection.commit()
            cursor.close()
            
            logging.info(f"Marked message {message_id} as processed")
            return True
            
        except Error as e:
            logging.error(f"Error marking message as processed: {e}")
            return False
    
    def get_unprocessed_messages(self, phone_number=None):
        """Get unprocessed user messages (for restart recovery)"""
        if not self.connection:
            return []
            
        try:
            cursor = self.connection.cursor(dictionary=True)
            
            if phone_number:
                cursor.execute('''
                    SELECT message_id, phone_number, message, message_type
                    FROM chat_messages 
                    WHERE phone_number = %s AND role = 'user' AND processed = FALSE
                    ORDER BY timestamp ASC
                ''', (phone_number,))
            else:
                cursor.execute('''
                    SELECT message_id, phone_number, message, message_type
                    FROM chat_messages 
                    WHERE role = 'user' AND processed = FALSE
                    ORDER BY timestamp ASC
                ''')
            
            messages = cursor.fetchall()
            cursor.close()
            
            logging.info(f"Found {len(messages)} unprocessed messages")
            return messages
            
        except Error as e:
            logging.error(f"Error getting unprocessed messages: {e}")
            return []
    
    def save_phone_number_to_followup_keep(self, phone_number):
        """
        Save phone number to chat_followup_keep table if it doesn't already exist
        Returns: True if saved (new entry), False if already exists, None if error
        """
        if not self.connection:
            logging.error("No database connection available for followup keep")
            return None
            
        try:
            cursor = self.connection.cursor()
            
            # Check if phone number already exists
            cursor.execute('''
                SELECT id FROM chat_followup_keep 
                WHERE phone_number = %s
            ''', (phone_number,))
            
            existing_record = cursor.fetchone()
            
            if existing_record:
                logging.info(f"Phone number {phone_number} already exists in chat_followup_keep")
                cursor.close()
                return False
            
            # FIXED: Let the table DEFAULT handle the timestamp
            cursor.execute('''
                INSERT INTO chat_followup_keep (phone_number, created_datetime)
                VALUES (%s, NOW())
            ''', (phone_number,))
            
            self.connection.commit()
            cursor.close()
            
            logging.info(f"Phone number {phone_number} saved to chat_followup_keep table with Malaysia time")
            return True
            
        except mysql.connector.IntegrityError as ie:
            if "Duplicate entry" in str(ie):
                logging.info(f"Phone number {phone_number} already exists (race condition)")
                return False
            else:
                logging.error(f"Integrity error saving to chat_followup_keep: {ie}")
                return None
                
        except Exception as e:
            logging.error(f"Error saving phone number to chat_followup_keep: {e}")
            return None
    
    def get_new_message_count(self, phone_number):
        """Get the current new_message_count for a specific user"""
        if not self.connection:
            return 0
            
        try:
            cursor = self.connection.cursor()
            cursor.execute('''
                SELECT new_message_count
                FROM chat_sessions 
                WHERE phone_number = %s 
                ORDER BY session_start DESC 
                LIMIT 1
            ''', (phone_number,))
            
            result = cursor.fetchone()
            cursor.close()
            
            if result:
                return result[0] or 0
            else:
                return 0
                
        except Error as e:
            logging.error(f"Error getting new message count: {e}")
            return 0

    def reset_new_message_count(self, phone_number):
        """Reset the new_message_count to 0 for a specific user"""
        if not self.connection:
            return False
            
        try:
            cursor = self.connection.cursor()
            cursor.execute('''
                UPDATE chat_sessions 
                SET new_message_count = 0,
                    last_activity = CONVERT_TZ(NOW(), @@session.time_zone, '+08:00')
                WHERE phone_number = %s
            ''', (phone_number,))
            
            self.connection.commit()
            cursor.close()
            
            logging.info(f"Reset new_message_count for {phone_number}")
            return True
            
        except Error as e:
            logging.error(f"Error resetting new message count: {e}")
            return False

    def get_users_with_new_messages(self, min_count=1):
        """Get list of users who have new messages above a certain threshold"""
        if not self.connection:
            return []
            
        try:
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute('''
                SELECT phone_number, new_message_count, last_activity,
                       CONVERT_TZ(last_activity, '+00:00', '+08:00') as malaysia_last_activity
                FROM chat_sessions 
                WHERE new_message_count >= %s
                ORDER BY new_message_count DESC, last_activity DESC
            ''', (min_count,))
            
            users = cursor.fetchall()
            cursor.close()
            return users
            
        except Error as e:
            logging.error(f"Error getting users with new messages: {e}")
            return []

    def update_session_flags(self, phone_number, manual_reply=None, appointment_section=None, 
                            human_section=None, human_reply=None, reset_new_count=False):
        """Enhanced update_session_flags method with option to reset new_message_count"""
        if not self.connection:
            return False
            
        try:
            cursor = self.connection.cursor()
            
            # Build dynamic update query
            updates = []
            params = []
            
            if manual_reply is not None:
                updates.append("manual_reply = %s")
                params.append(manual_reply)
            
            if appointment_section is not None:
                updates.append("appointment_section = %s")
                params.append(appointment_section)
            
            if human_section is not None:
                updates.append("human_section = %s")
                params.append(human_section)
            
            if human_reply is not None:
                updates.append("human_reply = %s")
                params.append(human_reply)
                
            if reset_new_count:
                updates.append("new_message_count = 0")
            
            if not updates:
                return True
            
            # Always update last_activity
            updates.append("last_activity = CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00')")
            params.append(phone_number)
            
            query = f'''
                UPDATE chat_sessions 
                SET {", ".join(updates)}
                WHERE phone_number = %s
            '''
            
            cursor.execute(query, params)
            self.connection.commit()
            cursor.close()
            
            reset_msg = " and new_message_count reset" if reset_new_count else ""
            logging.info(f"Updated session flags for {phone_number}: "
                    f"human_reply={human_reply}, manual_reply={manual_reply}, "
                    f"appointment_section={appointment_section}, human_section={human_section}{reset_msg}")
            return True
            
        except Error as e:
            logging.error(f"Error updating session flags: {e}")
            return False
        
    def get_user_chat_history(self, phone_number, limit=100):
        """Get chat history for a specific user"""
        if not self.connection:
            return []
            
        try:
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute('''
                SELECT m.message, m.role, 
                       CONVERT_TZ(m.timestamp, '+00:00', '+08:00') as malaysia_timestamp, 
                       m.message_type 
                FROM chat_messages m
                WHERE m.phone_number = %s 
                ORDER BY m.timestamp DESC 
                LIMIT %s
            ''', (phone_number, limit))
            
            messages = cursor.fetchall()
            cursor.close()
            return messages
            
        except Error as e:
            logging.error(f"Error getting chat history: {e}")
            return []

    def get_all_users(self):
        """Get list of all users who have chatted"""
        if not self.connection:
            return []
            
        try:
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute('''
                SELECT phone_number, 
                       COUNT(*) as message_count,
                       CONVERT_TZ(MIN(timestamp), '+00:00', '+08:00') as first_message,
                       CONVERT_TZ(MAX(timestamp), '+00:00', '+08:00') as last_message
                FROM chat_messages 
                GROUP BY phone_number 
                ORDER BY MAX(timestamp) DESC
            ''')
            
            users = cursor.fetchall()
            cursor.close()
            return users
            
        except Error as e:
            logging.error(f"Error getting users: {e}")
            return []

    def export_chat_history(self, phone_number=None, format='json'):
        """Export chat history to file"""
        if not self.connection:
            return None
            
        try:
            filename_prefix = "chat_history"
            if phone_number:
                query = '''
                    SELECT phone_number, message, role, timestamp, message_type
                    FROM chat_messages 
                    WHERE phone_number = %s
                    ORDER BY timestamp
                '''
                params = (phone_number,)
                filename = f"{filename_prefix}_{phone_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            else:
                query = '''
                    SELECT phone_number, message, role, timestamp, message_type
                    FROM chat_messages 
                    ORDER BY phone_number, timestamp
                '''
                params = ()
                filename = f"{filename_prefix}_all_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute(query, params)
            messages = cursor.fetchall()
            
            if format == 'json':
                with open(f"{filename}.json", 'w', encoding='utf-8') as f:
                    json.dump(messages, f, indent=2, ensure_ascii=False)
                return f"{filename}.json"
                
            elif format == 'csv':
                df = pd.DataFrame(messages)
                df.to_csv(f"{filename}.csv", index=False)
                return f"{filename}.csv"
                
            return None
            
        except Error as e:
            logging.error(f"Error exporting chat history: {e}")
            return None

    def get_chat_statistics(self):
        """Get overall chat statistics"""
        if not self.connection:
            return {}
            
        try:
            cursor = self.connection.cursor()
            
            # Total messages
            cursor.execute('SELECT COUNT(*) FROM chat_messages')
            total_messages = cursor.fetchone()[0]
            
            # Unique users
            cursor.execute('SELECT COUNT(DISTINCT phone_number) FROM chat_messages')
            unique_users = cursor.fetchone()[0]
            
            # Messages today
            cursor.execute('''
                SELECT COUNT(*) FROM chat_messages 
                WHERE DATE(timestamp) = DATE(CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00'))
            ''')
            messages_today = cursor.fetchone()[0]
            
            # Most active user
            cursor.execute('''
                SELECT phone_number, COUNT(*) as count
                FROM chat_messages 
                GROUP BY phone_number 
                ORDER BY count DESC 
                LIMIT 1
            ''')
            most_active = cursor.fetchone()
            
            # Media statistics
            cursor.execute('''
                SELECT message_type, COUNT(*) as count
                FROM chat_messages 
                WHERE message_type != 'text'
                GROUP BY message_type
            ''')
            media_stats = dict(cursor.fetchall())
            
            cursor.close()
            
            return {
                'total_messages': total_messages,
                'unique_users': unique_users,
                'messages_today': messages_today,
                'most_active_user': {
                    'phone_number': most_active[0] if most_active else None,
                    'message_count': most_active[1] if most_active else 0
                },
                'media_statistics': media_stats
            }
            
        except Error as e:
            logging.error(f"Error getting statistics: {e}")
            return {}

    def search_messages(self, search_term, phone_number=None):
        """Search for messages containing specific terms"""
        if not self.connection:
            return []
            
        try:
            cursor = self.connection.cursor(dictionary=True)
            
            if phone_number:
                cursor.execute('''
                    SELECT phone_number, message, role, timestamp
                    FROM chat_messages 
                    WHERE phone_number = %s AND message LIKE %s
                    ORDER BY timestamp DESC
                ''', (phone_number, f'%{search_term}%'))
            else:
                cursor.execute('''
                    SELECT phone_number, message, role, timestamp
                    FROM chat_messages 
                    WHERE message LIKE %s
                    ORDER BY timestamp DESC
                ''', (f'%{search_term}%',))
            
            results = cursor.fetchall()
            cursor.close()
            return results
            
        except Error as e:
            logging.error(f"Error searching messages: {e}")
            return []

    def get_frame_content(self):
        """Get frame content with updated structure for single SENSITIVE_CONTENT and SENSITIVE_REPLY"""
        if not self.connection:
            return None
            
        try:
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute('''
                SELECT FRAME_CONTENT, 
                       SENSITIVE_CONTENT,
                       SENSITIVE_REPLY,
                       ADDITIONAL_CONTENT,
                       APPOINTMENT_REPLY,
                       HUMAN_SERVICE_REPLY,
                       ATTACH_MEDIA_REPLY
                FROM chatbox_frame 
                WHERE PKKEY=1
            ''')
            frame_data = cursor.fetchone()
            cursor.close()
            return frame_data
        except Error as e:
            logging.error(f"Error getting frame content: {e}")
            return None

    def close_connection(self):
        """Close the database connection"""
        if self.connection:
            self.connection.close()
            logging.info("Database connection closed")

    def get_additional_content_files(self):
        """Get list of additional content files from database"""
        if not self.connection:
            return []
            
        try:
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute('''
                SELECT id, filename, original_name, file_size, mime_type, created_at
                FROM additional_content_files 
                ORDER BY created_at DESC
            ''')
            files = cursor.fetchall()
            cursor.close()
            return files
        except Error as e:
            logging.error(f"Error getting additional content files: {e}")
            return []

    def get_file_content_by_filename(self, filename):
        """Get file content from database by filename"""
        if not self.connection:
            return None
            
        try:
            cursor = self.connection.cursor()
            cursor.execute('''
                SELECT filename, original_name, mime_type, file_content
                FROM additional_content_files 
                WHERE filename = %s
            ''', (filename,))
            
            result = cursor.fetchone()
            cursor.close()
            
            if result:
                return {
                    'filename': result[0],
                    'original_name': result[1],
                    'mime_type': result[2],
                    'content': result[3]  # This should be BLOB data
                }
            return None
            
        except Error as e:
            logging.error(f"Error getting file content: {e}")
            return None

    def save_media_file(self, media_content, mime_type, phone_number, media_type, original_filename=None):
        """Save media file (image, video, document, audio) to filesystem and return file info"""
        try:
            # Generate unique filename based on content hash
            media_hash = hashlib.sha256(media_content).hexdigest()
            
            # Determine file extension and storage path
            if media_type == 'audio':
                extension_map = {
                    'audio/ogg': '.ogg',
                    'audio/mpeg': '.mp3',
                    'audio/mp4': '.m4a',
                    'audio/amr': '.amr',
                    'audio/wav': '.wav'
                }
                storage_path = self.audio_storage_path
            elif media_type == 'image':
                extension_map = {
                    'image/jpeg': '.jpg',
                    'image/png': '.png',
                    'image/gif': '.gif',
                    'image/webp': '.webp'
                }
                storage_path = self.image_storage_path
            elif media_type == 'video':
                extension_map = {
                    'video/mp4': '.mp4',
                    'video/avi': '.avi',
                    'video/mov': '.mov',
                    'video/quicktime': '.mov'
                }
                storage_path = self.video_storage_path
            elif media_type == 'document':
                extension_map = {
                    'application/pdf': '.pdf',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
                    'application/msword': '.doc',
                    'text/plain': '.txt',
                    'application/vnd.ms-excel': '.xls',
                    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx'
                }
                storage_path = self.document_storage_path
            else:
                raise ValueError(f"Unsupported media type: {media_type}")
            
            file_extension = extension_map.get(mime_type, '.bin')
            
            # Create organized directory structure: media_type/YYYY/MM/DD/
            malaysia_now = get_malaysia_time()
            date_path = malaysia_now.strftime("%Y/%m/%d")
            full_directory = storage_path / date_path
            full_directory.mkdir(parents=True, exist_ok=True)
            
            # Generate filename: hash_phoneNumber_timestamp.ext
            timestamp = malaysia_now.strftime("%H%M%S")
            filename = f"{media_hash[:12]}_{phone_number}_{timestamp}{file_extension}"
            file_path = full_directory / filename
            
            # Check if file already exists (deduplication)
            if file_path.exists():
                logging.info(f"{media_type.title()} file already exists: {file_path}")
                return str(file_path.relative_to(storage_path)), filename, media_hash, len(media_content)
            
            # Save file to disk
            with open(file_path, 'wb') as f:
                f.write(media_content)
            
            relative_path = str(file_path.relative_to(storage_path))
            
            logging.info(f"{media_type.title()} file saved: {relative_path}")
            return relative_path, filename, media_hash, len(media_content)
            
        except Exception as e:
            logging.error(f"Error saving {media_type} file: {e}")
            return None, None, None, None

    def save_message_with_media(self, phone_number, message, role, message_type='text', 
                           media_file_path=None, media_file_name=None, media_mime_type=None,
                           media_file_size=None, media_hash=None, extracted_text=None,
                           audio_duration=None):
        """Enhanced message saving with support for all media types (fallback method)"""
        if not self.connection:
            logging.error("No database connection available")
            return False
            
        try:
            cursor = self.connection.cursor()
            
            # Get or create session
            cursor.execute('''
                SELECT session_id FROM chat_sessions 
                WHERE phone_number = %s AND session_end IS NULL 
                ORDER BY session_start DESC LIMIT 1
            ''', (phone_number,))
            session = cursor.fetchone()
            
            if not session:
                cursor.execute('''
                    INSERT INTO chat_sessions (phone_number) 
                    VALUES (%s)
                ''', (phone_number,))
                session_id = cursor.lastrowid
            else:
                session_id = session[0]
            
            # Prepare columns and values based on message type
            base_columns = ['session_id', 'phone_number', 'message', 'role', 'message_type']
            base_values = [session_id, phone_number, message, role, message_type]
            
            # Add media-specific columns
            if message_type == 'audio':
                additional_columns = ['audio_file_path', 'audio_file_name', 'audio_mime_type', 
                                    'audio_file_size', 'audio_hash', 'transcribed_text', 'audio_duration']
                additional_values = [media_file_path, media_file_name, media_mime_type, 
                                media_file_size, media_hash, extracted_text, audio_duration]
            elif message_type == 'image':
                additional_columns = ['image_file_path', 'image_file_name', 'image_mime_type', 
                                    'image_file_size', 'image_hash']
                additional_values = [media_file_path, media_file_name, media_mime_type, 
                                media_file_size, media_hash]
            elif message_type == 'video':
                additional_columns = ['video_file_path', 'video_file_name', 'video_mime_type', 
                                    'video_file_size', 'video_hash']
                additional_values = [media_file_path, media_file_name, media_mime_type, 
                                media_file_size, media_hash]
            elif message_type == 'document':
                additional_columns = ['document_file_path', 'document_file_name', 'document_mime_type', 
                                    'document_file_size', 'document_hash', 'document_extracted_text']
                additional_values = [media_file_path, media_file_name, media_mime_type, 
                                media_file_size, media_hash, extracted_text]
            else:
                additional_columns = []
                additional_values = []
            
            all_columns = base_columns + additional_columns
            all_values = base_values + additional_values
            
            # Build query
            placeholders = ', '.join(['%s'] * len(all_columns))
            columns_str = ', '.join(all_columns)
            
            cursor.execute(f'''
                INSERT INTO chat_messages ({columns_str})
                VALUES ({placeholders})
            ''', all_values)
            
            message_id = cursor.lastrowid
            
            # Update session activity
            cursor.execute('''
                UPDATE chat_sessions 
                SET message_count = message_count + 1, 
                    last_activity = CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00')
                WHERE session_id = %s
            ''', (session_id,))
            
            self.connection.commit()
            cursor.close()
            return message_id
            
        except Error as e:
            logging.error(f"Error saving message with media: {e}")
            return False
        
    def save_message_with_media_and_caption(self, phone_number, message, role, message_type='text', 
                       media_file_path=None, media_file_name=None, media_mime_type=None,
                       media_file_size=None, media_hash=None, extracted_text=None,
                       audio_duration=None, media_caption=None):
        """Enhanced message saving with caption support for all media types"""
        if not self.connection:
            logging.error("No database connection available")
            return False
            
        try:
            cursor = self.connection.cursor()
            
            # Get or create session
            cursor.execute('''
                SELECT session_id FROM chat_sessions 
                WHERE phone_number = %s AND session_end IS NULL 
                ORDER BY session_start DESC LIMIT 1
            ''', (phone_number,))
            session = cursor.fetchone()
            
            if not session:
                cursor.execute('''
                    INSERT INTO chat_sessions (phone_number) 
                    VALUES (%s)
                ''', (phone_number,))
                session_id = cursor.lastrowid
            else:
                session_id = session[0]
            
            # Check if media_caption column exists before using it
            cursor.execute('''
                SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS 
                WHERE TABLE_SCHEMA = %s AND TABLE_NAME = 'chat_messages' AND COLUMN_NAME = 'media_caption'
            ''', (os.getenv("MYSQL_DATABASE", "railway"),))
            
            has_caption_column = cursor.fetchone() is not None
            
            # Prepare columns and values based on message type
            base_columns = ['session_id', 'phone_number', 'message', 'role', 'message_type']
            base_values = [session_id, phone_number, message, role, message_type]
            
            # Add media-specific columns
            if message_type == 'audio':
                additional_columns = ['audio_file_path', 'audio_file_name', 'audio_mime_type', 
                                    'audio_file_size', 'audio_hash', 'transcribed_text', 'audio_duration']
                additional_values = [media_file_path, media_file_name, media_mime_type, 
                                media_file_size, media_hash, extracted_text, audio_duration]
            elif message_type == 'image':
                additional_columns = ['image_file_path', 'image_file_name', 'image_mime_type', 
                                    'image_file_size', 'image_hash']
                additional_values = [media_file_path, media_file_name, media_mime_type, 
                                media_file_size, media_hash]
                # Add caption column only if it exists
                if has_caption_column and media_caption is not None:
                    additional_columns.append('media_caption')
                    additional_values.append(media_caption)
                    
            elif message_type == 'video':
                additional_columns = ['video_file_path', 'video_file_name', 'video_mime_type', 
                                    'video_file_size', 'video_hash']
                additional_values = [media_file_path, media_file_name, media_mime_type, 
                                media_file_size, media_hash]
                # Add caption column only if it exists
                if has_caption_column and media_caption is not None:
                    additional_columns.append('media_caption')
                    additional_values.append(media_caption)
                    
            elif message_type == 'document':
                additional_columns = ['document_file_path', 'document_file_name', 'document_mime_type', 
                                    'document_file_size', 'document_hash', 'document_extracted_text']
                additional_values = [media_file_path, media_file_name, media_mime_type, 
                                media_file_size, media_hash, extracted_text]
                # Add caption column only if it exists
                if has_caption_column and media_caption is not None:
                    additional_columns.append('media_caption')
                    additional_values.append(media_caption)
            else:
                additional_columns = []
                additional_values = []
            
            all_columns = base_columns + additional_columns
            all_values = base_values + additional_values
            
            # Build query
            placeholders = ', '.join(['%s'] * len(all_columns))
            columns_str = ', '.join(all_columns)
            
            cursor.execute(f'''
                INSERT INTO chat_messages ({columns_str})
                VALUES ({placeholders})
            ''', all_values)
            
            message_id = cursor.lastrowid
            
            # Update session activity
            cursor.execute('''
                UPDATE chat_sessions 
                SET message_count = message_count + 1, 
                    last_activity = CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00')
                WHERE session_id = %s
            ''', (session_id,))
            
            self.connection.commit()
            cursor.close()
            
            logging.info(f"Successfully saved {message_type} message with ID {message_id}")
            return message_id
            
        except Error as e:
            logging.error(f"Error saving message with media and caption: {e}")
            import traceback
            logging.error(f"Full traceback: {traceback.format_exc()}")
            return False

    def get_media_messages(self, phone_number=None, media_type=None, limit=50):
        """Get media messages (images, videos, documents, audio)"""
        if not self.connection:
            return []
            
        try:
            cursor = self.connection.cursor(dictionary=True)
            
            # Build query based on parameters
            where_conditions = []
            params = []
            
            if phone_number:
                where_conditions.append("phone_number = %s")
                params.append(phone_number)
            
            if media_type:
                where_conditions.append("message_type = %s")
                params.append(media_type)
            else:
                where_conditions.append("message_type IN ('image', 'video', 'document', 'audio')")
            
            where_clause = " AND ".join(where_conditions) if where_conditions else "1=1"
            params.append(limit)
            
            cursor.execute(f'''
                SELECT * FROM chat_messages 
                WHERE {where_clause}
                ORDER BY timestamp DESC 
                LIMIT %s
            ''', params)
            
            messages = cursor.fetchall()
            cursor.close()
            
            # Add media URLs to each message
            for message in messages:
                msg_type = message['message_type']
                if msg_type == 'audio' and message['audio_file_path']:
                    message['media_url'] = f"/api/audio/{message['audio_file_path']}"
                elif msg_type == 'image' and message['image_file_path']:
                    message['media_url'] = f"/api/images/{message['image_file_path']}"
                elif msg_type == 'video' and message['video_file_path']:
                    message['media_url'] = f"/api/videos/{message['video_file_path']}"
                elif msg_type == 'document' and message['document_file_path']:
                    message['media_url'] = f"/api/documents/{message['document_file_path']}"
            
            return messages
            
        except Error as e:
            logging.error(f"Error getting media messages: {e}")
            return []

    def set_opening_completed_by_phone(self, phone_number: str, completed: int = 1):
        """直接把最新一条该号码的会话的 opening_sequence_completed 置为 completed"""
        if not self.connection:
            return False
        try:
            cursor = self.connection.cursor()
            cursor.execute("""
                UPDATE chat_sessions
                SET opening_sequence_completed=%s,
                    last_activity=CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00')
                WHERE phone_number=%s
                ORDER BY session_start DESC
                LIMIT 1
            """, (completed, phone_number))
            self.connection.commit()
            cursor.close()
            return True
        except Exception as e:
            logging.error("set_opening_completed_by_phone error: %s", e)
            return False


class FileProcessor:
    """Class to handle PDF, Word document, and Excel processing"""
    
    @staticmethod
    def extract_pdf_text(file_content):
        """Extract text from PDF file content"""
        try:
            pdf_stream = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            text_content = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content += page.extract_text() + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logging.error(f"Error extracting PDF text: {e}")
            return None

    @staticmethod
    def extract_docx_text(file_content):
        """Extract text from Word document file content"""
        try:
            docx_stream = io.BytesIO(file_content)
            doc = docx.Document(docx_stream)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logging.error(f"Error extracting DOCX text: {e}")
            return None

    @staticmethod
    def extract_excel_text(file_content):
        """Extract text from Excel file content"""
        try:
            import pandas as pd
            from io import BytesIO
            
            excel_stream = BytesIO(file_content)
            
            # Try to read all sheets
            excel_file = pd.ExcelFile(excel_stream)
            text_content = ""
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(excel_stream, sheet_name=sheet_name)
                    
                    # Add sheet name as header
                    text_content += f"\n=== Sheet: {sheet_name} ===\n"
                    
                    # Convert DataFrame to string representation
                    # Handle NaN values and convert to string
                    df_cleaned = df.fillna('')  # Replace NaN with empty string
                    
                    # Add column headers
                    text_content += " | ".join(str(col) for col in df_cleaned.columns) + "\n"
                    text_content += "-" * 50 + "\n"
                    
                    # Add each row
                    for index, row in df_cleaned.iterrows():
                        row_text = " | ".join(str(value) for value in row.values)
                        text_content += row_text + "\n"
                    
                    text_content += "\n"
                    
                except Exception as sheet_error:
                    logging.warning(f"Error reading sheet '{sheet_name}': {sheet_error}")
                    text_content += f"Error reading sheet '{sheet_name}'\n"
            
            return text_content.strip() if text_content.strip() else None
            
        except Exception as e:
            logging.error(f"Error extracting Excel text: {e}")
            return None

    @staticmethod
    def extract_csv_text(file_content):
        """Extract text from CSV file content"""
        try:
            import pandas as pd
            from io import StringIO
            
            # Decode bytes to string
            csv_string = file_content.decode('utf-8', errors='ignore')
            csv_stream = StringIO(csv_string)
            
            # Read CSV
            df = pd.read_csv(csv_stream)
            
            # Convert to text format
            text_content = ""
            
            # Add column headers
            text_content += " | ".join(str(col) for col in df.columns) + "\n"
            text_content += "-" * 50 + "\n"
            
            # Add each row (limit to first 100 rows for large files)
            max_rows = min(100, len(df))
            for index, row in df.head(max_rows).iterrows():
                row_text = " | ".join(str(value) for value in row.values)
                text_content += row_text + "\n"
            
            if len(df) > max_rows:
                text_content += f"\n... (showing first {max_rows} rows of {len(df)} total rows)\n"
            
            return text_content.strip()
            
        except Exception as e:
            logging.error(f"Error extracting CSV text: {e}")
            return None

    @staticmethod
    def extract_text_from_file(file_content, mime_type):
        """Extract text from file based on MIME type"""
        if mime_type == 'application/pdf':
            return FileProcessor.extract_pdf_text(file_content)
        elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                          'application/msword']:
            return FileProcessor.extract_docx_text(file_content)
        elif mime_type in ['application/vnd.ms-excel', 
                          'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            return FileProcessor.extract_excel_text(file_content)
        elif mime_type == 'text/csv':
            return FileProcessor.extract_csv_text(file_content)
        elif mime_type == 'text/plain':
            try:
                return file_content.decode('utf-8')
            except:
                return file_content.decode('utf-8', errors='ignore')
        else:
            logging.warning(f"Unsupported file type: {mime_type}")
            return None

class MediaHandler:
    """Class to handle media files (images and videos)"""
    
    def __init__(self, media_base_path=None, media_base_url=None):
        # Set the correct path based on the type of media
        if media_base_url:
            self.media_base_url = media_base_url
        else:
            # Default to opening media URL
            self.media_base_url = "https://reactjs-appointmentwhatsapp-production.up.railway.app/media_opening"
        
        # For local development fallback
        if media_base_path:
            self.media_base_path = Path(media_base_path)
        else:
            # Local fallback directory
            self.media_base_path = Path("media_opening")
        
        # Ensure local media directory exists (for local development)
        self.media_base_path.mkdir(parents=True, exist_ok=True)
        
        # Log the media configuration
        logging.info(f"Media base path: {self.media_base_path.absolute()}")
        logging.info(f"Media base URL: {self.media_base_url}")

    def get_media_file_path(self, filename, media_type=None):
        """Get the full path to a media file"""
        return self.media_base_path / filename
    
    def file_exists(self, filename, media_type=None):
        """FIXED: Check both local filesystem AND remote URL availability"""
        try:
            # First, try to find the file locally
            local_path = self.get_media_file_path(filename, media_type)
            if local_path and local_path.exists():
                logging.info(f"File check for {filename}: EXISTS LOCALLY at {local_path}")
                return True
            
            # If not found locally, check remote URL availability
            if not hasattr(self, '_file_cache'):
                self._file_cache = {}
            
            cache_key = f"{filename}:{media_type}"
            if cache_key in self._file_cache:
                cached_result, cache_time = self._file_cache[cache_key]
                if time.time() - cache_time < 300:  # 5 minutes cache
                    if cached_result:
                        logging.info(f"File check for {filename}: EXISTS REMOTELY (cached)")
                    else:
                        logging.info(f"File check for {filename}: NOT FOUND (cached)")
                    return cached_result
            
            # Check remote URL with shorter timeout
            media_url = self.get_media_url(filename)
            response = requests.head(media_url, timeout=5)
            exists = response.status_code == 200
            
            # Cache the result
            self._file_cache[cache_key] = (exists, time.time())
            
            if exists:
                logging.info(f"File check for {filename}: EXISTS REMOTELY via URL")
            else:
                logging.info(f"File check for {filename}: NOT FOUND")
            return exists
            
        except requests.exceptions.Timeout:
            logging.warning(f"Timeout checking file existence for {filename}")
            return False
        except Exception as e:
            logging.error(f"Error checking file existence for {filename}: {e}")
            return False
    
    def get_media_url(self, filename, media_type=None):
        """Generate public URL for media file access"""
        media_url = f"{self.media_base_url}/{filename}"
        logging.info(f"Generated media URL: {media_url}")
        return media_url

class AutoGitHubMediaManager:
    """Enhanced media manager with automatic git operations"""
    
    def __init__(self, base_path="media"):
        self.base_path = Path(base_path)
        self.base_path.mkdir(parents=True, exist_ok=True)
        
        # Initialize automatic git manager
        self.auto_git = AutoGitManager()
        
        logging.info("AutoGitHubMediaManager initialized")
    
    def save_media_with_auto_git(self, media_content, mime_type, phone_number, media_type, original_filename=None):
        """Save media file and automatically commit to git"""
        try:
            # Save file locally first
            file_path, filename, file_hash, file_size = self._save_media_file_local(
                media_content, mime_type, phone_number, media_type, original_filename
            )
            
            if file_path and filename:
                # Add to git queue for automatic processing
                full_file_path = self.base_path / file_path
                commit_message = f"Add {media_type} file from {phone_number}: {filename}"
                
                self.auto_git.add_file_to_queue(str(full_file_path), commit_message)
                
                logging.info(f"Media file saved and queued for git: {filename}")
                return file_path, filename, file_hash, file_size
            
            return None, None, None, None
            
        except Exception as e:
            logging.error(f"Error saving media with auto git: {e}")
            return None, None, None, None
    
    def _save_media_file_local(self, media_content, mime_type, phone_number, media_type, original_filename=None):
        """Save media file locally with proper folder structure"""
        try:
            import hashlib
            from datetime import datetime
            
            # Generate unique filename based on content hash
            media_hash = hashlib.sha256(media_content).hexdigest()
            
            # Determine file extension and storage path
            extension_map = {
                'audio': {
                    'audio/ogg': '.ogg',
                    'audio/mpeg': '.mp3',
                    'audio/mp4': '.m4a',
                    'audio/amr': '.amr',
                    'audio/wav': '.wav'
                },
                'image': {
                    'image/jpeg': '.jpg',
                    'image/png': '.png',
                    'image/gif': '.gif',
                    'image/webp': '.webp'
                },
                'video': {
                    'video/mp4': '.mp4',
                    'video/avi': '.avi',
                    'video/mov': '.mov',
                    'video/quicktime': '.mov'
                },
                'document': {
                    'application/pdf': '.pdf',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
                    'application/msword': '.doc',
                    'text/plain': '.txt',
                    'application/vnd.ms-excel': '.xls',
                    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx'
                }
            }
            
            if media_type not in extension_map:
                raise ValueError(f"Unsupported media type: {media_type}")
            
            file_extension = extension_map[media_type].get(mime_type, '.bin')
            storage_path = self.base_path / media_type
            
            # Create organized directory structure
            date_path = datetime.now().strftime("%Y/%m/%d")
            full_directory = storage_path / date_path
            full_directory.mkdir(parents=True, exist_ok=True)
            
            # Generate filename
            timestamp = datetime.now().strftime("%H%M%S")
            filename = f"{media_hash[:12]}_{phone_number}_{timestamp}{file_extension}"
            file_path = full_directory / filename
            
            # Check if file already exists (deduplication)
            if file_path.exists():
                logging.info(f"{media_type.title()} file already exists: {file_path}")
                return str(file_path.relative_to(self.base_path)), filename, media_hash, len(media_content)
            
            # Save file to disk
            with open(file_path, 'wb') as f:
                f.write(media_content)
            
            relative_path = str(file_path.relative_to(self.base_path))
            
            logging.info(f"{media_type.title()} file saved locally: {relative_path}")
            return relative_path, filename, media_hash, len(media_content)
            
        except Exception as e:
            logging.error(f"Error saving {media_type} file locally: {e}")
            return None, None, None, None


# === NEW: minimal GitHub Contents API client ===
class GitHubRepoClient:
    def __init__(self, owner, repo, token, branch="master"):
        self.owner = owner
        self.repo = repo
        self.branch = branch
        self.base = f"https://api.github.com/repos/{owner}/{repo}/contents"
        self.session = requests.Session()
        self.session.headers.update({
            "Authorization": f"token {token}",
            "Accept": "application/vnd.github+json",
            "User-Agent": "python-chatbot-uploader"
        })

    def get(self, path):
        r = self.session.get(f"{self.base}/{path}", params={"ref": self.branch}, timeout=30)
        return r.json() if r.status_code == 200 else None

    def put(self, path, content_bytes, message, sha=None):
        payload = {
            "message": message,
            "branch": self.branch,
            "content": base64.b64encode(content_bytes).decode("ascii")
        }
        if sha:
            payload["sha"] = sha
        r = self.session.put(f"{self.base}/{path}", json=payload, timeout=30)
        if r.status_code not in (200, 201):
            raise RuntimeError(f"GitHub PUT failed: {r.status_code} {r.text}")
        return r.json()


# === REPLACE: SpecialReplyGitHubMediaManager ===
class SpecialReplyGitHubMediaManager:
    """
    Uploads outgoing special-reply media to the *reactjs-appointmentwhatsapp* repo
    under: public/send_media/<image|video>/YYYY/MM/DD/<filename>
    """

    def __init__(self, base_path="public/send_media"):
        from pathlib import Path
        self.base_path = Path(base_path)
        self.base_path.mkdir(parents=True, exist_ok=True)

        self.github_token = os.getenv("GITHUB_TOKEN")
        self.frontend_owner = os.getenv("FRONTEND_REPO_OWNER", "Welsh510")
        self.frontend_repo  = os.getenv("FRONTEND_REPO_NAME",  "reactjs-appointmentwhatsapp")
        self.frontend_branch = os.getenv("FRONTEND_REPO_BRANCH", "master")

        if not self.github_token:
            logging.error("GITHUB_TOKEN not found – cannot upload to GitHub.")
        self.client = GitHubRepoClient(
            owner=self.frontend_owner,
            repo=self.frontend_repo,
            token=self.github_token,
            branch=self.frontend_branch
        )

    def copy_and_upload_media_to_private_repo(self, source_media_path, phone_number, media_type, original_filename=None):
        """
        1) Read local file bytes (this file may have been downloaded seconds before).
        2) Create organized destination path in the React repo.
        3) PUT to GitHub Contents API (create or update).
        Returns: (send_media_path, send_filename, file_hash, file_size)
        """
        from datetime import datetime
        import hashlib

        if not source_media_path or not os.path.exists(source_media_path):
            raise FileNotFoundError(f"Source media not found: {source_media_path}")

        with open(source_media_path, "rb") as f:
            content = f.read()

        now = datetime.now()
        year  = now.strftime("%Y")
        month = now.strftime("%m")
        day   = now.strftime("%d")

        # Decide extension and build a clean filename
        src_name = original_filename or os.path.basename(source_media_path)
        stem, ext = os.path.splitext(src_name)
        if not ext:
            ext = ".jpg" if media_type == "image" else ".mp4"

        file_hash = hashlib.md5(f"{src_name}_{phone_number}_{time.time()}".encode()).hexdigest()
        send_filename = f"{stem}_{file_hash[:8]}{ext}"
        send_media_path = f"{media_type}/{year}/{month}/{day}/{send_filename}"

        github_path = f"public/send_media/{send_media_path}"

        # If file already exists at destination, update; else create
        existing = self.client.get(github_path)
        sha = existing["sha"] if existing and "sha" in existing else None

        msg = f"Auto-upload special_reply media: {send_filename}"
        self.client.put(github_path, content, msg, sha=sha)

        return send_media_path, send_filename, file_hash, len(content)

# Modified ChatHistoryManager class
class EnhancedChatHistoryManager(ChatHistoryManager):
    """Enhanced ChatHistoryManager with manual reply detection"""
    
    def __init__(self):
        super().__init__()
        # Use the automatic git media manager
        self.auto_git_media_manager = AutoGitHubMediaManager("media")
        self.message_analyzer = MessageAnalyzer()
        
        # UPDATED: Use the enhanced special reply GitHub manager for private repo
        self.special_reply_github_manager = SpecialReplyGitHubMediaManager("public/send_media")
        
    def setup_git_config(self):
        """Setup git configuration for private repository"""
        try:
            # Set git config for the project
            subprocess.run(["git", "config", "user.email", "chatbot@railway.app"], 
                         check=True, cwd=".", capture_output=True)
            subprocess.run(["git", "config", "user.name", "Railway ChatBot"], 
                         check=True, cwd=".", capture_output=True)
            
            # For private repositories, ensure we have the right remote configured
            # This will use the GitHub token from environment variables
            github_token = os.getenv("GITHUB_TOKEN")
            if github_token:
                # Update remote URL to use token authentication
                repo_url = f"https://{github_token}@github.com/Welsh510/python-chatbot.git"
                subprocess.run(["git", "remote", "set-url", "origin", repo_url], 
                             cwd=".", capture_output=True)
                logging.info("Git remote configured with token authentication")
            else:
                logging.warning("GITHUB_TOKEN not found in environment variables")
                
        except subprocess.CalledProcessError as e:
            logging.error(f"Error setting up git config: {e}")
        except Exception as e:
            logging.error(f"Unexpected error in git setup: {e}")
    
    def save_media_file(self, media_content, mime_type, phone_number, media_type, original_filename=None):
        """Save media file with automatic git operations"""
        try:
            # Use the automatic git media manager
            result = self.auto_git_media_manager.save_media_with_auto_git(
                media_content, mime_type, phone_number, media_type, original_filename
            )
            
            if result and result[0]:
                logging.info(f"Media file saved with automatic git: {result[1]}")
                return result
            else:
                logging.error(f"Failed to save {media_type} file with automatic git")
                # Fallback to parent class method
                return super().save_media_file(media_content, mime_type, phone_number, media_type, original_filename)
                
        except Exception as e:
            logging.error(f"Error in automatic git media save: {e}")
            # Fallback to parent class method
            return super().save_media_file(media_content, mime_type, phone_number, media_type, original_filename)
      
    def save_sent_special_reply_media_existing_db(self, phone_number, message_text, media_type, 
                                                send_media_info, mime_type=None):
        """Save sent special reply media using existing database columns - FIXED VERSION"""
        if not self.connection or not self.connection.is_connected():
            logging.warning("Database connection lost, attempting to reconnect...")
            self.connection = self.create_connection()
            
        if not self.connection:
            logging.error("No database connection available")
            return False
        
        # DEBUG: Log the send_media_info content
        logging.info(f"DEBUG - send_media_info received: {send_media_info}")
        logging.info(f"DEBUG - media_type: {media_type}")
        logging.info(f"DEBUG - mime_type: {mime_type}")
            
        try:
            cursor = self.connection.cursor()
            
            # Get or create session
            cursor.execute('''
                SELECT session_id FROM chat_sessions 
                WHERE phone_number = %s AND session_end IS NULL 
                ORDER BY session_start DESC LIMIT 1
            ''', (phone_number,))
            session = cursor.fetchone()
            
            if not session:
                cursor.execute('''
                    INSERT INTO chat_sessions (phone_number) 
                    VALUES (%s)
                ''', (phone_number,))
                session_id = cursor.lastrowid
                logging.info(f"Created new session {session_id} for {phone_number}")
            else:
                session_id = session[0]
            
            # Prepare base columns and values
            base_columns = ['session_id', 'phone_number', 'message', 'role', 'message_type', 'processed']
            base_values = [session_id, phone_number, message_text, "assistant", media_type, True]
            
            # FIXED: Add media-specific columns based on media type using existing database structure
            if media_type == 'image' and send_media_info:
                # Extract values with proper error checking
                file_path = send_media_info.get('send_media_path') or send_media_info.get('file_path')
                file_name = send_media_info.get('send_filename') or send_media_info.get('filename')
                file_hash = send_media_info.get('file_hash') or send_media_info.get('hash')
                file_size = send_media_info.get('file_size') or send_media_info.get('size')
                
                logging.info(f"DEBUG - Image data: path={file_path}, name={file_name}, hash={file_hash}, size={file_size}")
                
                additional_columns = ['image_file_path', 'image_file_name', 'image_mime_type', 
                                    'image_file_size', 'image_hash']
                additional_values = [
                    file_path,
                    file_name,
                    mime_type or 'image/jpeg',
                    file_size,
                    file_hash
                ]
                
            elif media_type == 'video' and send_media_info:
                # Extract values with proper error checking
                file_path = send_media_info.get('send_media_path') or send_media_info.get('file_path')
                file_name = send_media_info.get('send_filename') or send_media_info.get('filename')
                file_hash = send_media_info.get('file_hash') or send_media_info.get('hash')
                file_size = send_media_info.get('file_size') or send_media_info.get('size')
                
                logging.info(f"DEBUG - Video data: path={file_path}, name={file_name}, hash={file_hash}, size={file_size}")
                
                additional_columns = ['video_file_path', 'video_file_name', 'video_mime_type', 
                                    'video_file_size', 'video_hash']
                additional_values = [
                    file_path,
                    file_name,
                    mime_type or 'video/mp4',
                    file_size,
                    file_hash
                ]
            else:
                logging.warning(f"No send_media_info or unsupported media_type: {media_type}")
                additional_columns = []
                additional_values = []
            
            all_columns = base_columns + additional_columns
            all_values = base_values + additional_values
            
            # DEBUG: Log the final query data
            logging.info(f"DEBUG - Columns: {all_columns}")
            logging.info(f"DEBUG - Values: {all_values}")
            
            # Build query
            placeholders = ', '.join(['%s'] * len(all_columns))
            columns_str = ', '.join(all_columns)
            
            query = f'INSERT INTO chat_messages ({columns_str}) VALUES ({placeholders})'
            logging.info(f"DEBUG - SQL Query: {query}")
            
            cursor.execute(query, all_values)
            
            message_id = cursor.lastrowid
            
            # Update session activity
            cursor.execute('''
                UPDATE chat_sessions 
                SET message_count = message_count + 1, 
                    last_activity = CURRENT_TIMESTAMP 
                WHERE session_id = %s
            ''', (session_id,))
            
            self.connection.commit()
            cursor.close()
            
            logging.info(f"Successfully saved sent special reply media with ID: {message_id}")
            return message_id
            
        except Error as e:
            logging.error(f"Error saving sent special reply media: {e}")
            import traceback
            logging.error(f"Full traceback: {traceback.format_exc()}")
            return False
    
    def save_sent_special_reply_media_with_admin_role(
        self,
        phone_number: str,
        message_text: str,
        media_type: str,                     # 'image' or 'video'
        send_media_info: dict,               # from send_special_reply_media_with_private_repo_upload
        mime_type: str | None = None,
    ):
        """Save a sent special-reply media message; prefer role='admin' and mark processed=True."""
        if not self.connection or not self.connection.is_connected():
            logging.warning("DB connection lost, attempting reconnect.")
            self.connection = self.create_connection()
        if not self.connection:
            logging.error("No database connection.")
            return False

        try:
            cursor = self.connection.cursor()

            # Get or create session
            cursor.execute("""
                SELECT session_id FROM chat_sessions
                WHERE phone_number = %s AND session_end IS NULL
                ORDER BY session_start DESC LIMIT 1
            """, (phone_number,))
            session = cursor.fetchone()
            if not session:
                cursor.execute("INSERT INTO chat_sessions (phone_number) VALUES (%s)", (phone_number,))
                session_id = cursor.lastrowid
            else:
                session_id = session[0]

            role_to_use = "admin"  # per your requirement

            base_columns = ['session_id','phone_number','message','role','message_type','processed']
            base_values  = [session_id,  phone_number,  message_text, role_to_use, media_type,  True]

            # Map fields coming back from the uploader
            if media_type == 'image' and send_media_info:
                file_path = send_media_info.get('send_media_path') or send_media_info.get('file_path')
                file_name = send_media_info.get('send_filename')  or send_media_info.get('filename')
                file_hash = send_media_info.get('file_hash')      or send_media_info.get('hash')
                file_size = send_media_info.get('file_size')      or send_media_info.get('size')
                additional_columns = ['image_file_path','image_file_name','image_mime_type','image_file_size','image_hash']
                additional_values  = [file_path, file_name, mime_type or 'image/jpeg', file_size, file_hash]

            elif media_type == 'video' and send_media_info:
                file_path = send_media_info.get('send_media_path') or send_media_info.get('file_path')
                file_name = send_media_info.get('send_filename')  or send_media_info.get('filename')
                file_hash = send_media_info.get('file_hash')      or send_media_info.get('hash')
                file_size = send_media_info.get('file_size')      or send_media_info.get('size')
                additional_columns = ['video_file_path','video_file_name','video_mime_type','video_file_size','video_hash']
                additional_values  = [file_path, file_name, mime_type or 'video/mp4', file_size, file_hash]

            else:
                additional_columns = []
                additional_values  = []

            all_columns   = base_columns + additional_columns
            all_values    = base_values  + additional_values
            placeholders  = ', '.join(['%s'] * len(all_columns))
            columns_str   = ', '.join(all_columns)
            insert_sql    = f"INSERT INTO chat_messages ({columns_str}) VALUES ({placeholders})"

            # Try insert; if role enum doesn’t include 'admin', fall back gracefully
            try:
                cursor.execute(insert_sql, all_values)
            except Error as e:
                msg = str(e).lower()
                if "enum" in msg and "role" in msg and "admin" in msg:
                    logging.warning("chat_messages.role enum lacks 'admin'; falling back to 'assistant' for now.")
                    base_values[3] = "assistant"   # role index in base_values
                    all_values = base_values + additional_values
                    cursor.execute(insert_sql, all_values)
                else:
                    raise

            message_id = cursor.lastrowid

            # Update session activity
            cursor.execute("""
                UPDATE chat_sessions
                SET message_count = message_count + 1,
                    last_activity = CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00')
                WHERE session_id = %s
            """, (session_id,))

            self.connection.commit()
            cursor.close()
            logging.info(f"Saved special-reply {media_type} with admin role, ID={message_id}")
            return message_id

        except Error as e:
            logging.error(f"Error saving special-reply media (admin role): {e}")
            import traceback; logging.error(traceback.format_exc())
            return False
            
    def update_session_flags(self, phone_number, manual_reply=None, appointment_section=None, human_section=None, human_reply=None):
        """Update session flags for manual reply system"""
        if not self.connection:
            return False
            
        try:
            cursor = self.connection.cursor()
            
            # Build dynamic update query
            updates = []
            params = []
            
            if manual_reply is not None:
                updates.append("manual_reply = %s")
                params.append(manual_reply)
            
            if appointment_section is not None:
                updates.append("appointment_section = %s")
                params.append(appointment_section)
            
            if human_section is not None:
                updates.append("human_section = %s")
                params.append(human_section)
            
            if human_reply is not None:  # NEW FLAG HANDLING
                updates.append("human_reply = %s")
                params.append(human_reply)
            
            if not updates:
                return True
            
            # Always update last_activity
            updates.append("last_activity = CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00')")
            params.append(phone_number)
            
            query = f'''
                UPDATE chat_sessions 
                SET {", ".join(updates)}
                WHERE phone_number = %s
            '''
            
            cursor.execute(query, params)
            self.connection.commit()
            cursor.close()
            
            logging.info(f"Updated session flags for {phone_number}: "
                    f"human_reply={human_reply}, manual_reply={manual_reply}, "
                    f"appointment_section={appointment_section}, human_section={human_section}")
            return True
            
        except Error as e:
            logging.error(f"Error updating session flags: {e}")
            return False
    
    def get_session_flags(self, phone_number):
        """Get current session flags"""
        if not self.connection:
            return {
                "manual_reply": 0, 
                "appointment_section": 0, 
                "human_section": 0,
                "human_reply": 0  # NEW FLAG
            }
            
        try:
            cursor = self.connection.cursor()
            cursor.execute('''
                SELECT manual_reply, appointment_section, human_section, human_reply  # UPDATED
                FROM chat_sessions 
                WHERE phone_number = %s 
                ORDER BY session_start DESC 
                LIMIT 1
            ''', (phone_number,))
            
            result = cursor.fetchone()
            cursor.close()
            
            if result:
                return {
                    "manual_reply": result[0] or 0,
                    "appointment_section": result[1] or 0,
                    "human_section": result[2] or 0,
                    "human_reply": result[3] or 0  # NEW FLAG
                }
            else:
                return {
                    "manual_reply": 0, 
                    "appointment_section": 0, 
                    "human_section": 0,
                    "human_reply": 0  # NEW FLAG
                }
                
        except Error as e:
            logging.error(f"Error getting session flags: {e}")
            return {
                "manual_reply": 0, 
                "appointment_section": 0, 
                "human_section": 0,
                "human_reply": 0  # NEW FLAG
            }
            
class MessageAnalyzer:
    """Class to analyze messages for appointment and human service requests"""
    
    def __init__(self):
        # Keywords for appointment detection (multilingual)
        self.appointment_keywords = {
            'en': ['appointment', 'schedule', 'book', 'reserve', 'meeting', 'visit', 'consultation', 'booking'],
            'zh': ['预约', '预定', '约定', '安排', '会议', '访问', '咨询', '订约', '约见', '见面', '预订'],
            'ms': ['temujanji', 'jadual', 'tempahan', 'lawatan', 'konsultasi', 'masa', 'bertemu']
        }
        
        # Keywords for human service requests (multilingual)
        self.human_service_keywords = {
            'en': ['human', 'person', 'agent', 'staff', 'customer service', 'support', 'help desk', 'representative', 'operator', 'talk to someone'],
            'zh': ['人工', '客服', '人员', '工作人员', '代表', '服务员', '真人', '人类', '客户服务', '支持', '帮助台', '操作员', '找人'],
            'ms': ['manusia', 'pekerja', 'wakil', 'khidmat pelanggan', 'sokongan', 'bantuan', 'operator', 'staf']
        }
    
    def detect_appointment_request(self, message, detected_language='en'):
        """Detect if message contains appointment request"""
        if not message or not message.strip():
            return False
            
        message_lower = message.lower().strip()
        
        # Check keywords for detected language
        keywords = self.appointment_keywords.get(detected_language, [])
        
        # Also check English keywords as fallback
        if detected_language != 'en':
            keywords.extend(self.appointment_keywords.get('en', []))
        
        for keyword in keywords:
            if keyword in message_lower:
                logging.info(f"Appointment keyword detected: '{keyword}' in message: '{message[:50]}...'")
                return True
        
        return False
    
    def detect_human_service_request(self, message, detected_language='en'):
        """Detect if message contains human service request"""
        if not message or not message.strip():
            return False
            
        message_lower = message.lower().strip()
        
        # Check keywords for detected language
        keywords = self.human_service_keywords.get(detected_language, [])
        
        # Also check English keywords as fallback
        if detected_language != 'en':
            keywords.extend(self.human_service_keywords.get('en', []))
        
        for keyword in keywords:
            if keyword in message_lower:
                logging.info(f"Human service keyword detected: '{keyword}' in message: '{message[:50]}...'")
                return True
        
        return False

        
# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('chatbot.log'),
        logging.StreamHandler()
    ]
)

# Apply Malaysia time formatter to all handlers
for handler in logging.getLogger().handlers:
    handler.setFormatter(MalaysiaTimeFormatter('%(asctime)s - %(levelname)s - %(message)s'))

logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)

# Configure environment variables
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
DEFAULT_OPENAI_MODEL = "gpt-3.5-turbo"

META_WHATSAPP_TOKEN = os.getenv("META_WHATSAPP_TOKEN")
META_WHATSAPP_PHONE_ID = os.getenv("META_WHATSAPP_PHONE_ID")
META_VERIFY_TOKEN = os.getenv("META_VERIFY_TOKEN")
META_API_VERSION = os.getenv("META_API_VERSION", "v18.0")
META_APP_ID = os.getenv("META_APP_ID")
META_APP_SECRET = os.getenv("META_APP_SECRET")

# MySQL configuration
MYSQL_HOST = os.getenv("MYSQL_HOST", "mysql.railway.internal")
MYSQL_USER = os.getenv("MYSQL_USER", "root")
MYSQL_PASSWORD = os.getenv("MYSQL_PASSWORD")
MYSQL_DATABASE = os.getenv("MYSQL_DATABASE", "railway")

# WhatsApp API base URL
META_API_BASE_URL = f"https://graph.facebook.com/{META_API_VERSION}/{META_WHATSAPP_PHONE_ID}/messages"

# Chat history storage (in-memory cache)
chat_histories = {}
user_last_activity = {}

# Rate limiting
RATE_LIMIT_MESSAGES = 10  # messages per minute
RATE_LIMIT_WINDOW = 60  # seconds
user_message_counts = {}

SERVER_START_TIME = get_malaysia_time().timestamp()
# Global variable to track recent messages (place after other global variables)
recent_processed_messages = {}

class LanguageDetector:
    """Enhanced language detection and response matching"""
    
    def __init__(self):
        # Language mapping for better detection
        self.language_map = {
            'en': 'English',
            'zh': 'Chinese',
            'zh-cn': 'Chinese',
            'zh-tw': 'Chinese',
            'ms': 'Malay',
            'id': 'Malay'  # Indonesian is similar to Malay
        }
        
        # Common phrases for language detection
        self.language_phrases = {
            'zh': ['你好', '您好', '谢谢', '不好意思', '请问', '怎么', '什么', '哪里', '多少', '时间'],
            'en': ['hello', 'hi', 'thank', 'please', 'how', 'what', 'where', 'when', 'why', 'can'],
            'ms': ['hello', 'terima kasih', 'maaf', 'tolong', 'bagaimana', 'apa', 'di mana', 'bila', 'boleh']
        }
        
        # Address patterns for different regions
        self.address_patterns = [
            # Malaysian address patterns
            r'\b\d+[A-Za-z]?,?\s+(?:Jalan|Jln|Lorong|Lg|Persiaran|Lebuh|Taman|Kg|Kampung|Bandar|Shah Alam|Petaling Jaya|Kuala Lumpur|Johor Bahru|Penang|Ipoh|Klang|Subang Jaya|Ampang|Seremban|Melaka|Alor Setar|Kuching|Kota Kinabalu)[^.!?]*',
            
            # English address patterns  
            r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Lane|Ln|Drive|Dr|Boulevard|Blvd|Place|Pl|Court|Ct|Way|Circle|Cir)\b[^.!?]*',
            
            # Postal codes
            r'\b\d{5}(?:-\d{4})?\s+[A-Za-z\s]+\b',  # US format
            r'\b\d{5}\s+[A-Za-z\s]+\b',  # Malaysian format
            
            # Chinese address patterns (basic)
            r'[一-龯]+[市区县镇街道路巷弄号楼室栋层]\s*\d*[号楼室栋层]?[^.!?]*',
            
            # General patterns with numbers and location keywords
            r'\b(?:No\.?\s*|#)\d+[A-Za-z]?,?\s+[A-Za-z\s,]+(?:Town|City|State|Province|District|Area)[^.!?]*',
            
            # Building/complex names with numbers
            r'\b(?:Block|Blk|Unit|Suite|Apt|Apartment)\s*[A-Za-z]?\d+[A-Za-z]?,?\s+[A-Za-z\s,]+[^.!?]*'
        ]

    def detect_language(self, text):
        """Enhanced language detection"""
        if not text or not text.strip():
            return 'en'  # Default to English
            
        text_lower = text.lower().strip()
        
        # Check for Chinese characters
        if any('\u4e00' <= char <= '\u9fff' for char in text):
            return 'zh'
            
        # Check for common phrases
        for lang, phrases in self.language_phrases.items():
            for phrase in phrases:
                if phrase in text_lower:
                    return lang
        
        # Use langid as fallback
        try:
            import langid
            detected_lang, confidence = langid.classify(text)
            # Only trust langid if confidence is high
            if confidence > 0.7 and detected_lang in self.language_map:
                return detected_lang
            elif detected_lang in ['zh-cn', 'zh-tw']:
                return 'zh'
            elif detected_lang in ['id', 'ms']:
                return 'ms'
        except Exception as e:
            logging.error(f"Language detection error: {e}")
        
        # Default to English if uncertain
        return 'en'
    
    def extract_addresses(self, text):
        """Extract addresses from text and return them with their positions"""
        addresses = []
        
        for pattern in self.address_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                addresses.append({
                    'text': match.group(0),
                    'start': match.start(),
                    'end': match.end()
                })
        
        # Sort by start position and remove overlaps
        addresses = sorted(addresses, key=lambda x: x['start'])
        unique_addresses = []
        
        for addr in addresses:
            # Check if this address overlaps with any existing one
            overlap = False
            for existing in unique_addresses:
                if (addr['start'] < existing['end'] and addr['end'] > existing['start']):
                    # Keep the longer match
                    if len(addr['text']) > len(existing['text']):
                        unique_addresses.remove(existing)
                    else:
                        overlap = True
                    break
            
            if not overlap:
                unique_addresses.append(addr)
        
        return unique_addresses
    
    def reply_contains_address(self, reply_text: str) -> bool:
        if not reply_text or not str(reply_text).strip():
            return False
        try:
            addrs = self.extract_addresses(reply_text)  # uses your existing extractor
            return bool(addrs)
        except Exception:
            return False

    def is_address_query(self, user_text: str) -> bool:
        if not user_text:
            return False
        t = str(user_text).lower()
        keywords = [
            "地址","地点","地點","定位","在哪里","在哪裡",
            "address","location","map","google map","pin",
            "alamat","lokasi","peta"
        ]
        return any(k in t for k in keywords)

    def translate_preserving_addresses(self, text: str, target_lang: str) -> str:
        """
        Translate 'text' to target_lang but keep any detected address substrings EXACT.
        Strategy: mask address pieces with tokens, translate, then restore originals.
        """
        if not text:
            return text

        try:
            address_chunks = self.extract_addresses(text) or []
        except Exception:
            address_chunks = []

        if not address_chunks:
            # No address—normal translate
            try:
                return self.translate_text(text, target_lang) or text
            except Exception:
                return text

        # Step 1: mask each address chunk ONE-BY-ONE to avoid over-replacing duplicates
        masked_text = text
        original_map = []
        for i, addr in enumerate(address_chunks):
            if not addr:
                continue
            token = f"__ADDR_{i}__"
            # replace only the first occurrence left-to-right each time
            # (use regex escape to avoid special chars in address)
            import re
            masked_text, count = re.subn(re.escape(addr), token, masked_text, count=1)
            if count > 0:
                original_map.append((token, addr))

        # Step 2: translate masked text
        try:
            translated = self.translate_text(masked_text, target_lang) or masked_text
        except Exception:
            translated = masked_text

        # Step 3: restore tokens back to original address substrings (exact formatting)
        for token, addr in original_map:
            translated = translated.replace(token, addr)

        return translated

    def should_skip_address_variation(self, user_text: str, reply_text: str) -> bool:
        """
        When the reply contains addresses (or user asked for address), we avoid
        paraphrase/variation to prevent altering address formatting.
        """
        return self.reply_contains_address(reply_text) or self.is_address_query(user_text)
        
    def preserve_addresses_translate(self, text, target_language):
        """Translate text while preserving addresses in their original form"""
        if not text or not text.strip():
            return text
            
        # If target language is English and text appears to be English, return as-is
        if target_language == 'en' and all(ord(char) < 128 for char in text if char.isalpha()):
            return text
        
        try:
            # Extract addresses first
            addresses = self.extract_addresses(text)
            
            if not addresses:
                # No addresses found, translate normally
                return self.translate_text_basic(text, target_language)
            
            # Create placeholders for addresses
            text_with_placeholders = text
            address_placeholders = {}
            
            # Replace addresses with placeholders (work backwards to maintain positions)
            for i, addr in enumerate(reversed(addresses)):
                placeholder = f"__ADDRESS_PLACEHOLDER_{len(addresses)-i-1}__"
                text_with_placeholders = (
                    text_with_placeholders[:addr['start']] + 
                    placeholder + 
                    text_with_placeholders[addr['end']:]
                )
                address_placeholders[placeholder] = addr['text']
            
            # Translate text with placeholders
            translated_with_placeholders = self.translate_text_with_instruction(
                text_with_placeholders, target_language, preserve_placeholders=True
            )
            
            # Restore original addresses
            final_translated = translated_with_placeholders
            for placeholder, original_address in address_placeholders.items():
                final_translated = final_translated.replace(placeholder, original_address)
            
            logging.info(f"Translated with preserved addresses: {len(addresses)} addresses kept original")
            return final_translated
            
        except Exception as e:
            logging.error(f"Error in address-preserving translation: {e}")
            # Fallback to basic translation
            return self.translate_text_basic(text, target_language)
    
    def translate_text_with_instruction(self, text, target_language, preserve_placeholders=False):
        """Translate text with specific instructions about placeholder preservation"""
        try:
            import openai
            
            # Map language codes to full names for better translation
            language_names = {
                'en': 'English',
                'zh': 'Chinese (Simplified)', 
                'ms': 'Malay'
            }
            
            target_lang_name = language_names.get(target_language, 'English')
            
            # Enhanced system prompt
            system_content = f"You are a translator. Translate the following text to {target_lang_name}."
            
            if preserve_placeholders:
                system_content += " IMPORTANT: Keep any text that looks like '__ADDRESS_PLACEHOLDER_N__' exactly as-is without translating it. These are placeholders for addresses that should not be changed."
            
            system_content += " Only return the translated text, no explanations."
            
            client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",  # Use the same model as the main chatbot
                messages=[
                    {
                        "role": "system", 
                        "content": system_content
                    },
                    {
                        "role": "user", 
                        "content": text
                    }
                ],
                max_tokens=500,
                temperature=0.3
            )
            
            translated_text = response.choices[0].message.content.strip()
            return translated_text
            
        except Exception as e:
            logging.error(f"Translation with instruction error: {e}")
            return text
    
    def translate_text_basic(self, text, target_language):
        """Basic translation without address preservation (original method)"""
        try:
            import openai
            
            language_names = {
                'en': 'English',
                'zh': 'Chinese (Simplified)',
                'ms': 'Malay'
            }
            
            target_lang_name = language_names.get(target_language, 'English')
            
            client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {
                        "role": "system", 
                        "content": f"You are a translator. Translate the following text to {target_lang_name}. Only return the translated text, no explanations."
                    },
                    {
                        "role": "user", 
                        "content": text
                    }
                ],
                max_tokens=500,
                temperature=0.3
            )
            
            translated_text = response.choices[0].message.content.strip()
            return translated_text
            
        except Exception as e:
            logging.error(f"Basic translation error: {e}")
            return text
            
    def translate_text(self, text, target_language):
        """Main translation method - now uses address preservation by default"""
        return self.preserve_addresses_translate(text, target_language)
        
    def get_language_specific_prompt(self, base_prompt, detected_language):
        """Add language-specific instructions to the prompt"""
        language_instructions = {
            'en': "\n\nIMPORTANT: The user is communicating in English. Please respond in clear, natural English. Keep any addresses in their original format.",
            'zh': "\n\n重要提示：用户使用中文交流。请用自然、清晰的中文回复。保持地址的原始格式。",
            'ms': "\n\nPENTING: Pengguna berkomunikasi dalam Bahasa Melayu. Sila balas dalam Bahasa Melayu yang jelas dan natural. Kekalkan alamat dalam format asal."
        }
        
        instruction = language_instructions.get(detected_language, language_instructions['en'])
        return base_prompt + instruction

    def get_error_message(self, message_type, detected_language):
        """Get error messages in the detected language"""
        error_messages = {
            'unsupported_language': {
                'en': "We only support English, Chinese, and Malay. Please use one of these languages.",
                'zh': "我们只支持英文、中文和马来文。请使用其中一种语言。",
                'ms': "Kami hanya menyokong Bahasa Inggeris, Cina, dan Melayu. Sila gunakan salah satu daripada bahasa ini."
            },
            'sensitive_content': {
                'en': "This content cannot be processed. Please contact support.",
                'zh': "此内容无法处理。请联系客服。",
                'ms': "Kandungan ini tidak dapat diproses. Sila hubungi sokongan."
            },
            'rate_limit': {
                'en': "You're sending messages too quickly! Please wait a moment and try again.",
                'zh': "您发送消息太快了！请稍等片刻再试。",
                'ms': "Anda menghantar mesej terlalu cepat! Sila tunggu sebentar dan cuba lagi."
            },
            'service_busy': {
                'en': "Service is busy, please try again later!",
                'zh': "服务繁忙，请稍后再试！",
                'ms': "Perkhidmatan sibuk, sila cuba lagi nanti!"
            },
            'timeout': {
                'en': "Response timed out, please try again!",
                'zh': "响应超时，请重试！",
                'ms': "Respons tamat masa, sila cuba lagi!"
            },
            'service_unavailable': {
                'en': "Service is temporarily unavailable, please try again later.",
                'zh': "服务暂时不可用，请稍后再试。",
                'ms': "Perkhidmatan tidak tersedia buat sementara waktu, sila cuba lagi nanti."
            }
        }
        
        return error_messages.get(message_type, {}).get(detected_language, 
                                                       error_messages.get(message_type, {}).get('en', 'Error'))

class TextVariationGenerator:
    """Generate slight variations of text while preserving meaning"""
    
    def __init__(self):
        self.variation_cache = {}  # Cache to avoid generating same variations repeatedly
        self.cache_max_size = 100
        
    def generate_text_variation(self, original_text, language='en', variation_level='slight'):
        """
        Generate a slight variation of the original text while preserving meaning
        
        Args:
            original_text: The text to vary
            language: Language code ('en', 'zh', 'ms')
            variation_level: 'slight', 'moderate' - controls how much to vary
        """
        if not original_text or not original_text.strip():
            return original_text
            
        # Create cache key
        cache_key = f"{original_text[:50]}:{language}:{variation_level}"
        
        # Return cached variation if available
        if cache_key in self.variation_cache:
            return self.variation_cache[cache_key]
        
        try:
            # Clean cache if it gets too large
            if len(self.variation_cache) > self.cache_max_size:
                # Remove oldest 20 entries
                keys_to_remove = list(self.variation_cache.keys())[:20]
                for key in keys_to_remove:
                    del self.variation_cache[key]
            
            # Generate variation using OpenAI
            varied_text = self._generate_variation_with_openai(original_text, language, variation_level)
            
            # Cache the result
            self.variation_cache[cache_key] = varied_text
            
            return varied_text
            
        except Exception as e:
            logging.warning(f"Text variation failed for '{original_text[:30]}...': {e}")
            # Return original text if variation fails
            return original_text
    
    def _generate_variation_with_openai(self, text, language, variation_level):
        """Use OpenAI to generate text variations"""
        try:
            import openai
            
            # Language-specific instructions
            language_instructions = {
                'en': {
                    'name': 'English',
                    'instructions': 'Slightly rephrase this English text to make it sound different while keeping the exact same meaning. Use synonyms, different sentence structure, or slightly different phrasing.',
                    'warning': 'Keep it professional and natural.'
                },
                'zh': {
                    'name': 'Chinese',
                    'instructions': '请稍微改写这段中文文本，让它听起来不同，但保持完全相同的意思。可以使用同义词、不同的句子结构或稍微不同的表达方式。',
                    'warning': '非常重要：保持原意不变，只改变表达方式。确保语气和正式程度保持一致。'
                },
                'ms': {
                    'name': 'Malay',
                    'instructions': 'Sila ubah sedikit teks Bahasa Melayu ini supaya kedengaran berbeza sambil mengekalkan maksud yang tepat sama. Gunakan sinonim, struktur ayat yang berbeza, atau frasa yang sedikit berbeza.',
                    'warning': 'Pastikan ia kekal profesional dan semula jadi.'
                }
            }
            
            lang_info = language_instructions.get(language, language_instructions['en'])
            
            # Variation level settings
            if variation_level == 'slight':
                intensity = "very slight"
                additional_instruction = "Make minimal changes - just enough so it doesn't sound exactly the same."
            else:  # moderate
                intensity = "moderate" 
                additional_instruction = "Make noticeable but not dramatic changes."
            
            # Construct system prompt
            system_prompt = f"""You are a text variation expert. Your task is to create {intensity} variations of text in {lang_info['name']}.

INSTRUCTIONS: {lang_info['instructions']}

IMPORTANT RULES:
1. {lang_info['warning']}
2. {additional_instruction}
3. Keep the same tone, formality level, and politeness
4. Preserve all specific information (names, numbers, technical terms)
5. Do NOT change the core message or meaning
6. Return only the varied text, no explanations

For Chinese text: 特别注意保持原文的语气、正式程度和礼貌程度不变。"""

            client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Original text: {text}"}
                ],
                max_tokens=800,
                temperature=0.7  # Moderate creativity for variations
            )
            
            varied_text = response.choices[0].message.content.strip()
            
            # Basic validation - ensure we got a reasonable variation
            if len(varied_text) < len(text) * 0.5 or len(varied_text) > len(text) * 2:
                logging.warning(f"Variation length seems unreasonable, using original text")
                return text
                
            logging.info(f"Generated text variation for {language}: '{text[:30]}...' -> '{varied_text[:30]}...'")
            return varied_text
            
        except Exception as e:
            logging.error(f"OpenAI variation generation failed: {e}")
            return text

class ChatBot:
    def __init__(self):
        # self.history_manager = ChatHistoryManager()
        self.history_manager = EnhancedChatHistoryManager()  # Uses GitHub integration
        self.file_processor = FileProcessor()
        self.language_detector = LanguageDetector()  # Add language detector
        
        self.text_variation_generator = TextVariationGenerator()
        
        # Initialize MediaHandler with correct paths
        self.media_handler = MediaHandler(
            # media_base_path="../../ReactProject/appointmentwhatsapp/public/media_opening",
            # media_base_url="http://localhost:3000/media_opening"  # Adjust this to your actual server URL
            media_base_path="media_opening",  # Local fallback
            media_base_url="https://reactjs-appointmentwhatsapp-production.up.railway.app/media_opening"
        )
        
        self.special_reply_media_handler = MediaHandler(
            media_base_path="media_special_reply",
            media_base_url="https://reactjs-appointmentwhatsapp-production.up.railway.app/media_special_reply"
        )

        self.load_frame_content()
        self.current_token = META_WHATSAPP_TOKEN
        
        # Supported languages
        self.supported_languages = {
            'en': 'English',
            'zh': 'Chinese',
            'ms': 'Malay'
        }
        
        # Preload opening media files
        self.preload_opening_media()
    
    def preload_opening_media(self):
        """Preload and verify opening media files"""
        try:
            opening_messages = self.history_manager.get_opening_messages()
            media_files = []
            
            for msg in opening_messages:
                if msg['TYPE'] in [2, 3] and msg.get('MEDIANAME'):
                    media_type = 'image' if msg['TYPE'] == 2 else 'video'
                    media_files.append((msg['MEDIANAME'], media_type))
            
            logging.info(f"Preloading {len(media_files)} media files...")
            
            # Check all media files in parallel
            with ThreadPoolExecutor(max_workers=5) as executor:
                futures = {
                    executor.submit(self.media_handler.file_exists, filename, media_type): (filename, media_type)
                    for filename, media_type in media_files
                }
                
                for future in futures:
                    filename, media_type = futures[future]
                    try:
                        exists = future.result(timeout=10)
                        if not exists:
                            logging.error(f"Opening {media_type} file not accessible: {filename}")
                        else:
                            logging.info(f"Opening {media_type} file verified: {filename}")
                    except Exception as e:
                        logging.error(f"Error preloading {filename}: {e}")
            
            logging.info("Media preloading completed")
            
        except Exception as e:
            logging.error(f"Error during media preloading: {e}")
            
    def load_frame_content(self):
        """Load frame content with new structure"""
        frame_data = self.history_manager.get_frame_content()
        
        if frame_data:
            self.system_prompt = frame_data.get('FRAME_CONTENT', '') or self.default_system_prompt()
            
            # Load single sensitive content and reply
            self.sensitive_content = frame_data.get('SENSITIVE_CONTENT', '') or ''
            self.sensitive_reply = frame_data.get('SENSITIVE_REPLY', '') or ''
            
            self.additional_content = frame_data.get('ADDITIONAL_CONTENT', '') or ''

            # Load and process additional content files
            self.load_additional_files_content()

        else:
            # Fallback defaults
            self.system_prompt = self.default_system_prompt()
            self.sensitive_content = ''
            self.sensitive_reply = ''
            self.additional_content = ''
            self.files_content = ''

        logging.info("Frame content loaded successfully")
        logging.info(f"System prompt: {self.system_prompt[:100]}...")
        logging.info(f"Sensitive content: {self.sensitive_content}")
        logging.info(f"Sensitive reply: {self.sensitive_reply}")

    def load_additional_files_content(self):
        """Load and extract text from additional content files"""
        self.files_content = ""
        files = self.history_manager.get_additional_content_files()
        
        for file_info in files:
            filename = file_info['filename']
            original_name = file_info['original_name']
            mime_type = file_info['mime_type']
            
            # Get file content from database
            file_data = self.history_manager.get_file_content_by_filename(filename)
            
            if file_data and file_data['content']:
                # Extract text from file
                extracted_text = self.file_processor.extract_text_from_file(
                    file_data['content'], mime_type
                )
                
                if extracted_text:
                    self.files_content += f"\n\n--- Content from {original_name} ---\n"
                    self.files_content += extracted_text
                    logging.info(f"Loaded content from file: {original_name} ({len(extracted_text)} characters)")
                else:
                    logging.warning(f"Could not extract text from file: {original_name}")
            else:
                logging.warning(f"Could not load file content: {filename}")

    def default_system_prompt(self):
        return """You are a professional assistant. Please respond professionally to user queries and always respond in the same language as the user's message."""

    def detect_message_language(self, message):
        """Detect the language of the user message"""
        return self.language_detector.detect_language(message)

    def is_supported_language(self, text):
        """Check if text is in supported language using enhanced detection"""
        if not text.strip():
            return False
            
        detected_lang = self.detect_message_language(text)
        return detected_lang in self.supported_languages

    def contains_sensitive_words(self, text, detected_language=None):
        """Enhanced sensitive word detection with auto-translation"""
        if not text.strip():
            return False
            
        # If no sensitive content configured, skip check
        if not self.sensitive_content or not self.sensitive_content.strip():
            logging.info("No sensitive content configured, skipping sensitive word check")
            return False
            
        if not detected_language:
            detected_language = self.language_detector.detect_language(text)
            
        text_lower = text.lower().strip()
        
        # Split sensitive content into keywords
        sensitive_keywords = [keyword.strip() for keyword in self.sensitive_content.split(',') if keyword.strip()]
        
        if not sensitive_keywords:
            return False
        
        logging.info(f"Checking sensitive words for language '{detected_language}' in text: '{text[:50]}...'")
        
        for keyword in sensitive_keywords:
            if not keyword:
                continue
                
            # First check if keyword matches directly (same language)
            if keyword.lower() in text_lower:
                logging.info(f"Direct match found for sensitive keyword: '{keyword}'")
                return True
            
            # If detected language is not English, translate the keyword to detected language
            if detected_language != 'en':
                try:
                    translated_keyword = self.language_detector.translate_text(keyword, detected_language)
                    if translated_keyword and translated_keyword.lower() in text_lower:
                        logging.info(f"Translated match found: '{keyword}' -> '{translated_keyword}'")
                        return True
                except Exception as e:
                    logging.error(f"Error translating sensitive keyword '{keyword}': {e}")
            
            # If detected language is not the original language of keyword, also check reverse translation
            if detected_language == 'en':
                # Try translating to other languages and check
                for lang in ['zh', 'ms']:
                    try:
                        translated_keyword = self.language_detector.translate_text(keyword, lang)
                        if translated_keyword and translated_keyword.lower() in text_lower:
                            logging.info(f"Reverse translated match found: '{keyword}' -> '{translated_keyword}' ({lang})")
                            return True
                    except Exception as e:
                        logging.error(f"Error reverse translating sensitive keyword '{keyword}' to {lang}: {e}")
                        
        return False

    def get_sensitive_content_reply(self, detected_language):
        """Get sensitive content reply message in detected language with variation"""
        # If no sensitive reply configured, use default message
        if not self.sensitive_reply or not self.sensitive_reply.strip():
            logging.info("No sensitive reply configured, using default message")
            return self.language_detector.get_error_message('sensitive_content_default', detected_language)
        
        # If sensitive reply is configured, translate it to detected language if needed
        try:
            # Check if the reply is already in the detected language (basic heuristic)
            base_reply = self.sensitive_reply
            if detected_language == 'en':
                # For English, check if contains mostly ASCII characters
                if all(ord(char) < 128 for char in self.sensitive_reply if char.isalpha()):
                    base_reply = self.sensitive_reply
                else:
                    base_reply = self.language_detector.translate_text(self.sensitive_reply, detected_language)
            elif detected_language == 'zh':
                # For Chinese, check if contains Chinese characters
                if any('\u4e00' <= char <= '\u9fff' for char in self.sensitive_reply):
                    base_reply = self.sensitive_reply
                else:
                    base_reply = self.language_detector.translate_text(self.sensitive_reply, detected_language)
            else:
                # Translate the sensitive reply to detected language
                base_reply = self.language_detector.translate_text(self.sensitive_reply, detected_language)
            
            # Generate a variation of the reply
            varied_reply = self.text_variation_generator.generate_text_variation(
                base_reply, detected_language, 'slight'
            )
            
            logging.info(f"Generated varied sensitive reply for {detected_language}: '{varied_reply[:50]}...'")
            return varied_reply
            
        except Exception as e:
            logging.error(f"Error generating varied sensitive reply: {e}")
            # Fallback to default message if variation fails
            return self.language_detector.get_error_message('sensitive_content_default', detected_language)

    def handle_opening_sequence_by_keywords(self, phone_number: str, user_text: str, detected_language: str) -> bool:
        """首次消息：先用关键词尝试 default=0 开场（内部调用 handle_opening_sequence），成功则置完成标记。"""
        try:
            if not user_text or not str(user_text).strip():
                return False
            res = self.handle_opening_sequence(
                phone_number,
                user_language=detected_language or 'en',
                original_message=user_text,
                detected_language=detected_language or 'en'
            )
            if res is not None:
                try:
                    self.history_manager.set_opening_completed_by_phone(phone_number, 1)
                except Exception:
                    pass
                return True
            return False
        except Exception as e:
            logging.error("handle_opening_sequence_by_keywords error: %s", e)
            return False
        
    def handle_opening_sequence(self, phone_number, user_language='en', original_message=None, detected_language=None):
        """第一次用户触达时的开场逻辑：先匹配关键词分类(default=0)，否则发默认分类(default=1)"""
        try:
            # 只有第一次用户才发开场
            # if not self.history_manager.is_first_time_user(phone_number):
                # logging.info(f"User {phone_number} is not a first-time user, skip opening.")
                # return None

            category_id = None
            matched = False

            # 如提供原始消息，尝试关键词匹配 default=0 的分类
            if original_message:
                dl = detected_language or user_language or 'en'
                category_id = self.match_opening_category_by_keywords(original_message, detected_language=dl)
                matched = category_id is not None

            # 若未匹配，落入 default=1 的分类
            if not category_id:
                default_cat = self.history_manager.get_default_opening_category()
                if default_cat:
                    category_id = default_cat['PKKEY']

            if not category_id:
                logging.warning("No opening category found (neither keywords nor default).")
                return None

            # 取该分类下的开场消息序列
            opening_messages = self.history_manager.get_opening_messages_by_category(category_id)
            if not opening_messages:
                logging.warning(f"No opening messages under category {category_id}.")
                return {'category_id': category_id, 'matched': matched, 'sent': 0}

            logging.info(f"Opening category={category_id}, total msgs={len(opening_messages)}, lang={user_language}, matched={matched}")

            sent = 0
            for opening_msg in opening_messages:
                message_type = opening_msg['TYPE']

                if message_type == 1:  # 文本
                    text_content = opening_msg.get('TEXT_CONTENT') or ""
                    if not text_content:
                        continue

                    try:
                        prepared_text = self.language_detector.translate_preserving_addresses(
                            text_content, user_language
                        )

                        if self.language_detector.should_skip_address_variation(original_message, text_content):
                            text_to_send = prepared_text.strip()
                        else:
                            text_to_send = prepared_text

                    except Exception as e:
                        logging.error(f"Opening preserve-translate failed: {e}")
                        text_to_send = text_content.strip()

                    res = self.send_whatsapp_message(phone_number, text_to_send)
                    if res:
                        self.history_manager.save_message(
                            phone_number, text_to_send, "assistant", "text",
                            opening_message_id=opening_msg.get('PKKEY')
                        )
                        sent += 1

                elif message_type == 2:  # 图片
                    media_name = opening_msg.get('MEDIANAME')
                    if media_name and self.media_handler.file_exists(media_name):
                        res = self.send_whatsapp_media(phone_number, media_name, 'image')
                        if res:
                            self.history_manager.save_message(
                                phone_number, media_name, "assistant", "image",
                                opening_message_id=opening_msg['PKKEY']
                            )
                            sent += 1
                    else:
                        logging.error(f"Opening image not found: {media_name}")

                elif message_type == 3:  # 视频
                    media_name = opening_msg.get('MEDIANAME')
                    if media_name and self.media_handler.file_exists(media_name, 'video'):
                        res = self.send_whatsapp_media(phone_number, media_name, 'video')
                        if res:
                            self.history_manager.save_message(
                                phone_number, media_name, "assistant", "video",
                                opening_message_id=opening_msg['PKKEY']
                            )
                            sent += 1
                    else:
                        logging.error(f"Opening video not found: {media_name}")

                time.sleep(0.3)  # 轻微间隔，避免节流

            # 标记开场已完成（可按你需要更新 chat_sessions 的字段）
            self.history_manager.update_session_flags(phone_number, manual_reply=0)

            return {'category_id': category_id, 'matched': matched, 'sent': sent}

        except Exception as e:
            logging.error(f"Error in handle_opening_sequence: {e}")
            return None

    def send_whatsapp_media(self, to_number, media_filename, media_type):
        """Send media (image or video) via WhatsApp Business API"""
        try:
            headers = {
                "Authorization": f"Bearer {self.current_token}",
                "Content-Type": "application/json"
            }
            
            # Clean phone number format
            to_number = ''.join(filter(str.isdigit, to_number))
            
            # Check if media file exists
            if not self.media_handler.file_exists(media_filename, media_type):
                logging.error(f"Media file not found: {media_filename}")
                logging.error(f"Expected path: {self.media_handler.get_media_file_path(media_filename, media_type)}")
                return None
            
            # Get media URL - you need to implement this based on your setup
            media_url = self.media_handler.get_media_url(media_filename, media_type)
            
            # Prepare WhatsApp media payload
            if media_type == 'image':
                payload = {
                    "messaging_product": "whatsapp",
                    "recipient_type": "individual",
                    "to": to_number,
                    "type": "image",
                    "image": {
                        "link": media_url,
                        "caption": ""  # You can add caption if needed
                    }
                }
            elif media_type == 'video':
                payload = {
                    "messaging_product": "whatsapp",
                    "recipient_type": "individual", 
                    "to": to_number,
                    "type": "video",
                    "video": {
                        "link": media_url,
                        "caption": ""  # You can add caption if needed
                    }
                }
            else:
                logging.error(f"Unsupported media type: {media_type}")
                return None
            
            logging.info(f"Sending media payload: {json.dumps(payload, indent=2)}")

            response = requests.post(
                META_API_BASE_URL,
                headers=headers,
                json=payload,
                timeout=30
            )
            
            if response.status_code == 200:
                logging.info(f"Media sent successfully to {to_number}: {media_filename}")
                return response.json()
            else:
                logging.error(f"Failed to send media: {response.status_code} - {response.text}")
                # Log the full response for debugging
                try:
                    error_data = response.json()
                    logging.error(f"WhatsApp API Error Details: {json.dumps(error_data, indent=2)}")
                except:
                    logging.error(f"Raw response: {response.text}")
                return None
                
        except Exception as e:
            logging.error(f"Error sending media {media_filename} to {to_number}: {e}")
            import traceback
            logging.error(f"Full traceback: {traceback.format_exc()}")
            return None

    def is_rate_limited(self, phone_number):
        """Check if user is rate limited"""
        current_time = time.time()
        
        if phone_number not in user_message_counts:
            user_message_counts[phone_number] = []
        
        # Remove old messages outside the window
        user_message_counts[phone_number] = [
            msg_time for msg_time in user_message_counts[phone_number]
            if current_time - msg_time < RATE_LIMIT_WINDOW
        ]
        
        # Check if user exceeds rate limit
        if len(user_message_counts[phone_number]) >= RATE_LIMIT_MESSAGES:
            return True
        
        # Add current message time
        user_message_counts[phone_number].append(current_time)
        return False

    def refresh_access_token(self):
        """Refresh the access token using app credentials"""
        if not META_APP_ID or not META_APP_SECRET:
            logger.warning("Cannot refresh token: META_APP_ID or META_APP_SECRET not provided")
            return False
        
        try:
            # Get app access token
            app_token_url = f"https://graph.facebook.com/oauth/access_token"
            app_token_params = {
                'client_id': META_APP_ID,
                'client_secret': META_APP_SECRET,
                'grant_type': 'client_credentials'
            }
            
            response = requests.get(app_token_url, params=app_token_params, timeout=30)
            
            if response.status_code == 200:
                app_token = response.json().get('access_token')
                
                # Exchange for long-lived token
                exchange_url = f"https://graph.facebook.com/oauth/access_token"
                exchange_params = {
                    'grant_type': 'fb_exchange_token',
                    'client_id': META_APP_ID,
                    'client_secret': META_APP_SECRET,
                    'fb_exchange_token': self.current_token
                }
                
                exchange_response = requests.get(exchange_url, params=exchange_params, timeout=30)
                
                if exchange_response.status_code == 200:
                    new_token = exchange_response.json().get('access_token')
                    if new_token:
                        self.current_token = new_token
                        logger.info("Access token refreshed successfully")
                        return True
                        
            logger.error(f"Failed to refresh token: {response.status_code} - {response.text}")
            return False
            
        except Exception as e:
            logger.error(f"Error refreshing access token: {e}")
            return False

    def handle_token_error(self, response):
        """Handle token-related errors and attempt refresh"""
        if response.status_code == 401:
            error_data = response.json().get('error', {})
            error_code = error_data.get('code')
            error_subcode = error_data.get('error_subcode')
            
            logger.error(f"Token error - Code: {error_code}, Subcode: {error_subcode}, Message: {error_data.get('message')}")
            
            if error_code == 190:  # Invalid access token
                if error_subcode == 463:  # Session has expired
                    logger.error("Session has expired. This might indicate:")
                    logger.error("1. Token is not actually permanent")
                    logger.error("2. Token permissions were revoked")
                    logger.error("3. App or Business account issues")
                    
                logger.warning("Access token expired, attempting to refresh...")
                if self.refresh_access_token():
                    return True
                else:
                    logger.error("Failed to refresh access token. Please generate a new permanent token.")
                    logger.error("Visit: https://business.facebook.com/settings/system-users")
        return False

    def validate_token(self):
        """Validate the current access token and check its details"""
        try:
            # Check token info
            token_info_url = f"https://graph.facebook.com/me?access_token={self.current_token}"
            response = requests.get(token_info_url, timeout=30)
            
            if response.status_code == 200:
                token_data = response.json()
                logger.info(f"Token validation successful: {token_data}")
                
                # Check token permissions
                permissions_url = f"https://graph.facebook.com/me/permissions?access_token={self.current_token}"
                perm_response = requests.get(permissions_url, timeout=30)
                
                if perm_response.status_code == 200:
                    permissions = perm_response.json()
                    logger.info(f"Token permissions: {permissions}")
                    
                    # Check for WhatsApp permissions
                    required_perms = ['whatsapp_business_messaging', 'whatsapp_business_management']
                    granted_perms = [p['permission'] for p in permissions.get('data', []) if p.get('status') == 'granted']
                    
                    missing_perms = [p for p in required_perms if p not in granted_perms]
                    if missing_perms:
                        logger.warning(f"Missing permissions: {missing_perms}")
                        return False, f"Missing permissions: {missing_perms}"
                    
                    return True, "Token is valid and has required permissions"
                else:
                    logger.error(f"Failed to check permissions: {perm_response.status_code} - {perm_response.text}")
                    return False, "Could not verify token permissions"
            else:
                logger.error(f"Token validation failed: {response.status_code} - {response.text}")
                return False, f"Token validation failed: {response.text}"
                
        except Exception as e:
            logger.error(f"Error validating token: {e}")
            return False, f"Error validating token: {e}"

    def get_token_info(self):
        """Get detailed information about the current token"""
        try:
            # Get app access token to check the user token
            debug_url = f"https://graph.facebook.com/debug_token"
            params = {
                'input_token': self.current_token,
                'access_token': f"{META_APP_ID}|{META_APP_SECRET}" if META_APP_ID and META_APP_SECRET else self.current_token
            }
            
            response = requests.get(debug_url, params=params, timeout=30)
            
            if response.status_code == 200:
                debug_data = response.json()
                logger.info(f"Token debug info: {json.dumps(debug_data, indent=2)}")
                
                token_data = debug_data.get('data', {})
                expires_at = token_data.get('expires_at', 0)
                
                if expires_at == 0:
                    logger.info("Token is permanent (never expires)")
                else:
                    expiry_date = datetime.fromtimestamp(expires_at)
                    logger.info(f"Token expires at: {expiry_date}")
                    
                return True, debug_data
            else:
                logger.error(f"Failed to get token info: {response.status_code} - {response.text}")
                return False, response.text
                
        except Exception as e:
            logger.error(f"Error getting token info: {e}")

    def clean_chat_history(self, phone_number):
        """Clean old chat history to save memory"""
        if phone_number in chat_histories:
            # Keep only last 20 messages (10 pairs)
            chat_histories[phone_number] = chat_histories[phone_number][-20:]
    
    def send_whatsapp_message(self, to_number, message, retry=True):
        """Send a message using WhatsApp Business API with enhanced error handling"""
        headers = {
            "Authorization": f"Bearer {self.current_token}",
            "Content-Type": "application/json"
        }
        
        # Clean phone number format
        to_number = ''.join(filter(str.isdigit, to_number))
        
        # Split long messages
        if len(message) > 1600:  # WhatsApp limit is ~4096, but keep it shorter
            messages = [message[i:i+1600] for i in range(0, len(message), 1600)]
        else:
            messages = [message]
        
        responses = []
        for msg in messages:
            payload = {
                "messaging_product": "whatsapp",
                "recipient_type": "individual",
                "to": to_number,
                "type": "text",
                "text": {
                    "preview_url": False,
                    "body": msg
                }
            }
            
            try:
                response = requests.post(
                    META_API_BASE_URL,
                    headers=headers,
                    json=payload,
                    timeout=30
                )
                
                if response.status_code == 200:
                    logger.info(f"Message sent successfully to {to_number}")
                    responses.append(response.json())
                elif response.status_code == 401 and retry:
                    # Try to refresh token and retry once
                    if self.handle_token_error(response):
                        logger.info("Retrying message send with refreshed token...")
                        return self.send_whatsapp_message(to_number, message, retry=False)
                    else:
                        logger.error(f"Failed to send message: {response.status_code} - {response.text}")
                        return None
                else:
                    logger.error(f"Failed to send message: {response.status_code} - {response.text}")
                    return None
                    
                # Small delay between messages if sending multiple
                if len(messages) > 1:
                    time.sleep(1)
                    
            except requests.exceptions.Timeout:
                logger.error(f"Timeout sending message to {to_number}")
                return None
            except Exception as e:
                logger.error(f"Error sending message to {to_number}: {e}")
                return None
        
        return responses

# Enhanced chatbot with MySQL storage and opening sequence
class EnhancedChatBot(ChatBot):
    def __init__(self):
        super().__init__()
        # Use the enhanced history manager
        self.history_manager = EnhancedChatHistoryManager()
        self.message_analyzer = MessageAnalyzer()
        
        # Add possible media paths for special reply files
        self.possible_media_paths = [
            "media_special_reply",
            "public/media_special_reply", 
            "../reactjs-appointmentwhatsapp/public/media_special_reply"
        ]
    
    def get_manual_reply_message(self, detected_language='en', message_type='general'):
        """
        Get manual/auto reply message from chatbox_frame (via history_manager.get_frame_content),
        translate to the detected language while PRESERVING any address substrings exactly,
        and (only if safe) apply a slight variation.

        Requires LanguageDetector to implement:
          - translate_preserving_addresses(text, target_lang)
          - should_skip_address_variation(user_text, reply_text)

        Note:
          - We intentionally avoid paraphrasing when the reply contains addresses
            (or when the user asked for an address), to keep formatting exact.
        """
        try:
            frame_data = self.history_manager.get_frame_content()

            if not frame_data:
                # Fallback to original hardcoded messages if no database content
                return self._get_fallback_manual_reply(detected_language, message_type)

            # Pick the correct column by message_type
            if message_type == 'appointment':
                reply_text = frame_data.get('APPOINTMENT_REPLY', '')
            elif message_type == 'human_service':
                reply_text = frame_data.get('HUMAN_SERVICE_REPLY', '')
            elif message_type == 'media':
                reply_text = frame_data.get('ATTACH_MEDIA_REPLY', '')
            else:  # 'general' or unknown -> reuse a sensible default or add your own GENERAL_REPLY column
                reply_text = frame_data.get('APPOINTMENT_REPLY', '')

            if not reply_text or not str(reply_text).strip():
                # Fallback if the chosen DB field is empty
                return self._get_fallback_manual_reply(detected_language, message_type)

            # --- Address-aware translation (preserve address substrings exactly) ---
            try:
                prepared_reply = self.language_detector.translate_preserving_addresses(
                    reply_text, detected_language
                )
            except Exception as trans_err:
                logging.error(f"Manual/Auto preserve-translate failed: {trans_err}")
                prepared_reply = reply_text

            # --- Avoid variation if addresses involved or user asked for address ---
            # We don't have the user's raw text in this method, so pass None (the helper
            # will still skip variation if it detects address substrings in reply_text).
            try:
                if self.language_detector.should_skip_address_variation(None, reply_text):
                    logging.info("Manual/Auto reply contains address (or address intent); "
                                 "sending WITHOUT variation.")
                    return prepared_reply.strip()
            except Exception as addr_flag_err:
                logging.warning(f"Address-variation decision failed: {addr_flag_err}")

            # --- Safe to vary slightly if no addresses are involved ---
            try:
                varied_reply = self.text_variation_generator.generate_text_variation(
                    prepared_reply, detected_language, 'slight'
                )
                if varied_reply and str(varied_reply).strip():
                    logging.info(f"Generated varied manual reply ({message_type}) for {detected_language}.")
                    return str(varied_reply).strip()
                return prepared_reply.strip()
            except Exception as var_err:
                logging.warning(f"Manual/Auto variation failed: {var_err}")
                return prepared_reply.strip()

        except Exception as e:
            logging.error(f"Error getting manual reply from database: {e}")
            # Fallback to original hardcoded messages
            return self._get_fallback_manual_reply(detected_language, message_type)

    def _get_fallback_manual_reply(self, detected_language='en', message_type='general'):
        """Fallback manual reply messages (original hardcoded ones)"""
        messages = {
            'appointment': {
                'en': "Thank you for your appointment request. Our team will review your request and contact you shortly to confirm the details and schedule.",
                'zh': "感谢您的预约申请。我们的专业团队将尽快审核您的需求，并主动联系您确认详细安排。",
                'ms': "Terima kasih atas permohonan temujanji anda. Pasukan kami akan mengkaji permintaan anda dan menghubungi anda tidak lama lagi untuk mengesahkan butiran dan jadual."
            },
            'human_service': {
                'en': "Thank you for your request. Our customer service team will prioritize your inquiry and connect with you as soon as possible.",
                'zh': "感谢您的咨询需求。我们的客服团队将优先处理您的问题，并尽快安排专人与您取得联系。",
                'ms': "Terima kasih atas permintaan anda. Pasukan khidmat pelanggan kami akan mengutamakan pertanyaan anda dan berhubung dengan anda secepat mungkin."
            },
            'media': {
                'en': "Thank you for sharing the media content. Our support team will review it and provide you with appropriate assistance shortly.",
                'zh': "感谢您分享的媒体内容。我们的技术支持团队将仔细查看并尽快为您提供针对性的专业协助。",
                'ms': "Terima kasih kerana berkongsi kandungan media. Pasukan sokongan kami akan menyemaknya dan memberikan anda bantuan yang sesuai tidak lama lagi."
            },
            'general': {
                'en': "Thank you for your message. Our team will arrange for a specialist to review your inquiry and respond to you promptly.",
                'zh': "感谢您的留言。我们将尽快安排专业人员审核您的咨询并及时回复您。",
                'ms': "Terima kasih atas mesej anda. Pasukan kami akan mengatur pakar untuk menyemak pertanyaan anda dan membalas anda dengan segera."
            }
        }
        
        return messages.get(message_type, messages['general']).get(detected_language, messages['general']['en'])
    
    def _get_fallback_message_for_intent(self, detected_language, is_appointment_request):
        """Fallback messages when AI fails for appointment/human service requests"""
        if is_appointment_request:
            messages = {
                'en': "Thank you for your appointment inquiry. I'll help you with that right away!",
                'zh': "感谢您的预约咨询。我会立即为您提供帮助！",
                'ms': "Terima kasih atas pertanyaan temujanji anda. Saya akan membantu anda dengan segera!"
            }
        else:  # human service request
            messages = {
                'en': "I understand you'd like to speak with our team. Let me assist you with your inquiry.",
                'zh': "我理解您希望与我们的团队交流。让我来协助您解决您的问题。",
                'ms': "Saya faham anda ingin bercakap dengan pasukan kami. Biar saya bantu dengan pertanyaan anda."
            }
        
        return messages.get(detected_language, messages['en'])
        
    # FIXED: Add the corrected message processing flow
    def process_user_message(self, message, phone_number):
        """
        ENHANCED: New process flow with human_reply flag check at the beginning:
        0.0 Check human_reply flag - if set, skip AI processing completely
        5.1 Check First Time Message
        5.2 Check get_manual_reply_message (appointment/human service) 
        5.3 Check Special Reply Message
        5.4 Normal AI processing (only if all above are skipped)
        """
        try:
            # 0.0: PRIORITY CHECK - Human reply flag (HIGHEST PRIORITY)
            session_flags = self.history_manager.get_session_flags(phone_number)
            if session_flags.get('human_reply', 0) == 1:
                logging.info(f"Human reply flag active for {phone_number}, skipping ALL automated processing")
                
                # Save the user message but mark as processed since no AI response will be generated
                user_message_id = self.history_manager.save_message(phone_number, message, "user")
                if user_message_id:
                    self.history_manager.mark_message_processed(user_message_id)
                    logging.info(f"User message saved and marked as processed due to human_reply flag: {user_message_id}")
                
                # Return None to indicate no automated response should be sent
                return None
            
            # 0.1: Save phone number to followup keep table
            try:
                followup_saved = self.history_manager.save_phone_number_to_followup_keep(phone_number)
                if followup_saved:
                    logging.info(f"New phone number added to followup keep: {phone_number}")
                elif followup_saved is False:
                    logging.debug(f"Phone number already in followup keep: {phone_number}")
                else:
                    logging.warning(f"Failed to save phone number to followup keep: {phone_number}")
            except Exception as followup_error:
                logging.error(f"Error in followup keep save: {followup_error}")
                # Continue processing even if followup save fails
            
            # Detect language first
            detected_language = self.detect_message_language(message)
            logging.info(f"Detected language: {detected_language}")
            
            # 5.1: First-time user check (HIGHEST PRIORITY after human_reply)
            is_first_time = self.history_manager.is_first_time_user(phone_number)
            
            if is_first_time:
                logging.info(f"5.1: First-time user detected: {phone_number}")
                
                # Save user message first
                user_message_id = self.history_manager.save_message(phone_number, message, "user")
                
                # Handle opening sequence
                opening_result = self.handle_opening_sequence_by_keywords(
                    phone_number, message, detected_language
                )
                
                if not opening_result:
                    # Fallback to default opening
                    opening_result = self.handle_opening_sequence(
                        phone_number,
                        user_language=detected_language,
                        original_message=message,
                        detected_language=detected_language
                    )
                
                # Mark as processed and complete
                if user_message_id:
                    self.history_manager.mark_message_processed(user_message_id)
                
                # Mark opening as completed
                self.history_manager.set_opening_completed_by_phone(phone_number, 1)
                
                return None  # Opening sequence sent, no additional response needed

            # 5.2: Check appointment/human service requests (SECOND PRIORITY)
            try:
                is_appointment_request = self.message_analyzer.detect_appointment_request(message, detected_language)
                is_human_service_request = self.message_analyzer.detect_human_service_request(message, detected_language)
            except Exception as _intent_e:
                logging.error(f"Intent detection failed: {_intent_e}")
                is_appointment_request = False
                is_human_service_request = False

            if is_appointment_request or is_human_service_request:
                logging.info(f"5.2: Appointment/Human service detected - using AI auto-reply - appointment: {is_appointment_request}, human_service: {is_human_service_request}")
                
                # Save user message first
                user_message_id = self.history_manager.save_message(phone_number, message, "user")
                
                # Update session flags for tracking (keep original functionality)
                if is_appointment_request:
                    self.history_manager.update_session_flags(phone_number, appointment_section=1)
                    logging.info(f"Updated appointment_section flag for {phone_number}")
                else:
                    self.history_manager.update_session_flags(phone_number, human_section=1)
                    logging.info(f"Updated human_section flag for {phone_number}")
                
                # Use AI auto-reply instead of manual database reply
                ai_response = self._process_normal_ai_response(message, phone_number, detected_language, user_message_id)
                
                if ai_response:
                    logging.info(f"AI auto-reply generated for {phone_number}: {ai_response[:100]}...")
                    return ai_response
                else:
                    # Fallback if AI fails
                    logging.warning(f"AI auto-reply failed for {phone_number}, using fallback message")
                    fallback_message = self._get_fallback_message_for_intent(detected_language, is_appointment_request)
                    self.history_manager.save_message(phone_number, fallback_message, "assistant")
                    if user_message_id:
                        self.history_manager.mark_message_processed(user_message_id)
                    return fallback_message

            # 5.3: Special reply keyword check (THIRD PRIORITY)
            special_reply_id = self.match_special_reply_by_keywords_with_cooldown(message, detected_language, phone_number)
            
            if special_reply_id:
                logging.info(f"5.3: Special reply keyword matched: {special_reply_id}")
                
                # Save user message
                user_message_id = self.history_manager.save_message(phone_number, message, "user")
                
                # Handle special reply sequence
                special_result = self.handle_special_reply_sequence(
                    phone_number, 
                    detected_language, 
                    special_reply_id, 
                    user_message_id
                )
                
                if special_result and special_result.get('sent', 0) > 0:
                    logging.info(f"Special reply sent successfully")
                    return None  # Special reply sent, no additional response needed
                else:
                    logging.warning(f"Special reply failed, falling back to AI")
                    # Continue to normal AI processing
            
            # 5.4: Normal AI processing (LOWEST PRIORITY - only if all above are skipped)
            logging.info(f"5.4: Processing with normal AI response")
            return self.get_normal_ai_response(message, phone_number, detected_language)
            
        except Exception as e:
            logging.error(f"Error in process_user_message: {e}")
            
            # Fallback error message
            if 'detected_language' in locals():
                if detected_language == 'zh':
                    return "抱歉，处理您的消息时出现错误。请重试。"
                elif detected_language == 'ms':
                    return "Maaf, terdapat ralat memproses mesej anda. Sila cuba lagi."
                else:
                    return "Sorry, there was an error processing your message. Please try again."
            else:
                return "Sorry, there was an error processing your message. Please try again."

    # FIXED: Simplified AI response method (rename the existing get_ai_response to this)
    def get_normal_ai_response(self, message, phone_number, detected_language):
        """
        FIXED: Handle normal AI responses without duplication
        This replaces the complex get_ai_response logic
        """
        # Check human_reply flag first
        session_flags = self.history_manager.get_session_flags(phone_number)
        if session_flags.get('human_reply', 0) == 1:
            logging.info(f"Human reply active for {phone_number}, skipping AI")
            user_message_id = self.history_manager.save_message(phone_number, message, "user")
            if user_message_id:
                self.history_manager.mark_message_processed(user_message_id)
            return None
        
        # Check manual reply flags
        if session_flags.get('manual_reply', 0) == 1:
            logging.info(f"Manual reply active for {phone_number}")
            user_message_id = self.history_manager.save_message(phone_number, message, "user")
            
            # Determine reply type
            if session_flags.get('appointment_section', 0) == 1:
                manual_reply = self.get_manual_reply_message(detected_language, 'appointment')
            elif session_flags.get('human_section', 0) == 1:
                manual_reply = self.get_manual_reply_message(detected_language, 'human_service')
            else:
                manual_reply = self.get_manual_reply_message(detected_language, 'general')
            
            self.history_manager.save_message(phone_number, manual_reply, "assistant")
            if user_message_id:
                self.history_manager.mark_message_processed(user_message_id)
            
            return manual_reply
        
        # Normal AI processing
        return self._process_normal_ai_response(message, phone_number, detected_language, None)

    def _process_normal_ai_response(self, message, phone_number, detected_language, user_message_id):
        """Process normal AI response (extracted from original get_ai_response method)"""
        
        # Check language support
        if not self.is_supported_language(message):
            unsupported_msg = self.language_detector.get_error_message('unsupported_language', detected_language)
            try:
                self.history_manager.save_message(phone_number, unsupported_msg, "assistant")
            except:
                logging.warning("Failed to save assistant response, but continuing")
            if user_message_id:
                try:
                    self.history_manager.mark_message_processed(user_message_id)
                except:
                    logging.warning("Failed to mark message as processed")
            return unsupported_msg
        
        # Enhanced sensitive words check with auto-translation
        if self.contains_sensitive_words(message, detected_language):
            sensitive_msg = self.get_sensitive_content_reply(detected_language)
            try:
                self.history_manager.save_message(phone_number, sensitive_msg, "assistant")
            except:
                logging.warning("Failed to save assistant response, but continuing")
            if user_message_id:
                try:
                    self.history_manager.mark_message_processed(user_message_id)
                except:
                    logging.warning("Failed to mark message as processed")
            return sensitive_msg
    
        # Check rate limiting
        if self.is_rate_limited(phone_number):
            rate_msg = self.language_detector.get_error_message('rate_limit', detected_language)
            try:
                self.history_manager.save_message(phone_number, rate_msg, "assistant")
            except:
                logging.warning("Failed to save assistant response, but continuing")
            if user_message_id:
                try:
                    self.history_manager.mark_message_processed(user_message_id)
                except:
                    logging.warning("Failed to mark message as processed")
            return rate_msg
        
        # Save user message if not already saved
        if not user_message_id:
            user_message_id = self.history_manager.save_message(phone_number, message, "user")
        
        # Initialize or get chat history
        if phone_number not in chat_histories:
            chat_histories[phone_number] = []
        
        # Clean old history
        self.clean_chat_history(phone_number)
        
        # Add user message to history
        chat_histories[phone_number].append({"role": "user", "content": message})
        
        # Update last activity time
        user_last_activity[phone_number] = datetime.now()
        
        # Prepare OpenAI API messages with language-specific prompt
        base_system_prompt = self.system_prompt

        # Add additional content
        if hasattr(self, 'additional_content') and self.additional_content:
            base_system_prompt += f"\n\n{self.additional_content}"
        
        # Add files content
        if hasattr(self, 'files_content') and self.files_content:
            base_system_prompt += f"\n\nAdditional Reference Materials:\n{self.files_content}"
        
        current_myt_time = get_malaysia_time()
        current_time_str = current_myt_time.strftime('%Y-%m-%d %H:%M')
        current_day_of_week = current_myt_time.strftime('%A')  # Monday, Tuesday, etc.
        
        base_system_prompt += f"\n\n【重要】当前真实系统时间（马来西亚时间 MYT）：{current_time_str}"
        base_system_prompt += f"\n今天是：{current_day_of_week}"
        base_system_prompt += f"\n\n你必须使用上述时间作为\"当前时间\"进行所有日期和时间相关的计算和判断。"
        base_system_prompt += f"\n当用户询问\"今天\"、\"明天\"或使用相对时间表达时，请基于上述时间计算。"
    
        # Add language-specific instructions to the system prompt
        system_prompt_with_language = self.language_detector.get_language_specific_prompt(
            base_system_prompt, detected_language
        )

        messages = [
            {"role": "system", "content": system_prompt_with_language},
        ] + chat_histories[phone_number][-10:]  
        
        try:
            # Call OpenAI API
            client = openai.OpenAI(api_key=OPENAI_API_KEY)
            response = client.chat.completions.create(
                model=DEFAULT_OPENAI_MODEL,
                messages=messages,
                max_tokens=1200,
                temperature=0.7
            )
            
            # Extract assistant reply
            ai_message = response.choices[0].message.content.strip()
            
            # Add AI reply to history
            chat_histories[phone_number].append({"role": "assistant", "content": ai_message})
            
            # Try to save AI response to database
            try:
                self.history_manager.save_message(phone_number, ai_message, "assistant")
            except Exception as save_error:
                logging.warning(f"Failed to save AI response: {save_error}")
            
            # Mark user message as processed after successful AI response
            if user_message_id:
                try:
                    self.history_manager.mark_message_processed(user_message_id)
                except Exception as mark_error:
                    logging.warning(f"Failed to mark message as processed: {mark_error}")
            
            logging.info(f"AI response generated for: {phone_number} in language: {detected_language}")
            return ai_message
        
        except openai.RateLimitError:
            logging.warning(f"OpenAI rate limit: {phone_number}")
            error_msg = self.language_detector.get_error_message('service_busy', detected_language)
            try:
                self.history_manager.save_message(phone_number, error_msg, "assistant")
                if user_message_id:
                    self.history_manager.mark_message_processed(user_message_id)
            except:
                logging.warning("Failed to save error response to database")
            return error_msg
        
        except openai.APITimeoutError:
            logging.warning(f"OpenAI timeout: {phone_number}")
            error_msg = self.language_detector.get_error_message('timeout', detected_language)
            try:
                self.history_manager.save_message(phone_number, error_msg, "assistant")
                if user_message_id:
                    self.history_manager.mark_message_processed(user_message_id)
            except:
                logging.warning("Failed to save error response to database")
            return error_msg
        
        except Exception as e:
            logging.error(f"Error getting AI response: {phone_number}: {e}")
            error_msg = self.language_detector.get_error_message('service_unavailable', detected_language)
            try:
                self.history_manager.save_message(phone_number, error_msg, "assistant")
                if user_message_id:
                    self.history_manager.mark_message_processed(user_message_id)
            except:
                logging.warning("Failed to save error response to database")
            return error_msg

    def _normalize_text_for_matching(self, text: str) -> str:
        """Normalize text for consistent matching"""
        import re
        
        if not text:
            return ""
        
        # Convert to lowercase
        text = text.lower().strip()
        
        # Keep Chinese characters, alphanumeric, and spaces
        # Remove other punctuation but keep essential characters
        text = re.sub(r"[^\w\s\u4e00-\u9fff]", " ", text)
        
        # Normalize multiple spaces
        text = re.sub(r"\s+", " ", text).strip()
        
        return text

    def _tokenize_text(self, text: str) -> set:
        """
        改进的文本分词方法，对中文更友好
        """
        normalized = self._normalize_text_for_matching(text)
        if not normalized:
            return set()
        
        tokens = set()
        
        # 对于包含中文的文本，既按空格分割，也按字符分割
        if any('\u4e00' <= char <= '\u9fff' for char in normalized):
            # 按空格分割（处理混合内容）
            space_tokens = normalized.split()
            tokens.update(space_tokens)
            
            # 对于纯中文部分，按字符分割（但组合成有意义的词）
            chinese_chars = ''.join(char if '\u4e00' <= char <= '\u9fff' else ' ' for char in normalized)
            chinese_words = chinese_chars.split()
            for word in chinese_words:
                if len(word) >= 1:  # 中文单字也算一个token
                    tokens.add(word)
                    # 同时添加双字组合
                    if len(word) >= 2:
                        for i in range(len(word) - 1):
                            tokens.add(word[i:i+2])
        else:
            # 非中文文本按空格分割
            tokens.update(normalized.split())
        
        return tokens

    def _translate_to_all_languages(self, text: str, source_language: str = None) -> set:
        """
        Translate text to all supported languages and return variants
        Includes original text and all translations
        """
        if not text or not text.strip():
            return set()
        
        # Always include the original normalized text
        variants = {self._normalize_text_for_matching(text)}
        
        # Supported target languages
        target_languages = ['en', 'zh', 'ms']
        
        for target_lang in target_languages:
            # Skip if source and target are the same
            if source_language and source_language == target_lang:
                continue
                
            try:
                # Use the existing translation method
                translated = self.language_detector.translate_text(text, target_lang)
                if translated and translated.strip():
                    normalized_translation = self._normalize_text_for_matching(translated)
                    if normalized_translation:
                        variants.add(normalized_translation)
                        
            except Exception as e:
                # Log but don't stop - continue with other translations
                logging.debug(f"Translation failed for '{text}' to {target_lang}: {e}")
                continue
        
        return variants

    def _get_cached_keyword_variants(self, keyword: str) -> set:
        """Get multilingual variants of keyword with caching for performance"""
        
        # Initialize cache if not exists
        if not hasattr(self, '_keyword_variants_cache'):
            self._keyword_variants_cache = {}
        
        # Use normalized keyword as cache key
        cache_key = self._normalize_text_for_matching(keyword)
        
        # Return cached result if available
        if cache_key in self._keyword_variants_cache:
            return self._keyword_variants_cache[cache_key]
        
        # Generate variants for all languages
        variants = self._translate_to_all_languages(keyword)
        
        # Cache the result
        self._keyword_variants_cache[cache_key] = variants
        
        return variants

    def _calculate_keyword_match_score(self, user_message_variants: set, keyword_variants: set) -> int:
        """
        Calculate matching score between user message and keyword variants
        Returns: 0 (no match), 1 (substring match), or 3 (word boundary match)
        """
        
        # Get tokenized versions for word boundary matching
        user_tokens = set()
        for variant in user_message_variants:
            user_tokens.update(self._tokenize_text(variant))
        
        keyword_tokens = set()
        for variant in keyword_variants:
            keyword_tokens.update(self._tokenize_text(variant))
        
        # Priority 1: Word boundary matching (exact token match)
        if keyword_tokens and user_tokens and (keyword_tokens & user_tokens):
            return 3  # High score for exact word matches
        
        # Priority 2: Substring matching
        for kw_variant in keyword_variants:
            if not kw_variant:
                continue
            for user_variant in user_message_variants:
                if not user_variant:
                    continue
                # Check both directions: keyword in message or message in keyword
                if kw_variant in user_variant or user_variant in kw_variant:
                    return 1  # Lower score for substring matches
        
        return 0  # No match

    def _is_translation_too_generic(self, original_text: str, translated_text: str) -> bool:
        """
        AI-based detection of whether a translation is too generic for precise keyword matching
        Uses OpenAI to evaluate translation specificity
        """
        try:
            # Cache to avoid repeated API calls for same translations
            if not hasattr(self, '_generic_translation_cache'):
                self._generic_translation_cache = {}
            
            cache_key = f"{original_text}:{translated_text}"
            if cache_key in self._generic_translation_cache:
                return self._generic_translation_cache[cache_key]
            
            # Use AI to evaluate translation specificity
            client = openai.OpenAI(api_key=OPENAI_API_KEY)
            
            prompt = f"""Analyze this translation for keyword matching precision:

    Original: "{original_text}"
    Translation: "{translated_text}"

    Question: Is the English translation too generic/broad and likely to cause false positive matches when used for keyword detection?

    Consider:
    1. Is the translation a very common/general term that could match many unrelated concepts?
    2. Could this translation accidentally match content about different topics?
    3. Does the translation preserve the specific meaning of the original?

    Examples of problematic generic translations:
    - "牙医" -> "doctor" (too generic, should be "dentist")
    - "手机" -> "device" (too generic, should be "phone" or "mobile")
    - "汽车" -> "vehicle" (acceptable, but "car" is more specific)

    Response format: 
    - "GENERIC" if the translation is too broad/generic for precise matching
    - "SPECIFIC" if the translation is appropriately specific
    - Brief reason (1 sentence)

    Response:"""

            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=100,
                temperature=0.1  # Low temperature for consistent evaluation
            )
            
            ai_response = response.choices[0].message.content.strip()
            
            # Parse AI response
            is_generic = "GENERIC" in ai_response.upper()
            
            # Cache the result
            self._generic_translation_cache[cache_key] = is_generic
            
            logging.info(f"AI translation evaluation: '{original_text}' -> '{translated_text}' = {'GENERIC' if is_generic else 'SPECIFIC'}")
            logging.debug(f"AI reasoning: {ai_response}")
            
            return is_generic
            
        except Exception as e:
            logging.warning(f"AI generic translation detection failed for '{original_text}' -> '{translated_text}': {e}")
            
            # Fallback: simple heuristic check
            translated_lower = translated_text.lower().strip()
            
            # Very basic fallback for common over-generic terms
            basic_generic_terms = ['thing', 'item', 'stuff', 'service', 'product', 'place', 'person']
            is_basic_generic = any(term == translated_lower for term in basic_generic_terms)
            
            if is_basic_generic:
                logging.info(f"Fallback detected generic term: '{translated_text}'")
            
            return is_basic_generic

    def _is_precise_chinese_match(self, keyword: str, user_message: str) -> bool:
        """
        More precise Chinese matching to avoid false positives
        """
        if len(keyword) < 2:
            return False
        
        # Find all occurrences of the keyword in the message
        start = 0
        while True:
            pos = user_message.find(keyword, start)
            if pos == -1:
                break
                
            # Check if this is a meaningful match (not part of a different word)
            # For Chinese, we check surrounding characters
            before_char = user_message[pos-1] if pos > 0 else ' '
            after_char = user_message[pos+len(keyword)] if pos+len(keyword) < len(user_message) else ' '
            
            # If surrounded by non-Chinese characters or at boundaries, it's a good match
            if (not ('\u4e00' <= before_char <= '\u9fff') or 
                not ('\u4e00' <= after_char <= '\u9fff') or
                pos == 0 or pos + len(keyword) == len(user_message)):
                return True
                
            start = pos + 1
        
        return False

    def match_opening_category_by_keywords(self, original_message, detected_language='en'):
        """
        FIXED VERSION: More precise keyword matching to prevent false positives
        - AI-based generic translation detection (no hardcoded terms)
        - Stricter translation variants
        - Higher threshold for matches
        - Better Chinese character handling
        """
        try:
            # Enhanced filtering for simple messages
            message_clean = original_message.strip().lower()
            word_count = len(original_message.split())
            char_count = len(message_clean)
            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in original_message)
            
            # Skip matching if message is too simple
            is_too_simple = (
                word_count <= 1 and 
                char_count <= 5 and 
                detected_language == 'en' and 
                not has_chinese and
                message_clean.isalpha()
            )
            
            if is_too_simple:
                logging.info(f"Message too simple for keyword matching: '{original_message}'")
                return None
                    
            import re

            def normalize_text(s: str) -> str:
                s = (s or "").lower().strip()
                s = re.sub(r"[^\w\s\u4e00-\u9fff]", " ", s)
                s = re.sub(r"\s+", " ", s).strip()
                return s

            def tokenize_chinese_aware(s: str):
                """Improved tokenization for Chinese"""
                tokens = set()
                
                # For mixed content, split by spaces first
                space_tokens = s.split()
                tokens.update(space_tokens)
                
                # For Chinese content, create meaningful segments
                if any('\u4e00' <= char <= '\u9fff' for char in s):
                    chinese_chars = ''.join(char if '\u4e00' <= char <= '\u9fff' else ' ' for char in s)
                    chinese_words = [w for w in chinese_chars.split() if len(w) >= 1]
                    
                    for word in chinese_words:
                        if len(word) >= 2:  # Only add multi-character Chinese words
                            tokens.add(word)
                            
                return tokens

            def get_conservative_translation_variants(text: str, source_lang: str = None):
                """AI-powered conservative translation to avoid false positives"""
                if not text or not text.strip():
                    return set()
                
                # Always include the original normalized text
                variants = {normalize_text(text)}
                
                # CONSERVATIVE: Only translate if really necessary and with strict checks
                if source_lang == 'zh' and len(text) >= 3:  # Only translate longer Chinese text
                    try:
                        english_translation = self.language_detector.translate_text(text, 'en')
                        if english_translation and len(english_translation) >= 3:
                            # AI-based generic term detection instead of hardcoded list
                            if not self._is_translation_too_generic(text, english_translation):
                                variants.add(normalize_text(english_translation))
                            else:
                                logging.debug(f"Skipped generic translation: '{text}' -> '{english_translation}'")
                    except Exception as e:
                        logging.debug(f"Conservative translation failed for '{text}': {e}")
                
                return variants

            # Process user message with conservative approach
            msg_raw = (original_message or "").strip()
            if not msg_raw:
                return None

            msg_variants = get_conservative_translation_variants(msg_raw, detected_language)
            logging.info(f"🔍 DEBUG - Original message: '{original_message}'")
            logging.info(f"🔍 DEBUG - User message variants: {msg_variants}")
            
            msg_tokens_union = set()
            for mv in msg_variants:
                msg_tokens_union |= tokenize_chinese_aware(mv)

            # Enhanced caching for keyword variants
            if not hasattr(self, "_conservative_kw_cache"):
                self._conservative_kw_cache = {}

            def get_keyword_variants_cached(kw: str):
                if kw in self._conservative_kw_cache:
                    return self._conservative_kw_cache[kw]
                
                variants = get_conservative_translation_variants(kw, 'zh' if any('\u4e00' <= c <= '\u9fff' for c in kw) else None)
                self._conservative_kw_cache[kw] = variants
                return variants

            # Process opening categories (excluding default ones)
            categories = self.history_manager.get_opening_categories(default_only=0)
            best_cat, best_score = None, 0

            for cat in categories:
                raw = (cat.get('KEYWORDS') or '').replace('\n', ',').replace(';', ',')
                kws = [k.strip() for k in raw.split(',') if k.strip()]
                if not kws:
                    continue
                
                logging.info(f"🔍 DEBUG - Checking category {cat.get('PKKEY')}: {cat.get('TITLE')}")
                logging.info(f"🔍 DEBUG - Keywords: {kws}")
            
                category_score = 0
                matched_keywords = []
                
                for kw in kws:
                    kw_variants = get_keyword_variants_cached(kw)
                    logging.info(f"🔍 DEBUG - Keyword '{kw}' variants: {kw_variants}")
                    
                    keyword_score = 0
                    
                    # PRIORITY 1: Exact Chinese character sequence match (highest precision)
                    if any('\u4e00' <= char <= '\u9fff' for char in kw) and len(kw) >= 2:
                        for mv in msg_variants:
                            if kw in mv and len(kw) >= 2:
                                # Additional verification: ensure it's not a substring of a different word
                                if self._is_precise_chinese_match(kw, mv):
                                    keyword_score = 3  # Highest score for exact Chinese match
                                    logging.info(f"✅ Exact Chinese match: '{kw}' in '{mv}'")
                                    break
                    
                    # PRIORITY 2: Token boundary matching (medium precision)  
                    if keyword_score == 0:
                        kw_tokens = set()
                        for kv in kw_variants:
                            kw_tokens |= tokenize_chinese_aware(kv)
                        
                        if kw_tokens and msg_tokens_union:
                            # STRICTER: Require significant overlap (at least 70% of keyword tokens)
                            matching_tokens = kw_tokens & msg_tokens_union
                            if matching_tokens and len(matching_tokens) >= len(kw_tokens) * 0.7:
                                keyword_score = 2
                                logging.info(f"✅ Token match: {matching_tokens} (threshold: 70%)")

                    # PRIORITY 3: Substring matching (lowest precision, very strict)
                    if keyword_score == 0 and len(kw) >= 3:  # Only for longer keywords
                        for kv in kw_variants:
                            if not kv or len(kv) < 3:
                                continue
                            for mv in msg_variants:
                                if not mv:
                                    continue
                                # VERY STRICT: keyword must be substantial portion of message
                                if kv in mv and len(kv) / len(mv) >= 0.4:  # At least 40% of message
                                    keyword_score = 1
                                    logging.info(f"✅ Strict substring match: '{kv}' in '{mv}' (ratio: {len(kv)/len(mv):.2f})")
                                    break
                            if keyword_score > 0:
                                break

                    if keyword_score > 0:
                        category_score += keyword_score
                        matched_keywords.append(kw)
                        logging.info(f"Keyword '{kw}' matched with score {keyword_score}")

                # INCREASED THRESHOLD: Require higher score to prevent false positives
                if category_score >= 2:  # Increased from 1 to 2
                    logging.info(f"Category {cat.get('PKKEY')} total score: {category_score}, matched keywords: {matched_keywords}")
                    
                    if category_score > best_score:
                        best_score = category_score
                        best_cat = cat

            return best_cat['PKKEY'] if (best_cat and best_score >= 2) else None

        except Exception as e:
            logging.error(f"match_opening_category_by_keywords error: {e}")
            return None
        
    def get_special_replies(self):
        """Get all active special replies from database"""
        return self.history_manager.get_special_replies()

    def match_special_reply_by_keywords_with_cooldown(self, original_message: str, detected_language: str, phone_number: str):
        """
        Special reply matching with per-user cooldown timer support
        Each user has independent cooldown timers for each special reply
        """
        try:
            # Validate input
            if not original_message or not original_message.strip():
                return None

            user_message_normalized = self._normalize_text_for_matching(original_message)
            if not user_message_normalized:
                return None

            user_message_variants = self._get_limited_translation_variants(original_message, detected_language)
            
            if not user_message_variants:
                return None

            # Get all special replies from database
            special_replies = self.get_special_replies()
            if not special_replies:
                return None

            best_special_reply = None
            best_total_score = 0
            cooldown_blocked_replies = []
            
            logging.info(f"Matching message for user {phone_number}: '{original_message}' (language: {detected_language})")

            # Evaluate each special reply
            for special_reply in special_replies:
                special_reply_id = special_reply.get('PKKEY')
                
                # Check per-user cooldown timer
                is_in_cooldown, remaining_hours = self.history_manager.is_special_reply_in_cooldown_for_user(phone_number, special_reply_id)
                
                if is_in_cooldown:
                    cooldown_blocked_replies.append({
                        'id': special_reply_id, 
                        'remaining_hours': remaining_hours
                    })
                    logging.info(f"Special reply {special_reply_id} blocked by per-user cooldown for {phone_number} ({remaining_hours:.1f}h remaining)")
                    continue
                
                # Parse keywords from the database field
                keywords_raw = (special_reply.get('KEYWORDS') or '').replace('\n', ',').replace(';', ',')
                keywords = [kw.strip() for kw in keywords_raw.split(',') if kw.strip()]
                
                if not keywords:
                    continue

                total_score = 0
                matched_keywords = []
                
                # Check each keyword with existing strict matching
                for keyword in keywords:
                    match_score = self._calculate_strict_keyword_match_score(
                        user_message_variants, keyword, detected_language
                    )
                    
                    if match_score > 0:
                        total_score += match_score
                        matched_keywords.append(keyword)
                        logging.info(f"Keyword '{keyword}' matched with score {match_score} for special reply ID {special_reply_id}")

                # Only consider if there's a clear match
                if total_score >= 2:
                    logging.info(f"Special reply ID {special_reply_id} total score: {total_score}, matched keywords: {matched_keywords}")
                    
                    if total_score > best_total_score:
                        best_total_score = total_score
                        best_special_reply = special_reply

            # Log cooldown status for debugging
            if cooldown_blocked_replies:
                logging.info(f"User {phone_number} has {len(cooldown_blocked_replies)} special replies in cooldown")

            if best_special_reply and best_total_score >= 2:
                best_id = best_special_reply.get('PKKEY')
                logging.info(f"Best special reply match for user {phone_number}: ID {best_id} with score {best_total_score}")
                return best_id
            
            return None

        except Exception as e:
            logging.error(f"Error in per-user cooldown special reply matching: {e}")
            return None

    def _get_limited_translation_variants(self, text: str, source_language: str = None) -> set:
        """
        获取有限的翻译变体，减少误匹配可能性
        """
        if not text or not text.strip():
            return set()
        
        # 总是包含原文的标准化版本
        variants = {self._normalize_text_for_matching(text)}
        
        # 如果检测到的语言不是英文，只翻译成英文（减少变体数量）
        if source_language and source_language != 'en':
            try:
                english_translation = self.language_detector.translate_text(text, 'en')
                if english_translation and english_translation.strip():
                    variants.add(self._normalize_text_for_matching(english_translation))
            except Exception as e:
                logging.debug(f"Translation to English failed for '{text}': {e}")
        
        # 如果原文是英文，翻译成用户的主要语言（中文或马来文）
        elif source_language == 'en':
            # 只翻译成中文，避免过多变体
            try:
                chinese_translation = self.language_detector.translate_text(text, 'zh')
                if chinese_translation and chinese_translation.strip():
                    variants.add(self._normalize_text_for_matching(chinese_translation))
            except Exception as e:
                logging.debug(f"Translation to Chinese failed for '{text}': {e}")
        
        return variants
    
    def _calculate_strict_keyword_match_score(self, user_message_variants: set, keyword: str, detected_language: str = None) -> int:
        """
        更严格的关键词匹配评分，减少误匹配
        Returns: 0 (no match), 1 (partial match), 2 (good match), 3 (exact match)
        """
        if not keyword or not keyword.strip():
            return 0
        
        # 标准化关键词
        keyword_normalized = self._normalize_text_for_matching(keyword)
        if not keyword_normalized:
            return 0
        
        # 获取关键词的有限翻译变体
        keyword_variants = self._get_limited_translation_variants(keyword, detected_language)
        
        # 优先级1: 完全精确匹配（整个关键词完全匹配用户消息的某个部分）
        for kw_variant in keyword_variants:
            if not kw_variant:
                continue
            for user_variant in user_message_variants:
                if not user_variant:
                    continue
                
                # 完全匹配：关键词就是整个用户消息，或者是用户消息中的完整词
                if kw_variant == user_variant:
                    logging.debug(f"🎯 Exact full match: '{kw_variant}' == '{user_variant}'")
                    return 3
                
                # 词边界匹配：关键词作为完整的词出现在用户消息中
                if self._is_complete_word_match(kw_variant, user_variant):
                    logging.debug(f"🎯 Complete word match: '{kw_variant}' in '{user_variant}'")
                    return 3
        
        # 优先级2: Token 级别的精确匹配
        user_tokens = set()
        for variant in user_message_variants:
            user_tokens.update(self._tokenize_text(variant))
        
        keyword_tokens = set()
        for variant in keyword_variants:
            keyword_tokens.update(self._tokenize_text(variant))
        
        if keyword_tokens and user_tokens:
            # 检查关键词的所有token是否都在用户消息中
            if keyword_tokens.issubset(user_tokens):
                logging.debug(f"🎯 All keyword tokens match: {keyword_tokens} ⊆ {user_tokens}")
                return 2
            
            # 检查是否有重要的token匹配（至少50%的关键词token匹配）
            matching_tokens = keyword_tokens & user_tokens
            if matching_tokens and len(matching_tokens) >= len(keyword_tokens) * 0.5:
                logging.debug(f"🎯 Significant token overlap: {matching_tokens}")
                return 1
        
        # 优先级3: 严格的子串匹配（只有在关键词是用户消息的明确子串时才匹配）
        for kw_variant in keyword_variants:
            if not kw_variant or len(kw_variant) < 2:  # 避免太短的关键词造成误匹配
                continue
            for user_variant in user_message_variants:
                if not user_variant:
                    continue
                
                # 只检查关键词是否在用户消息中（不检查反向）
                if kw_variant in user_variant and len(kw_variant) >= 2:
                    # 额外检查：确保这不是偶然的字符重叠
                    if self._is_meaningful_substring_match(kw_variant, user_variant):
                        logging.debug(f"🎯 Meaningful substring match: '{kw_variant}' in '{user_variant}'")
                        return 1
        
        return 0  # No meaningful match
    
    
    def _is_complete_word_match(self, keyword: str, user_message: str) -> bool:
        """
        检查关键词是否作为完整的词出现在用户消息中
        对中文做更宽松的处理：只要包含即可（关键字需 >=2 个字符），以适配“我要洗牙”等句式。
        """
        import re
        
        # 对于中文：放宽规则——包含即视为匹配（避免被前后汉字“夹住”导致匹配失败）
        if any('\u4e00' <= char <= '\u9fff' for char in keyword):
            if len(keyword) >= 2 and keyword in user_message:
                return True
            return False
        
        # 英/马文：仍按单词边界匹配
        pattern = r'\b' + re.escape(keyword) + r'\b'
        return re.search(pattern, user_message) is not None

    def _is_meaningful_substring_match(self, keyword: str, user_message: str) -> bool:
        """
        检查子串匹配是否有意义（避免偶然的字符重叠）
        """
        # 关键词太短，可能造成误匹配
        if len(keyword) < 2:
            return False
        
        # 关键词长度相对于用户消息来说应该是合理的
        if len(keyword) / len(user_message) < 0.3:  # 关键词至少应该是消息长度的30%
            return True
        
        # 如果关键词很长且几乎就是整个消息，也是有意义的匹配
        return True

    def download_media_file(self, media_id):
        """Download media file from WhatsApp Business API"""
        try:
            # First, get the media URL
            media_url = f"https://graph.facebook.com/{META_API_VERSION}/{media_id}"
            headers = {
                "Authorization": f"Bearer {self.current_token}"
            }
            
            response = requests.get(media_url, headers=headers, timeout=30)
            
            if response.status_code == 200:
                media_info = response.json()
                file_url = media_info.get('url')
                mime_type = media_info.get('mime_type', '')
                
                if file_url:
                    # Download the actual file
                    file_response = requests.get(file_url, headers=headers, timeout=60)
                    
                    if file_response.status_code == 200:
                        return file_response.content, mime_type
                    else:
                        logger.error(f"Failed to download media file: {file_response.status_code}")
                        return None, None
                else:
                    logger.error("No file URL in media response")
                    return None, None
            else:
                logger.error(f"Failed to get media info: {response.status_code} - {response.text}")
                return None, None
                
        except Exception as e:
            logger.error(f"Error downloading media file: {e}")
            return None, None
    
    def transcribe_audio(self, audio_content, mime_type):
        """Transcribe audio using OpenAI Whisper API"""
        try:
            # Determine file extension based on mime type
            extension_map = {
                'audio/ogg': '.ogg',
                'audio/mpeg': '.mp3',
                'audio/mp4': '.m4a',
                'audio/amr': '.amr',
                'audio/wav': '.wav'
            }
            
            file_extension = extension_map.get(mime_type, '.ogg')  # Default to .ogg for WhatsApp voice messages
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
                temp_file.write(audio_content)
                temp_file_path = temp_file.name
            
            try:
                # Transcribe using OpenAI Whisper
                client = openai.OpenAI(api_key=OPENAI_API_KEY)
                
                with open(temp_file_path, 'rb') as audio_file:
                    transcript = client.audio.transcriptions.create(
                        model="whisper-1",
                        file=audio_file,
                        response_format="text"
                    )
                
                logger.info(f"Audio transcription successful: {transcript[:100]}...")
                return transcript.strip()
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            logger.error(f"Error transcribing audio: {e}")
            return None
    
    def handle_media_message_with_text(self, media_message, phone_number, media_type):
        """Enhanced media message handling with manual reply system"""
        try:
            media_id = media_message.get('id')
            media_caption = media_message.get('caption', '')
            
            if not media_id:
                logging.error(f"No media ID found in {media_type} message")
                return f"❌ Sorry, I couldn't process your {media_type}. Please try again."
            
            logging.info(f"Processing {media_type} message with ID: {media_id}")
            if media_caption:
                logging.info(f"Media has caption: {media_caption[:100]}...")
            
            # Detect language from caption if available
            detected_language = 'en'
            if media_caption.strip():
                detected_language = self.detect_message_language(media_caption)
                logging.info(f"Detected language from caption: {detected_language}")
            
            # Set flags for media message - this triggers manual reply
            logging.info(f"Media message received, setting human_section and manual_reply flags for {phone_number}")
            self.history_manager.update_session_flags(
                phone_number, 
                manual_reply=1, 
                human_section=1
            )
            
            # Download and save the media file
            media_content, mime_type = self.download_media_file(media_id)
            
            if not media_content:
                error_msg = "Sorry, I couldn't download your media. Please try sending it again."
                if detected_language == 'zh':
                    error_msg = "抱歉，我无法下载您的媒体文件。请重新发送。"
                elif detected_language == 'ms':
                    error_msg = "Maaf, saya tidak dapat memuat turun media anda. Sila hantar semula."
                return f"❌ {error_msg}"
            
            logging.info(f"Downloaded {media_type} file, mime type: {mime_type}, size: {len(media_content)} bytes")
            
            # Save media file to disk
            file_path, filename, file_hash, file_size = self.history_manager.save_media_file(
                media_content, mime_type, phone_number, media_type
            )
            
            if not file_path:
                error_msg = "Sorry, I couldn't save your media. Please try again."
                if detected_language == 'zh':
                    error_msg = "抱歉，我无法保存您的媒体文件。请重试。"
                elif detected_language == 'ms':
                    error_msg = "Maaf, saya tidak dapat menyimpan media anda. Sila cuba lagi."
                return f"❌ {error_msg}"
            
            # Extract text for documents
            extracted_text = None
            if media_type == 'document':
                extracted_text = self.file_processor.extract_text_from_file(media_content, mime_type)
                if extracted_text:
                    logging.info(f"Extracted text from document: {len(extracted_text)} characters")
            
            # Get original filename if available
            original_filename = media_message.get('filename', filename)
            
            # Save the media message with caption
            try:
                media_message_id = self.history_manager.save_message_with_media_and_caption(
                    phone_number=phone_number,
                    message=f"[{media_type.title()}: {original_filename}]",
                    role="user",
                    message_type=media_type,
                    media_file_path=file_path,
                    media_file_name=filename,
                    media_mime_type=mime_type,
                    media_file_size=file_size,
                    media_hash=file_hash,
                    extracted_text=extracted_text,
                    media_caption=media_caption if media_caption else None
                )
                
                if media_message_id:
                    # Mark as processed since we're sending manual reply
                    self.history_manager.mark_message_processed(media_message_id)
            except Exception as save_error:
                logging.error(f"Error saving media message: {save_error}")
            
            logging.info(f"Saved {media_type} message from {phone_number}: {filename}")
            
            # Return manual reply message for media
            manual_reply = self.get_manual_reply_message(detected_language, 'media')
            
            # Save manual reply
            self.history_manager.save_message(phone_number, manual_reply, "assistant")
            
            return manual_reply
                
        except Exception as e:
            logging.error(f"Error handling {media_type} message: {e}")
            import traceback
            logging.error(f"Full traceback: {traceback.format_exc()}")
            
            # Error message in appropriate language
            if detected_language == 'zh':
                error_msg = "抱歉，处理您的媒体文件时出错。请尝试发送文本消息。"
            elif detected_language == 'ms':
                error_msg = "Maaf, terdapat ralat semasa memproses media anda. Sila cuba hantar mesej teks."
            else:
                error_msg = "Sorry, there was an error processing your media. Please try sending a text message instead."
            
            return f"❌ {error_msg}"
    
    def handle_audio_message(self, audio_message, phone_number):
        """Handle audio message by transcribing and responding with language detection - FIXED to prevent duplicate saves"""
        try:
            media_id = audio_message.get('id')
            if not media_id:
                logger.error("No media ID found in audio message")
                return "Sorry, I couldn't process your voice message. Please try again."
            
            logger.info(f"Processing audio message with ID: {media_id}")
            
            # Download the audio file
            audio_content, mime_type = self.download_media_file(media_id)
            
            if not audio_content:
                return "Sorry, I couldn't download your voice message. Please try sending it again."
            
            logger.info(f"Downloaded audio file, mime type: {mime_type}, size: {len(audio_content)} bytes")
            
            # Save audio file to disk
            file_path, filename, file_hash, file_size = self.history_manager.save_media_file(
                audio_content, mime_type, phone_number, 'audio'
            )
            
            if not file_path:
                return "Sorry, I couldn't save your voice message. Please try again."
            
            # Transcribe the audio
            transcribed_text = self.transcribe_audio(audio_content, mime_type)
            
            # Detect language from transcribed text
            detected_language = self.detect_message_language(transcribed_text)
            logging.info(f"Detected language from audio transcription: {detected_language}")
            
            # Check if this is a first-time user BEFORE processing
            is_first_time = self.history_manager.is_first_time_user(phone_number)
            
            # Determine language for opening sequence
            if detected_language not in self.supported_languages:
                opening_language = 'en'  # Use English for unsupported languages
            else:
                opening_language = detected_language
            
            if is_first_time:
                logging.info(f"First-time user detected: {phone_number}, starting opening sequence in {opening_language}")
                try:
                    opening_result = self.handle_opening_sequence(phone_number, opening_language)  # ✅ Pass the language parameter
                    if opening_result:
                        logging.info(f"Opening sequence completed for first-time audio user: {phone_number}")
                        time.sleep(1)
                except Exception as opening_error:
                    logging.error(f"Error in opening sequence: {opening_error}")
                
            if not transcribed_text:
                # Save audio message without transcription and mark as processed
                audio_message_id = self.history_manager.save_message_with_media(
                    phone_number=phone_number,
                    message="[Audio message - transcription failed]",
                    role="user",
                    message_type="audio",
                    media_file_path=file_path,
                    media_file_name=filename,
                    media_mime_type=mime_type,
                    media_file_size=file_size,
                    media_hash=file_hash,
                    extracted_text=None
                )
                
                # Mark as processed immediately since we can't process it further
                if audio_message_id:
                    self.history_manager.mark_message_processed(audio_message_id)
                
                return "Sorry, I couldn't understand your voice message. The audio was saved, but please try speaking more clearly or send a text message."
            
            logging.info(f"Transcribed audio from {phone_number}: {transcribed_text}")
            
            # Detect language from transcribed text
            detected_language = self.detect_message_language(transcribed_text)
            logging.info(f"Detected language from audio transcription: {detected_language}")
            
            # SAVE THE AUDIO MESSAGE ONLY ONCE HERE with processed=FALSE
            audio_message_id = self.history_manager.save_message_with_media(
                phone_number=phone_number,
                message=transcribed_text,  # Use transcribed text as the message content
                role="user",
                message_type="audio",
                media_file_path=file_path,
                media_file_name=filename,
                media_mime_type=mime_type,
                media_file_size=file_size,
                media_hash=file_hash,
                extracted_text=transcribed_text
            )
            
            if not audio_message_id:
                logging.error("Failed to save audio message to database")
                return "Sorry, there was an error processing your voice message."
            
            # Now process the transcribed text for AI response WITHOUT saving user message again
            try:
                # Check if this is a first-time user BEFORE processing
                is_first_time = self.history_manager.is_first_time_user(phone_number)
                
                if is_first_time:
                    logging.info(f"First-time user detected: {phone_number}, starting opening sequence")
                    try:
                        opening_result = self.handle_opening_sequence(phone_number)
                        if opening_result:
                            logging.info(f"Opening sequence completed for first-time audio user: {phone_number}")
                            time.sleep(1)
                    except Exception as opening_error:
                        logging.error(f"Error in opening sequence: {opening_error}")
                
                # Generate AI response using the enhanced method (handles language detection, sensitive content, etc.)
                ai_response = self._generate_ai_response_for_transcribed_audio(
                    transcribed_text, phone_number, detected_language
                )
                
                if ai_response:
                    # Mark the audio message as processed after successful AI response
                    self.history_manager.mark_message_processed(audio_message_id)
                    
                    # Add voice message indicator to response in appropriate language
                    if detected_language == 'zh':
                        voice_indicator = f"🎤➡️📝 我听到了: \"{transcribed_text[:50]}{'...' if len(transcribed_text) > 50 else ''}\"\n\n{ai_response}"
                    elif detected_language == 'ms':
                        voice_indicator = f"🎤➡️📝 Saya dengar: \"{transcribed_text[:50]}{'...' if len(transcribed_text) > 50 else ''}\"\n\n{ai_response}"
                    else:
                        voice_indicator = f"🎤➡️📝 I heard: \"{transcribed_text[:50]}{'...' if len(transcribed_text) > 50 else ''}\"\n\n{ai_response}"
                    
                    return voice_indicator
                else:
                    # Mark as processed even if AI response fails
                    self.history_manager.mark_message_processed(audio_message_id)
                    
                    error_msg = "Sorry, I couldn't generate a response to your voice message. Please try again."
                    if detected_language == 'zh':
                        error_msg = "抱歉，我无法对您的语音消息生成回复。请重试。"
                    elif detected_language == 'ms':
                        error_msg = "Maaf, saya tidak dapat menjana respons kepada mesej suara anda. Sila cuba lagi."
                    return f"❌ {error_msg}"
                    
            except Exception as processing_error:
                logging.error(f"Error processing transcribed audio: {processing_error}")
                # Mark as processed to prevent infinite loops
                if audio_message_id:
                    self.history_manager.mark_message_processed(audio_message_id)
                
                error_msg = "Sorry, there was an error processing your voice message."
                if detected_language == 'zh':
                    error_msg = "抱歉，处理您的语音消息时出错。"
                elif detected_language == 'ms':
                    error_msg = "Maaf, terdapat ralat semasa memproses mesej suara anda."
                return f"❌ {error_msg}"
                
        except Exception as e:
            logger.error(f"Error handling audio message: {e}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return "Sorry, there was an error processing your voice message. Please try sending a text message instead."

    def _generate_ai_response_for_transcribed_audio(self, transcribed_text, phone_number, detected_language):
        """Generate AI response for transcribed audio without saving user message again"""
        try:
            # Check language support
            if not self.is_supported_language(transcribed_text):
                unsupported_msg = self.language_detector.get_error_message('unsupported_language', detected_language)
                self.history_manager.save_message(phone_number, unsupported_msg, "assistant")
                return unsupported_msg
            
            # Enhanced sensitive words check
            if self.contains_sensitive_words(transcribed_text, detected_language):
                sensitive_msg = self.get_sensitive_content_reply(detected_language)
                self.history_manager.save_message(phone_number, sensitive_msg, "assistant")
                return sensitive_msg

            # Check rate limiting
            if self.is_rate_limited(phone_number):
                rate_msg = self.language_detector.get_error_message('rate_limit', detected_language)
                self.history_manager.save_message(phone_number, rate_msg, "assistant")
                return rate_msg
            
            # Initialize or get chat history
            if phone_number not in chat_histories:
                chat_histories[phone_number] = []
            
            # Clean old history
            self.clean_chat_history(phone_number)
            
            # Add user message to in-memory history (transcribed text)
            chat_histories[phone_number].append({"role": "user", "content": transcribed_text})
            
            # Update last activity time
            user_last_activity[phone_number] = datetime.now()
            
            # Prepare OpenAI API messages with language-specific prompt
            base_system_prompt = self.system_prompt

            # Add additional content
            if hasattr(self, 'additional_content') and self.additional_content:
                base_system_prompt += f"\n\n{self.additional_content}"
            
            # Add files content
            if hasattr(self, 'files_content') and self.files_content:
                base_system_prompt += f"\n\nAdditional Reference Materials:\n{self.files_content}"

            # Add language-specific instructions to the system prompt
            system_prompt_with_language = self.language_detector.get_language_specific_prompt(
                base_system_prompt, detected_language
            )

            messages = [
                {"role": "system", "content": system_prompt_with_language},
            ] + chat_histories[phone_number][-10:]  
            
            try:
                # Call OpenAI API
                client = openai.OpenAI(api_key=OPENAI_API_KEY)
                response = client.chat.completions.create(
                    model=DEFAULT_OPENAI_MODEL,
                    messages=messages,
                    max_tokens=1200,
                    temperature=0.7,
                    timeout=30
                )
                
                # Extract assistant reply
                ai_message = response.choices[0].message.content.strip()
                
                # Add AI reply to in-memory history
                chat_histories[phone_number].append({"role": "assistant", "content": ai_message})
                
                # Save AI response to database
                self.history_manager.save_message(phone_number, ai_message, "assistant")
                
                logger.info(f"AI response generated for audio message from: {phone_number} in language: {detected_language}")
                return ai_message
            
            except openai.RateLimitError:
                logger.warning(f"OpenAI rate limit for audio message: {phone_number}")
                error_msg = self.language_detector.get_error_message('service_busy', detected_language)
                self.history_manager.save_message(phone_number, error_msg, "assistant")
                return error_msg
            
            except openai.APITimeoutError:
                logger.warning(f"OpenAI timeout for audio message: {phone_number}")
                error_msg = self.language_detector.get_error_message('timeout', detected_language)
                self.history_manager.save_message(phone_number, error_msg, "assistant")
                return error_msg
            
            except Exception as e:
                logger.error(f"Error getting AI response for audio: {phone_number}: {e}")
                error_msg = self.language_detector.get_error_message('service_unavailable', detected_language)
                self.history_manager.save_message(phone_number, error_msg, "assistant")
                return error_msg
                
        except Exception as outer_e:
            logger.error(f"Outer error in AI response generation for audio: {outer_e}")
            return None

    def handle_special_reply_sequence(self, phone_number: str, user_language: str, special_id: int, user_message_id: int = None):
        """Enhanced special reply sequence handler with text variation for text messages"""
        try:
            # Record that this user triggered this special reply (before processing)
            usage_recorded = self.history_manager.record_special_reply_usage(phone_number, special_id)
            if not usage_recorded:
                logging.warning(f"Failed to record special reply usage for user {phone_number}, special_reply {special_id}")

            # Get special reply sub-items
            subs = self.history_manager.get_special_reply_subs(special_id)
            if not subs:
                logging.warning(f"No special_reply_sub records found for SPECIAL_ID={special_id}")
                
                if user_message_id:
                    self.history_manager.mark_message_processed(user_message_id)
                
                return {'special_id': special_id, 'sent': 0, 'user': phone_number}

            sent_count = 0
            
            logging.info(f"Processing special reply {special_id} for user {phone_number} with {len(subs)} items")
            
            for sub in subs:
                message_type = sub['TYPE']

                try:
                    if message_type == 1:  # Text message - APPLY VARIATION HERE
                        text_content = sub.get('TEXT_CONTENT') or ""
                        if not text_content:
                            continue

                        try:
                            # Translate but preserve detected address substrings
                            prepared_text = self.language_detector.translate_preserving_addresses(
                                text_content, user_language
                            )

                            # If it involves addresses or user asked for address, avoid variation
                            if self.language_detector.should_skip_address_variation(user_text, text_content):
                                text_to_send = prepared_text.strip()
                            else:
                                # Safe to vary a bit
                                try:
                                    text_to_send = self.text_variation_generator.generate_text_variation(
                                        prepared_text, user_language, 'slight'
                                    )
                                except Exception as var_error:
                                    logging.warning(f"Text variation failed (special): {var_error}")
                                    text_to_send = prepared_text

                        except Exception as e:
                            logging.warning(f"Special text preserve-translate failed: {e}")
                            text_to_send = text_content.strip()

                        res = self.send_whatsapp_message(phone_number, text_to_send)
                        if res:
                            self.history_manager.save_message(phone_number, text_to_send, "assistant", "text")
                            sent += 1

                    elif message_type == 2:  # Image - no variation needed
                        media_name = sub.get('MEDIANAME')
                        if media_name and self.special_reply_media_handler.file_exists(media_name, 'image'):
                            result, send_media_info = self.send_special_reply_media_with_private_repo_upload(phone_number, media_name, 'image')
                            if result:
                                mime_type = 'image/jpeg'  # Default
                                if media_name.lower().endswith('.png'):
                                    mime_type = 'image/png'
                                elif media_name.lower().endswith('.gif'):
                                    mime_type = 'image/gif'
                                
                                # Save media message
                                if send_media_info:
                                    message_id = self.history_manager.save_sent_special_reply_media_with_admin_role(
                                        phone_number=phone_number,
                                        message_text=f"Special Image: {media_name}",
                                        media_type="image",
                                        send_media_info=send_media_info,
                                        mime_type=mime_type
                                    )
                                    
                                sent_count += 1
                                logging.info(f"Special image sent to user {phone_number}: {media_name}")
                            else:
                                logging.error(f"Failed to send special image to user {phone_number}: {media_name}")
                        else:
                            logging.error(f"Special image not found: {media_name}")

                    elif message_type == 3:  # Video - no variation needed
                        media_name = sub.get('MEDIANAME')
                        if media_name and self.special_reply_media_handler.file_exists(media_name, 'video'):
                            result, send_media_info = self.send_special_reply_media_with_private_repo_upload(phone_number, media_name, 'video')
                            if result:
                                mime_type = 'video/mp4'  # Default
                                if media_name.lower().endswith('.mov'):
                                    mime_type = 'video/quicktime'
                                elif media_name.lower().endswith('.avi'):
                                    mime_type = 'video/avi'
                                
                                # Save media message
                                if send_media_info:
                                    message_id = self.history_manager.save_sent_special_reply_media_with_admin_role(
                                        phone_number=phone_number,
                                        message_text=f"Special Video: {media_name}",
                                        media_type="video",
                                        send_media_info=send_media_info,
                                        mime_type=mime_type
                                    )
                                sent_count += 1
                                logging.info(f"Special video sent to user {phone_number}: {media_name}")
                            else:
                                logging.error(f"Failed to send special video to user {phone_number}: {media_name}")
                        else:
                            logging.error(f"Special video not found: {media_name}")

                    # Small delay between messages
                    if sent_count > 0:
                        time.sleep(0.5)

                except Exception as item_error:
                    logging.error(f"Error processing special reply sub-item {sub.get('PKKEY', 'unknown')} for user {phone_number}: {item_error}")
                    continue

            # Mark user message as processed
            if user_message_id:
                try:
                    self.history_manager.mark_message_processed(user_message_id)
                    logging.info(f"Marked user message {user_message_id} as processed after special reply for user {phone_number}")
                except Exception as mark_error:
                    logging.warning(f"Failed to mark user message as processed: {mark_error}")

            # Get cooldown info for response
            cooldown_hours = 24  # Default
            try:
                cursor = self.history_manager.connection.cursor()
                cursor.execute('SELECT COOLDOWN_HOURS FROM special_reply WHERE PKKEY = %s', (special_id,))
                result = cursor.fetchone()
                if result:
                    cooldown_hours = result[0] or 24
                cursor.close()
            except:
                pass

            logging.info(f"Special reply sequence completed for user {phone_number}: {sent_count} messages sent (with text variations), {cooldown_hours}h cooldown applied")
            
            return {
                'special_id': special_id, 
                'sent': sent_count, 
                'total_items': len(subs),
                'user': phone_number,
                'cooldown_hours': cooldown_hours,
                'per_user_cooldown': True,
                'text_variation_applied': True
            }

        except Exception as e:
            logging.error(f"Error in handle_special_reply_sequence for user {phone_number}: {e}")
            
            if user_message_id:
                try:
                    self.history_manager.mark_message_processed(user_message_id)
                except:
                    pass
            
            return {'special_id': special_id, 'sent': 0, 'error': str(e), 'user': phone_number}

    def send_special_reply_media_with_private_repo_upload(self, to_number, media_filename, media_type):
        """FIXED: Download remote files locally before GitHub upload"""
        try:
            headers = {
                "Authorization": f"Bearer {self.current_token}",
                "Content-Type": "application/json"
            }
            
            to_number = ''.join(filter(str.isdigit, to_number))
            
            # STEP 1: Check file availability
            source_media_path = self.get_special_reply_media_path(media_filename, media_type)
            file_exists_remotely = self.special_reply_media_handler.file_exists(media_filename, media_type)
            
            if not source_media_path and not file_exists_remotely:
                logging.error(f"Special reply media file not found anywhere: {media_filename}")
                return None, None
            
            # STEP 2: NEW - Download remote file if only exists remotely
            if not source_media_path and file_exists_remotely:
                logging.info(f"File exists remotely only, downloading for GitHub upload: {media_filename}")
                
                try:
                    # Download the remote file
                    remote_url = self.special_reply_media_handler.get_media_url(media_filename, media_type)
                    response = requests.get(remote_url, timeout=30)
                    
                    if response.status_code == 200:
                        # Create local directory structure
                        local_dir = Path("media_special_reply") / media_type
                        local_dir.mkdir(parents=True, exist_ok=True)
                        
                        # Save locally
                        local_file_path = local_dir / media_filename
                        with open(local_file_path, 'wb') as f:
                            f.write(response.content)
                        
                        source_media_path = str(local_file_path)
                        logging.info(f"Successfully downloaded remote file to: {source_media_path}")
                    else:
                        logging.error(f"Failed to download remote file: {response.status_code}")
                        
                except Exception as download_error:
                    logging.error(f"Error downloading remote file: {download_error}")
                    # Continue with remote-only processing as fallback
            
            # STEP 3: Process based on file availability (prioritize local files)
            send_media_path = None
            send_filename = None
            file_hash = None
            file_size = None
            repo_uploaded = False
            file_location = "unknown"
            
            if source_media_path and os.path.exists(source_media_path):
                # Case A: File exists locally (original or downloaded) - can upload to GitHub
                file_location = "local"
                logging.info(f"Processing LOCAL file for GitHub upload: {source_media_path}")
                
                send_media_path, send_filename, file_hash, file_size = self.history_manager.special_reply_github_manager.copy_and_upload_media_to_private_repo(
                    source_media_path, to_number, media_type, media_filename
                )
                repo_uploaded = True
                
            else:
                # Case B: Still remote only (download failed) - generate metadata without upload
                file_location = "remote"
                logging.warning(f"File remains remote-only, no GitHub upload: {media_filename}")
                
                # Generate organized metadata for remote file
                import hashlib
                from datetime import datetime
                
                now = datetime.now()
                current_year = now.strftime("%Y")
                current_month = now.strftime("%m")
                current_day = now.strftime("%d")
                
                content_for_hash = f"{media_filename}_{to_number}_{time.time()}"
                file_hash = hashlib.md5(content_for_hash.encode()).hexdigest()
                
                file_extension = os.path.splitext(media_filename)[1] or ('.jpg' if media_type == 'image' else '.mp4')
                clean_filename = os.path.splitext(media_filename)[0]
                send_filename = f"{clean_filename}_{file_hash[:8]}{file_extension}"
                
                send_media_path = f"{media_type}/{current_year}/{current_month}/{current_day}/{send_filename}"
                file_size = 0
                repo_uploaded = False
            
            # STEP 4: Get media URL for WhatsApp (unchanged)
            media_url = self.special_reply_media_handler.get_media_url(media_filename, media_type)
            if not media_url:
                # Fallback to GitHub raw URL
                media_url = f"https://raw.githubusercontent.com/Welsh510/reactjs-appointmentwhatsapp/master/public/media_special_reply/{media_filename}"
            
            # STEP 5: Prepare WhatsApp payload (unchanged)
            if media_type == 'image':
                payload = {
                    "messaging_product": "whatsapp",
                    "recipient_type": "individual",
                    "to": to_number,
                    "type": "image",
                    "image": {"link": media_url, "caption": ""}
                }
            elif media_type == 'video':
                payload = {
                    "messaging_product": "whatsapp",
                    "recipient_type": "individual", 
                    "to": to_number,
                    "type": "video",
                    "video": {"link": media_url, "caption": ""}
                }
            else:
                logging.error(f"Unsupported media type: {media_type}")
                return None, None
            
            # STEP 6: Send to WhatsApp
            response = requests.post(META_API_BASE_URL, headers=headers, json=payload, timeout=30)
            
            if response.status_code == 200:
                logging.info(f"Special reply media sent successfully: {media_filename} (file_location: {file_location}, repo_uploaded: {repo_uploaded})")
                
                # Return detailed info for database saving
                send_media_info = {
                    'send_media_path': send_media_path,
                    'send_filename': send_filename, 
                    'file_hash': file_hash,
                    'file_size': file_size,
                    'original_filename': media_filename,
                    'private_repo_uploaded': repo_uploaded,
                    'file_location': file_location,
                    'media_url': media_url,
                    'github_repo': 'reactjs-appointmentwhatsapp',
                    'organized_structure': True,
                    # Compatibility keys
                    'file_path': send_media_path,
                    'filename': send_filename,
                    'hash': file_hash,
                    'size': file_size
                }
                
                logging.info(f"GitHub upload status: {repo_uploaded}, location: {file_location}, path: {send_media_path}")
                return response.json(), send_media_info
            else:
                logging.error(f"WhatsApp send failed: {response.status_code} - {response.text}")
                return None, None
                
        except Exception as e:
            logging.error(f"Error sending special reply media: {e}")
            import traceback
            logging.error(f"Full traceback: {traceback.format_exc()}")
            return None, None

    def get_special_reply_media_path(self, filename, media_type):
        """ENHANCED: Return clear status about file location"""
        import os

        # Define possible local paths
        possible_paths = [
            f"media_special_reply/{filename}",
            f"public/media_special_reply/{filename}",
            f"../reactjs-appointmentwhatsapp/public/media_special_reply/{filename}",
            f"/app/media_special_reply/{filename}",
            f"/app/public/media_special_reply/{filename}"
        ]
        
        # Try each local path
        for path in possible_paths:
            if os.path.exists(path):
                logging.info(f"Found special reply media LOCALLY: {path}")
                return path
        
        # ENHANCED: Check if file exists remotely but not locally
        if self.special_reply_media_handler.file_exists(filename, media_type):
            logging.info(f"Special reply media {filename} EXISTS REMOTELY but NOT FOUND LOCALLY")
            logging.info(f"Will serve from remote URL: {self.special_reply_media_handler.get_media_url(filename, media_type)}")
        else:
            logging.warning(f"Special reply media {filename} NOT FOUND anywhere (local or remote)")
            logging.info(f"Checked local paths: {possible_paths}")
        
        return None  # File not available locally (but may exist remotely)

chatbot = EnhancedChatBot()

def handle_image_message_in_webhook(image_data, sender_id):
    """Handle image message in webhook - process fully but check human_reply for response"""
    logger.info(f"Received image message from {sender_id}")
    
    # Save to followup keep table
    try:
        chatbot.history_manager.save_phone_number_to_followup_keep(sender_id)
    except Exception as e:
        logging.error(f"Error saving to followup keep in image handler: {e}")
    
    # Check human_reply flag
    session_flags = chatbot.history_manager.get_session_flags(sender_id)
    human_reply_active = session_flags.get('human_reply', 0) == 1
    
    if human_reply_active:
        logger.info(f"Human reply flag active for {sender_id}, will process but not auto-reply")
    
    # Detect language from caption
    caption = image_data.get('caption', '')
    detected_language = 'en'
    if caption.strip():
        detected_language = chatbot.detect_message_language(caption)
        if detected_language not in chatbot.supported_languages:
            detected_language = 'en'
    
    # Check if first-time user
    is_first_time = chatbot.history_manager.is_first_time_user(sender_id)
    
    # Handle first-time user opening sequence (skip if human_reply is active)
    if is_first_time and not human_reply_active:
        logger.info(f"First-time user sent image, starting opening sequence for {sender_id} in {detected_language}")
        opening_result = chatbot.handle_opening_sequence(sender_id, detected_language)
        if opening_result:
            # Set manual reply flags after opening
            chatbot.history_manager.update_session_flags(sender_id, manual_reply=1, human_section=1)
            logger.info(f"Opening sequence sent to first-time image user: {sender_id}")
    
    # Process the image completely
    try:
        media_id = image_data.get('id')
        if not media_id:
            logging.error(f"No media ID found in image message")
            if not human_reply_active:
                error_msg = "❌ Sorry, I couldn't process your image. Please try again."
                chatbot.send_whatsapp_message(sender_id, error_msg)
            return False
        
        logging.info(f"Processing image message with ID: {media_id}")
        if caption:
            logging.info(f"Image has caption: {caption[:100]}...")
        
        # Set flags for media message (only if human_reply not active)
        if not human_reply_active and not is_first_time:
            logging.info(f"Image message received, setting human_section and manual_reply flags for {sender_id}")
            chatbot.history_manager.update_session_flags(sender_id, manual_reply=1, human_section=1)
        
        # Download the image file
        media_content, mime_type = chatbot.download_media_file(media_id)
        
        if not media_content:
            logging.error("Failed to download image")
            if not human_reply_active:
                error_msg = "Sorry, I couldn't download your image. Please try sending it again."
                if detected_language == 'zh':
                    error_msg = "抱歉，我无法下载您的图片。请重新发送。"
                elif detected_language == 'ms':
                    error_msg = "Maaf, saya tidak dapat memuat turun imej anda. Sila hantar semula."
                chatbot.send_whatsapp_message(sender_id, f"❌ {error_msg}")
            return False
        
        logging.info(f"Downloaded image file, mime type: {mime_type}, size: {len(media_content)} bytes")
        
        # Save image file to disk and GitHub
        file_path, filename, file_hash, file_size = chatbot.history_manager.save_media_file(
            media_content, mime_type, sender_id, 'image'
        )
        
        if not file_path:
            logging.error("Failed to save image file")
            if not human_reply_active:
                error_msg = "Sorry, I couldn't save your image. Please try again."
                if detected_language == 'zh':
                    error_msg = "抱歉，我无法保存您的图片。请重试。"
                elif detected_language == 'ms':
                    error_msg = "Maaf, saya tidak dapat menyimpan imej anda. Sila cuba lagi."
                chatbot.send_whatsapp_message(sender_id, f"❌ {error_msg}")
            return False
        
        # Get original filename if available
        original_filename = image_data.get('filename', filename)
        
        # Save complete image message to database
        media_message_id = chatbot.history_manager.save_message_with_media_and_caption(
            phone_number=sender_id,
            message=f"[Image: {original_filename}]",
            role="user",
            message_type="image",
            media_file_path=file_path,
            media_file_name=filename,
            media_mime_type=mime_type,
            media_file_size=file_size,
            media_hash=file_hash,
            extracted_text=None,
            media_caption=caption if caption else None
        )
        
        if media_message_id:
            logger.info(f"Image message saved to database with ID: {media_message_id}")
            
            # Mark as processed
            chatbot.history_manager.mark_message_processed(media_message_id)
            
            # If human_reply is active, skip sending response
            if human_reply_active:
                logger.info(f"Image processed but no auto-reply sent due to human_reply flag")
                return True
            
            # Send manual reply response (only if human_reply is not active)
            manual_reply = chatbot.get_manual_reply_message(detected_language, 'media')
            chatbot.history_manager.save_message(sender_id, manual_reply, "assistant")
            result = chatbot.send_whatsapp_message(sender_id, manual_reply)
            return result is not None
        
        return False
        
    except Exception as e:
        logging.error(f"Error handling image message: {e}")
        import traceback
        logging.error(f"Full traceback: {traceback.format_exc()}")
        
        if not human_reply_active:
            error_msg = "Sorry, there was an error processing your image. Please try sending a text message instead."
            if detected_language == 'zh':
                error_msg = "抱歉，处理您的图片时出错。请尝试发送文本消息。"
            elif detected_language == 'ms':
                error_msg = "Maaf, terdapat ralat semasa memproses imej anda. Sila cuba hantar mesej teks."
            chatbot.send_whatsapp_message(sender_id, f"❌ {error_msg}")
        return False

# Updated webhook handler for video messages
def handle_video_message_in_webhook(video_data, sender_id):
    """Handle video message in webhook - process fully but check human_reply for response"""
    logger.info(f"Received video message from {sender_id}")
    
    # Save to followup keep table
    try:
        chatbot.history_manager.save_phone_number_to_followup_keep(sender_id)
    except Exception as e:
        logging.error(f"Error saving to followup keep in video handler: {e}")
    
    # Check human_reply flag
    session_flags = chatbot.history_manager.get_session_flags(sender_id)
    human_reply_active = session_flags.get('human_reply', 0) == 1
    
    if human_reply_active:
        logger.info(f"Human reply flag active for {sender_id}, will process but not auto-reply")
    
    # Detect language from caption
    caption = video_data.get('caption', '')
    detected_language = 'en'
    if caption.strip():
        detected_language = chatbot.detect_message_language(caption)
        if detected_language not in chatbot.supported_languages:
            detected_language = 'en'
    
    # Check if first-time user
    is_first_time = chatbot.history_manager.is_first_time_user(sender_id)
    
    # Handle first-time user opening sequence (skip if human_reply is active)
    if is_first_time and not human_reply_active:
        logger.info(f"First-time user sent video, starting opening sequence for {sender_id} in {detected_language}")
        opening_result = chatbot.handle_opening_sequence(sender_id, detected_language)
        if opening_result:
            # Set manual reply flags after opening
            chatbot.history_manager.update_session_flags(sender_id, manual_reply=1, human_section=1)
            logger.info(f"Opening sequence sent to first-time video user: {sender_id}")
    
    # Process the video completely
    try:
        media_id = video_data.get('id')
        if not media_id:
            logging.error(f"No media ID found in video message")
            if not human_reply_active:
                error_msg = "❌ Sorry, I couldn't process your video. Please try again."
                chatbot.send_whatsapp_message(sender_id, error_msg)
            return False
        
        logging.info(f"Processing video message with ID: {media_id}")
        if caption:
            logging.info(f"Video has caption: {caption[:100]}...")
        
        # Set flags for media message (only if human_reply not active)
        if not human_reply_active and not is_first_time:
            logging.info(f"Video message received, setting human_section and manual_reply flags for {sender_id}")
            chatbot.history_manager.update_session_flags(sender_id, manual_reply=1, human_section=1)
        
        # Download the video file
        media_content, mime_type = chatbot.download_media_file(media_id)
        
        if not media_content:
            logging.error("Failed to download video")
            if not human_reply_active:
                error_msg = "Sorry, I couldn't download your video. Please try sending it again."
                if detected_language == 'zh':
                    error_msg = "抱歉，我无法下载您的视频。请重新发送。"
                elif detected_language == 'ms':
                    error_msg = "Maaf, saya tidak dapat memuat turun video anda. Sila hantar semula."
                chatbot.send_whatsapp_message(sender_id, f"❌ {error_msg}")
            return False
        
        logging.info(f"Downloaded video file, mime type: {mime_type}, size: {len(media_content)} bytes")
        
        # Save video file to disk and GitHub
        file_path, filename, file_hash, file_size = chatbot.history_manager.save_media_file(
            media_content, mime_type, sender_id, 'video'
        )
        
        if not file_path:
            logging.error("Failed to save video file")
            if not human_reply_active:
                error_msg = "Sorry, I couldn't save your video. Please try again."
                if detected_language == 'zh':
                    error_msg = "抱歉，我无法保存您的视频。请重试。"
                elif detected_language == 'ms':
                    error_msg = "Maaf, saya tidak dapat menyimpan video anda. Sila cuba lagi."
                chatbot.send_whatsapp_message(sender_id, f"❌ {error_msg}")
            return False
        
        # Get original filename if available
        original_filename = video_data.get('filename', filename)
        
        # Save complete video message to database
        media_message_id = chatbot.history_manager.save_message_with_media_and_caption(
            phone_number=sender_id,
            message=f"[Video: {original_filename}]",
            role="user",
            message_type="video",
            media_file_path=file_path,
            media_file_name=filename,
            media_mime_type=mime_type,
            media_file_size=file_size,
            media_hash=file_hash,
            extracted_text=None,
            media_caption=caption if caption else None
        )
        
        if media_message_id:
            logger.info(f"Video message saved to database with ID: {media_message_id}")
            
            # Mark as processed
            chatbot.history_manager.mark_message_processed(media_message_id)
            
            # If human_reply is active, skip sending response
            if human_reply_active:
                logger.info(f"Video processed but no auto-reply sent due to human_reply flag")
                return True
            
            # Send manual reply response (only if human_reply is not active)
            manual_reply = chatbot.get_manual_reply_message(detected_language, 'media')
            chatbot.history_manager.save_message(sender_id, manual_reply, "assistant")
            result = chatbot.send_whatsapp_message(sender_id, manual_reply)
            return result is not None
        
        return False
        
    except Exception as e:
        logging.error(f"Error handling video message: {e}")
        import traceback
        logging.error(f"Full traceback: {traceback.format_exc()}")
        
        if not human_reply_active:
            error_msg = "Sorry, there was an error processing your video. Please try sending a text message instead."
            if detected_language == 'zh':
                error_msg = "抱歉，处理您的视频时出错。请尝试发送文本消息。"
            elif detected_language == 'ms':
                error_msg = "Maaf, terdapat ralat semasa memproses video anda. Sila cuba hantar mesej teks."
            chatbot.send_whatsapp_message(sender_id, f"❌ {error_msg}")
        return False

# Updated webhook handler for document messages
def handle_document_message_in_webhook(document_data, sender_id):
    """Handle document message in webhook - process fully but check human_reply for response"""
    logger.info(f"Received document message from {sender_id}")
    
    # Save to followup keep table
    try:
        chatbot.history_manager.save_phone_number_to_followup_keep(sender_id)
    except Exception as e:
        logging.error(f"Error saving to followup keep in document handler: {e}")
    
    # Check human_reply flag
    session_flags = chatbot.history_manager.get_session_flags(sender_id)
    human_reply_active = session_flags.get('human_reply', 0) == 1
    
    if human_reply_active:
        logger.info(f"Human reply flag active for {sender_id}, will process but not auto-reply")
    
    # Detect language from caption
    caption = document_data.get('caption', '')
    detected_language = 'en'
    if caption.strip():
        detected_language = chatbot.detect_message_language(caption)
        if detected_language not in chatbot.supported_languages:
            detected_language = 'en'
    
    # Check if first-time user
    is_first_time = chatbot.history_manager.is_first_time_user(sender_id)
    
    # Handle first-time user opening sequence (skip if human_reply is active)
    if is_first_time and not human_reply_active:
        logger.info(f"First-time user sent document, starting opening sequence for {sender_id} in {detected_language}")
        opening_result = chatbot.handle_opening_sequence(sender_id, detected_language)
        if opening_result:
            # Set manual reply flags after opening
            chatbot.history_manager.update_session_flags(sender_id, manual_reply=1, human_section=1)
            logger.info(f"Opening sequence sent to first-time document user: {sender_id}")
    
    # Process the document completely
    try:
        media_id = document_data.get('id')
        if not media_id:
            logging.error(f"No media ID found in document message")
            if not human_reply_active:
                error_msg = "❌ Sorry, I couldn't process your document. Please try again."
                chatbot.send_whatsapp_message(sender_id, error_msg)
            return False
        
        logging.info(f"Processing document message with ID: {media_id}")
        if caption:
            logging.info(f"Document has caption: {caption[:100]}...")
        
        # Set flags for media message (only if human_reply not active)
        if not human_reply_active and not is_first_time:
            logging.info(f"Document message received, setting human_section and manual_reply flags for {sender_id}")
            chatbot.history_manager.update_session_flags(sender_id, manual_reply=1, human_section=1)
        
        # Download the document file
        media_content, mime_type = chatbot.download_media_file(media_id)
        
        if not media_content:
            logging.error("Failed to download document")
            if not human_reply_active:
                error_msg = "Sorry, I couldn't download your document. Please try sending it again."
                if detected_language == 'zh':
                    error_msg = "抱歉，我无法下载您的文档。请重新发送。"
                elif detected_language == 'ms':
                    error_msg = "Maaf, saya tidak dapat memuat turun dokumen anda. Sila hantar semula."
                chatbot.send_whatsapp_message(sender_id, f"❌ {error_msg}")
            return False
        
        logging.info(f"Downloaded document file, mime type: {mime_type}, size: {len(media_content)} bytes")
        
        # Save document file to disk and GitHub
        file_path, filename, file_hash, file_size = chatbot.history_manager.save_media_file(
            media_content, mime_type, sender_id, 'document'
        )
        
        if not file_path:
            logging.error("Failed to save document file")
            if not human_reply_active:
                error_msg = "Sorry, I couldn't save your document. Please try again."
                if detected_language == 'zh':
                    error_msg = "抱歉，我无法保存您的文档。请重试。"
                elif detected_language == 'ms':
                    error_msg = "Maaf, saya tidak dapat menyimpan dokumen anda. Sila cuba lagi."
                chatbot.send_whatsapp_message(sender_id, f"❌ {error_msg}")
            return False
        
        # Extract text from document
        extracted_text = None
        if mime_type:
            extracted_text = chatbot.file_processor.extract_text_from_file(media_content, mime_type)
            if extracted_text:
                logging.info(f"Extracted text from document: {len(extracted_text)} characters")
        
        # Get original filename if available
        original_filename = document_data.get('filename', filename)
        
        # Save complete document message to database
        media_message_id = chatbot.history_manager.save_message_with_media_and_caption(
            phone_number=sender_id,
            message=f"[Document: {original_filename}]",
            role="user",
            message_type="document",
            media_file_path=file_path,
            media_file_name=filename,
            media_mime_type=mime_type,
            media_file_size=file_size,
            media_hash=file_hash,
            extracted_text=extracted_text,
            media_caption=caption if caption else None
        )
        
        if media_message_id:
            logger.info(f"Document message saved to database with ID: {media_message_id}")
            
            # Mark as processed
            chatbot.history_manager.mark_message_processed(media_message_id)
            
            # If human_reply is active, skip sending response
            if human_reply_active:
                logger.info(f"Document processed but no auto-reply sent due to human_reply flag")
                return True
            
            # Send manual reply response (only if human_reply is not active)
            manual_reply = chatbot.get_manual_reply_message(detected_language, 'media')
            chatbot.history_manager.save_message(sender_id, manual_reply, "assistant")
            result = chatbot.send_whatsapp_message(sender_id, manual_reply)
            return result is not None
        
        return False
        
    except Exception as e:
        logging.error(f"Error handling document message: {e}")
        import traceback
        logging.error(f"Full traceback: {traceback.format_exc()}")
        
        if not human_reply_active:
            error_msg = "Sorry, there was an error processing your document. Please try sending a text message instead."
            if detected_language == 'zh':
                error_msg = "抱歉，处理您的文档时出错。请尝试发送文本消息。"
            elif detected_language == 'ms':
                error_msg = "Maaf, terdapat ralat semasa memproses dokumen anda. Sila cuba hantar mesej teks."
            chatbot.send_whatsapp_message(sender_id, f"❌ {error_msg}")
        return False

def handle_audio_message_in_webhook(audio_data, sender_id):
    """Handle audio message in webhook - process fully but check human_reply for response"""
    logger.info(f"Received audio message from {sender_id}")
    
    # Save to followup keep table
    try:
        chatbot.history_manager.save_phone_number_to_followup_keep(sender_id)
    except Exception as e:
        logging.error(f"Error saving to followup keep in audio handler: {e}")
    
    # Check human_reply flag
    session_flags = chatbot.history_manager.get_session_flags(sender_id)
    human_reply_active = session_flags.get('human_reply', 0) == 1
    
    if human_reply_active:
        logger.info(f"Human reply flag active for {sender_id}, will process but not auto-reply")
    
    # Check if first-time user
    is_first_time = chatbot.history_manager.is_first_time_user(sender_id)
    
    # Handle first-time user opening sequence (skip if human_reply is active)
    if is_first_time and not human_reply_active:
        logger.info(f"First-time user detected: {sender_id}, starting opening sequence")
        opening_result = chatbot.handle_opening_sequence(sender_id, 'en')
        if opening_result:
            logger.info(f"Opening sequence sent to first-time audio user: {sender_id}")
    
    # Process audio message completely (transcribe, save, etc.)
    try:
        media_id = audio_data.get('id')
        if not media_id:
            logger.error("No media ID found in audio message")
            if not human_reply_active:
                error_response = "Sorry, I couldn't process your voice message. Please try again."
                chatbot.send_whatsapp_message(sender_id, error_response)
            return False
        
        logger.info(f"Processing audio message with ID: {media_id}")
        
        # Download the audio file
        audio_content, mime_type = chatbot.download_media_file(media_id)
        
        if not audio_content:
            logger.error("Failed to download audio file")
            if not human_reply_active:
                error_response = "Sorry, I couldn't download your voice message. Please try sending it again."
                chatbot.send_whatsapp_message(sender_id, error_response)
            return False
        
        logger.info(f"Downloaded audio file, mime type: {mime_type}, size: {len(audio_content)} bytes")
        
        # Save audio file to disk and GitHub
        file_path, filename, file_hash, file_size = chatbot.history_manager.save_media_file(
            audio_content, mime_type, sender_id, 'audio'
        )
        
        if not file_path:
            logger.error("Failed to save audio file")
            if not human_reply_active:
                error_response = "Sorry, I couldn't save your voice message. Please try again."
                chatbot.send_whatsapp_message(sender_id, error_response)
            return False
        
        # Transcribe the audio
        transcribed_text = chatbot.transcribe_audio(audio_content, mime_type)
        
        # Detect language from transcribed text
        detected_language = 'en'
        if transcribed_text:
            detected_language = chatbot.detect_message_language(transcribed_text)
            logger.info(f"Detected language from audio transcription: {detected_language}")
        
        # Save complete audio message to database
        audio_message_id = chatbot.history_manager.save_message_with_media(
            phone_number=sender_id,
            message=transcribed_text if transcribed_text else "[Audio message - transcription failed]",
            role="user",
            message_type="audio",
            media_file_path=file_path,
            media_file_name=filename,
            media_mime_type=mime_type,
            media_file_size=file_size,
            media_hash=file_hash,
            extracted_text=transcribed_text
        )
        
        if not audio_message_id:
            logger.error("Failed to save audio message to database")
            return False
        
        logger.info(f"Audio message saved to database with ID: {audio_message_id}")
        
        # If human_reply is active, mark as processed and skip response
        if human_reply_active:
            chatbot.history_manager.mark_message_processed(audio_message_id)
            logger.info(f"Audio message processed but no auto-reply sent due to human_reply flag")
            return True
        
        # Generate and send AI response (only if human_reply is not active)
        if transcribed_text:
            ai_response = chatbot._generate_ai_response_for_transcribed_audio(
                transcribed_text, sender_id, detected_language
            )
            
            if ai_response:
                # Mark as processed after successful AI response
                chatbot.history_manager.mark_message_processed(audio_message_id)
                
                # Add voice message indicator
                if detected_language == 'zh':
                    voice_indicator = f"🎤➡️📝 我听到了: \"{transcribed_text[:50]}{'...' if len(transcribed_text) > 50 else ''}\"\n\n{ai_response}"
                elif detected_language == 'ms':
                    voice_indicator = f"🎤➡️📝 Saya dengar: \"{transcribed_text[:50]}{'...' if len(transcribed_text) > 50 else ''}\"\n\n{ai_response}"
                else:
                    voice_indicator = f"🎤➡️📝 I heard: \"{transcribed_text[:50]}{'...' if len(transcribed_text) > 50 else ''}\"\n\n{ai_response}"
                
                result = chatbot.send_whatsapp_message(sender_id, voice_indicator)
                return result is not None
        
        # Mark as processed even if no transcription
        chatbot.history_manager.mark_message_processed(audio_message_id)
        
        # Send error message if transcription failed
        error_msg = "Sorry, I couldn't understand your voice message. The audio was saved, but please try speaking more clearly or send a text message."
        if detected_language == 'zh':
            error_msg = "抱歉，我无法理解您的语音消息。音频已保存，但请尝试更清楚地说话或发送文本消息。"
        elif detected_language == 'ms':
            error_msg = "Maaf, saya tidak dapat memahami mesej suara anda. Audio telah disimpan, tetapi sila cuba bercakap dengan lebih jelas atau hantar mesej teks."
        
        result = chatbot.send_whatsapp_message(sender_id, error_msg)
        return result is not None
        
    except Exception as e:
        logger.error(f"Error handling audio message: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        
        if not human_reply_active:
            error_response = "Sorry, there was an error processing your voice message. Please try sending a text message instead."
            chatbot.send_whatsapp_message(sender_id, error_response)
        return False
      
def is_duplicate_recent_message(sender_id, message_content, message_timestamp):
    """Check if we recently processed this exact message from this user"""
    global recent_processed_messages
    
    # Clean old entries (older than 10 minutes)
    current_time = time.time()
    cutoff_time = current_time - 600  # 10 minutes
    
    # Clean old entries
    for key in list(recent_processed_messages.keys()):
        if recent_processed_messages[key] < cutoff_time:
            del recent_processed_messages[key]
    
    # Create unique key for this message
    message_key = f"{sender_id}:{message_content[:100]}:{message_timestamp}"
    
    # Check if we've seen this message recently
    if message_key in recent_processed_messages:
        return True
    
    # Mark this message as processed
    recent_processed_messages[message_key] = current_time
    return False
    
def handle_commands(command, sender_id):
    """Handle special commands with language detection"""
    # Detect language preference for commands (default to English)
    detected_language = 'en'  # Default for commands
    
    if command == '/help':
        if detected_language == 'zh':
            return """📊 *项目助手命令:*

/help - 显示此帮助信息
/clear - 清除我们的聊天记录
/info - 获取关于此机器人的信息
/project - 询问项目相关问题

💬 *如何使用:*
• 发送文本消息获得即时AI回复
• 发送语音消息 - 我会转录并回复！🎤
• 发送图片、视频、文档 - 我会帮助您联系人工客服
• 使用以 / 开头的命令

只需给我发送任何消息，我就会用AI助手回复您！😊"""
        elif detected_language == 'ms':
            return """📊 *Arahan Pembantu Projek:*

/help - Tunjukkan mesej bantuan ini
/clear - Kosongkan sejarah sembang kita
/info - Dapatkan maklumat tentang bot ini
/project - Tanya soalan berkaitan projek

💬 *Cara menggunakan:*
• Hantar mesej teks untuk respons AI segera
• Hantar mesej suara - Saya akan transkripsi dan balas! 🎤
• Hantar gambar, video, dokumen - Saya akan bantu hubungkan anda dengan sokongan manusia
• Gunakan arahan yang bermula dengan /

Hanya hantar saya sebarang mesej dan saya akan balas dengan bantuan AI! 😊"""
        else:
            return """📊 *Project Assistant Commands:*

/help - Show this help message
/clear - Clear our chat history
/info - Get information about this bot
/project - Ask project-related questions

💬 *How to use:*
• Send text messages for instant AI responses
• Send voice messages - I'll transcribe and respond! 🎤
• Send images, videos, documents - I'll help connect you to human support
• Use commands starting with /

Just send me any message and I'll respond with AI assistance! 😊"""
    
    elif command == '/clear':
        if sender_id in chat_histories:
            del chat_histories[sender_id]
        if detected_language == 'zh':
            return "🗑️ 聊天记录已清除！我们重新开始。"
        elif detected_language == 'ms':
            return "🗑️ Sejarah sembang telah dikosongkan! Kita bermula semula."
        else:
            return "🗑️ Chat history cleared! We're starting fresh."
    
    elif command == '/info':
        if detected_language == 'zh':
            return """ℹ️ *关于项目助手:*
        
我是专门从事项目管理的AI助手。我可以帮助您：
• 项目规划
• 任务调度
• 资源分配
• 进度跟踪

对于复杂问题，请咨询认证项目经理。"""
        elif detected_language == 'ms':
            return """ℹ️ *Tentang Pembantu Projek:*
        
Saya adalah pembantu AI yang pakar dalam pengurusan projek. Saya boleh membantu anda dengan:
• Perancangan projek
• Penjadualan tugas
• Peruntukan sumber
• Penjejakan kemajuan

Untuk isu kompleks, sila rujuk pengurus projek yang diperakui."""
        else:
            return """ℹ️ *About Project Assistant:*
        
I'm an AI assistant specialized in project management. I can help you with:
• Project planning
• Task scheduling
• Resource allocation
• Progress tracking

For complex issues, please consult a certified project manager."""
    
    elif command == '/project':
        if detected_language == 'zh':
            return "📊 请描述您的项目相关问题："
        elif detected_language == 'ms':
            return "📊 Sila huraikan soalan berkaitan projek anda:"
        else:
            return "📊 Please describe your project-related question:"
    
    else:
        if detected_language == 'zh':
            return "❓ 未知命令。发送 /help 查看可用命令。"
        elif detected_language == 'ms':
            return "❓ Arahan tidak diketahui. Hantar /help untuk melihat arahan yang tersedia."
        else:
            return "❓ Unknown command. Send /help to see available commands."

@app.route('/webhook', methods=['GET', 'POST'])
def webhook():
    if request.method == 'GET':
        # Verification logic remains the same
        mode = request.args.get('hub.mode')
        verify_token = request.args.get('hub.verify_token')
        challenge = request.args.get('hub.challenge')
        
        logger.info(f"Received verification request: mode={mode}, token={verify_token}, challenge={challenge}")
        
        if None in [mode, verify_token, challenge]:
            missing = [name for name, value in [('hub.mode', mode), 
                                               ('hub.verify_token', verify_token),
                                               ('hub.challenge', challenge)] if value is None]
            logger.warning(f"Missing parameters: {', '.join(missing)}")
            return "Missing required verification parameters", 400
            
        if mode == 'subscribe' and verify_token == META_VERIFY_TOKEN:
            logger.info("Webhook verified successfully!")
            return challenge, 200
        else:
            logger.warning(f"Verification failed | Expected token: {META_VERIFY_TOKEN} | Received: {verify_token}")
            return "Verification failed", 403
    
    elif request.method == 'POST':
        try:
            data = request.get_json()
            logger.info(f"Received webhook data: {json.dumps(data, indent=2)}")
            
            if data and 'entry' in data:
                for entry in data['entry']:
                    if 'changes' in entry:
                        for change in entry['changes']:
                            if change.get('field') == 'messages':
                                value = change.get('value', {})
                                messages = value.get('messages', [])
                                
                                for message in messages:
                                    sender_id = message.get('from')
                                    message_type = message.get('type')
                                    message_timestamp = message.get('timestamp')
                                    
                                    if not sender_id or not message_type:
                                        continue
                                    
                                    # Check if message is too old (same logic as before)
                                    if message_timestamp:
                                        try:
                                            msg_time = int(message_timestamp)
                                            current_time = time.time()
                                            msg_age_seconds = current_time - msg_time
                                            
                                            if msg_time < (SERVER_START_TIME - 180):
                                                logger.warning(f"IGNORING OLD MESSAGE from {sender_id}: Message sent {msg_age_seconds} seconds ago, before server started")
                                                continue
                                            
                                            if msg_age_seconds > 300:
                                                logger.warning(f"IGNORING VERY OLD MESSAGE from {sender_id}: Message is {msg_age_seconds} seconds old")
                                                continue
                                                
                                        except (ValueError, TypeError) as e:
                                            logger.error(f"Error parsing message timestamp {message_timestamp}: {e}")
                                            if time.time() - SERVER_START_TIME < 120:
                                                logger.warning(f"IGNORING MESSAGE due to timestamp parse error and recent server start")
                                                continue
                                    
                                    logger.info(f"Processing message from {sender_id}, type: {message_type}")
                                    
                                    # Duplicate check (same as before)
                                    message_content = ""
                                    if message_type == 'text':
                                        text_data = message.get('text', {})
                                        text_content = text_data.get('body', '')
                                        message_content = f"[Text message - {text_content[:50]}]"
                                    elif message_type == 'audio':
                                        message_content = f"[Audio message - {message.get('audio', {}).get('id', 'unknown')}]"
                                    elif message_type == 'image':
                                        message_content = f"[Image message - {message.get('image', {}).get('id', 'unknown')}]"
                                    elif message_type == 'video':
                                        message_content = f"[Video message - {message.get('video', {}).get('id', 'unknown')}]"
                                    elif message_type == 'document':
                                        message_content = f"[Document message - {message.get('document', {}).get('id', 'unknown')}]"
                                    
                                    if is_duplicate_recent_message(sender_id, message_content, message_timestamp):
                                        logger.warning(f"IGNORING DUPLICATE MESSAGE from {sender_id}: {message_content[:50]}...")
                                        continue
                                    
                                    # Handle different message types - FIXED TEXT MESSAGE HANDLER
                                    if message_type == 'text':
                                        text_data = message.get('text', {})
                                        text_content = text_data.get('body', '')
                                        logger.info(f"Received text message: {text_content}")
                                        
                                        # Check for commands first
                                        if text_content.startswith('/'):
                                            response_text = handle_commands(text_content, sender_id)
                                            
                                            # Save command and response
                                            command_message_id = chatbot.history_manager.save_message(
                                                sender_id, text_content, "user", "text"
                                            )
                                            if response_text:
                                                chatbot.history_manager.save_message(
                                                    sender_id, response_text, "assistant", "text"
                                                )
                                            if command_message_id:
                                                chatbot.history_manager.mark_message_processed(command_message_id)
                                        else:
                                            # FIXED: Use the corrected message processing flow
                                            response_text = chatbot.process_user_message(text_content, sender_id)
                                        
                                        # Send response back to WhatsApp
                                        if response_text:
                                            result = chatbot.send_whatsapp_message(sender_id, response_text)
                                            if result:
                                                logger.info(f"Response sent successfully to {sender_id}")
                                            else:
                                                logger.error(f"Failed to send response to {sender_id}")
                                        else:
                                            logger.info(f"No response generated for {sender_id}")
                                    
                                    # Handle audio messages (still get AI responses)
                                    elif message_type == 'audio':
                                        handle_audio_message_in_webhook(message.get('audio', {}), sender_id)
                                    
                                    # Handle media messages (trigger manual reply)
                                    elif message_type == 'image':
                                        handle_image_message_in_webhook(message.get('image', {}), sender_id)
                                    elif message_type == 'video':
                                        handle_video_message_in_webhook(message.get('video', {}), sender_id)
                                    elif message_type == 'document':
                                        handle_document_message_in_webhook(message.get('document', {}), sender_id)
            
            return jsonify({"status": "success"}), 200
            
        except Exception as e:
            logger.error(f"Error processing webhook POST request: {e}")
            return jsonify({"status": "error", "message": str(e)}), 200
    
    return jsonify({"status": "method not allowed"}), 405

@app.route('/admin/process-unprocessed', methods=['POST'])
def process_unprocessed_messages():
    """Process any unprocessed messages (manual recovery)"""
    try:
        unprocessed = chatbot.history_manager.get_unprocessed_messages()
        
        if not unprocessed:
            return jsonify({"status": "success", "message": "No unprocessed messages found"})
        
        processed_count = 0
        for msg in unprocessed:
            try:
                # Process the message
                response = chatbot.get_ai_response(msg['message'], msg['phone_number'])
                if response:
                    # Send response
                    chatbot.send_whatsapp_message(msg['phone_number'], response)
                    processed_count += 1
                    
            except Exception as e:
                logging.error(f"Error processing unprocessed message {msg['message_id']}: {e}")
                # Mark as processed even if failed to prevent infinite loops
                chatbot.history_manager.mark_message_processed(msg['message_id'])
        
        return jsonify({
            "status": "success", 
            "message": f"Processed {processed_count} out of {len(unprocessed)} unprocessed messages"
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
        
@app.route('/token-info', methods=['GET'])
def token_info():
    """Get information about the current token"""
    is_valid, info = chatbot.validate_token()
    debug_success, debug_info = chatbot.get_token_info()
    
    return jsonify({
        "token_valid": is_valid,
        "validation_info": info,
        "debug_success": debug_success,
        "debug_info": debug_info,
        "current_token_preview": chatbot.current_token[:20] + "..." if chatbot.current_token else "No token"
    })

@app.route('/stats', methods=['GET'])
def get_stats():
    """Get bot statistics"""
    return jsonify({
        "active_chats": len(chat_histories),
        "total_messages": sum(len(history) for history in chat_histories.values()),
        "users_last_24h": len([
            phone for phone, last_activity in user_last_activity.items()
            if datetime.now() - last_activity < timedelta(days=1)
        ])
    })

@app.route('/test', methods=['GET'])
def test_endpoint():
    """Enhanced test endpoint to verify server is running"""
    return jsonify({
        "status": "Server is running!",
        "timestamp": datetime.now().isoformat(),
        "webhook_url": request.host_url + "webhook",
        "mysql_connected": chatbot.history_manager.connection is not None and chatbot.history_manager.connection.is_connected(),
        "environment_vars": {
            "PORT": os.getenv('PORT'),
            "MYSQL_HOST": os.getenv('MYSQL_HOST'),
            "MYSQL_DATABASE": os.getenv('MYSQL_DATABASE'),
            "MYSQL_URL_EXISTS": bool(os.getenv('MYSQL_URL')),
            "OPENAI_API_KEY_EXISTS": bool(os.getenv('OPENAI_API_KEY')),
            "META_WHATSAPP_TOKEN_EXISTS": bool(os.getenv('META_WHATSAPP_TOKEN'))
        }
    })

# Continue with the rest of the endpoint definitions...
@app.route('/admin/users', methods=['GET'])
def get_all_users():
    """Get list of all users who have chatted"""
    try:
        users = chatbot.history_manager.get_all_users()
        return jsonify({
            "status": "success",
            "users": users,
            "total_users": len(users)
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/chat/<phone_number>', methods=['GET'])
def get_user_chat(phone_number):
    """Get chat history for a specific user"""
    try:
        limit = request.args.get('limit', 100, type=int)
        chat_history = chatbot.history_manager.get_user_chat_history(phone_number, limit)
        
        return jsonify({
            "status": "success",
            "phone_number": phone_number,
            "messages": chat_history,
            "total_messages": len(chat_history)
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/export', methods=['GET'])
def export_chat_history():
    """Export chat history"""
    try:
        phone_number = request.args.get('phone_number')
        format_type = request.args.get('format', 'json')
        
        filename = chatbot.history_manager.export_chat_history(phone_number, format_type)
        
        if filename:
            return jsonify({
                "status": "success",
                "message": f"Chat history exported to {filename}",
                "filename": filename
            })
        else:
            return jsonify({"status": "error", "message": "Failed to export chat history"}), 500
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/stats', methods=['GET'])
def get_admin_stats():
    """Get detailed chat statistics"""
    try:
        stats = chatbot.history_manager.get_chat_statistics()
        return jsonify({
            "status": "success",
            "statistics": stats
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/new-message-count/<phone_number>', methods=['GET'])
def get_new_message_count_endpoint(phone_number):
    """Get the new_message_count for a specific user"""
    try:
        count = chatbot.history_manager.get_new_message_count(phone_number)
        return jsonify({
            "status": "success",
            "phone_number": phone_number,
            "new_message_count": count
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/reset-new-message-count/<phone_number>', methods=['POST'])
def reset_new_message_count_endpoint(phone_number):
    """Reset the new_message_count to 0 for a specific user"""
    try:
        success = chatbot.history_manager.reset_new_message_count(phone_number)
        if success:
            return jsonify({
                "status": "success",
                "message": f"New message count reset for {phone_number}"
            })
        else:
            return jsonify({
                "status": "error", 
                "message": "Failed to reset new message count"
            }), 500
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/users-with-new-messages', methods=['GET'])
def get_users_with_new_messages_endpoint():
    """Get list of users who have new unread messages"""
    try:
        min_count = request.args.get('min_count', 1, type=int)
        users = chatbot.history_manager.get_users_with_new_messages(min_count)
        
        return jsonify({
            "status": "success",
            "users_with_new_messages": users,
            "total_users": len(users),
            "min_count_threshold": min_count
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/new-message-stats', methods=['GET'])
def get_new_message_stats():
    """Get statistics about new messages across all users"""
    try:
        if not chatbot.history_manager.connection:
            return jsonify({
                "status": "error", 
                "message": "Database connection not available"
            }), 500
        
        cursor = chatbot.history_manager.connection.cursor(dictionary=True)
        
        # Get overall statistics
        cursor.execute('''
            SELECT 
                COUNT(*) as total_active_sessions,
                SUM(new_message_count) as total_new_messages,
                AVG(new_message_count) as avg_new_messages_per_user,
                MAX(new_message_count) as max_new_messages,
                COUNT(CASE WHEN new_message_count > 0 THEN 1 END) as users_with_new_messages
            FROM chat_sessions 
            WHERE session_end IS NULL
        ''')
        
        stats = cursor.fetchone()
        
        # Get distribution
        cursor.execute('''
            SELECT 
                new_message_count,
                COUNT(*) as user_count
            FROM chat_sessions 
            WHERE session_end IS NULL
            GROUP BY new_message_count
            ORDER BY new_message_count
        ''')
        
        distribution = cursor.fetchall()
        cursor.close()
        
        return jsonify({
            "status": "success",
            "new_message_statistics": stats,
            "distribution": distribution
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
        
@app.route('/api/send_media/<path:file_path>')
def serve_send_media_file(file_path):
    """Serve sent media files for verification/download"""
    try:
        full_path = Path("send_media") / file_path
        
        if not full_path.exists():
            return jsonify({"error": "Sent media file not found"}), 404
        
        # Get mime type
        mime_type = mimetypes.guess_type(str(full_path))[0] or 'application/octet-stream'
        
        return send_file(
            str(full_path),
            mimetype=mime_type,
            as_attachment=False,
            download_name=full_path.name
        )
        
    except Exception as e:
        logging.error(f"Error serving sent media file: {e}")
        return jsonify({"error": "Error serving sent media file"}), 500

@app.route('/admin/sent-media-messages', methods=['GET'])
def get_sent_media_messages_endpoint():
    """API endpoint to get all sent media messages (special replies) using existing database columns"""
    try:
        phone_number = request.args.get('phone_number')
        media_type = request.args.get('media_type')  # 'image', 'video'
        limit = request.args.get('limit', 50, type=int)
        
        if not chatbot.history_manager.connection:
            return jsonify({"status": "error", "message": "Database connection not available"}), 500
        
        cursor = chatbot.history_manager.connection.cursor(dictionary=True)
        
        # Build query based on parameters using existing columns
        where_conditions = ["role = 'assistant'"]
        params = []
        
        if phone_number:
            where_conditions.append("phone_number = %s")
            params.append(phone_number)
        
        if media_type == 'image':
            where_conditions.append("image_file_path IS NOT NULL")
        elif media_type == 'video':
            where_conditions.append("video_file_path IS NOT NULL")
        elif media_type:
            return jsonify({"status": "error", "message": "Invalid media_type. Use 'image' or 'video'"}), 400
        else:
            where_conditions.append("(image_file_path IS NOT NULL OR video_file_path IS NOT NULL)")
        
        where_clause = " AND ".join(where_conditions)
        params.append(limit)
        
        query = f'''
            SELECT message_id, phone_number, message, message_type, timestamp,
                   image_file_path, image_file_name, image_mime_type, image_file_size, image_hash,
                   video_file_path, video_file_name, video_mime_type, video_file_size, video_hash
            FROM chat_messages 
            WHERE {where_clause}
            ORDER BY timestamp DESC 
            LIMIT %s
        '''
        
        cursor.execute(query, params)
        messages = cursor.fetchall()
        cursor.close()
        
        # Add media URLs to each message
        for message in messages:
            if message['image_file_path']:
                message['media_url'] = f"/api/send_media/{message['image_file_path']}"
                message['media_type'] = 'image'
                message['media_filename'] = message['image_file_name']
                message['media_size'] = message['image_file_size']
                message['media_hash'] = message['image_hash']
            elif message['video_file_path']:
                message['media_url'] = f"/api/send_media/{message['video_file_path']}"
                message['media_type'] = 'video'
                message['media_filename'] = message['video_file_name']
                message['media_size'] = message['video_file_size']
                message['media_hash'] = message['video_hash']
        
        return jsonify({
            "status": "success",
            "sent_media_messages": messages,
            "total": len(messages),
            "media_type": media_type or "all"
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/github-upload-stats', methods=['GET'])
def get_github_upload_stats():
    """Get GitHub upload statistics for sent media using existing database columns"""
    try:
        if not chatbot.history_manager.connection:
            return jsonify({"status": "error", "message": "Database connection not available"}), 500
        
        cursor = chatbot.history_manager.connection.cursor()
        
        # Total sent media messages (images and videos with file paths)
        cursor.execute('''
            SELECT COUNT(*) FROM chat_messages 
            WHERE role = 'assistant' 
            AND (image_file_path IS NOT NULL OR video_file_path IS NOT NULL)
        ''')
        total_sent_media = cursor.fetchone()[0]
        
        # Count by media type
        cursor.execute('''
            SELECT 'image' as media_type, COUNT(*) as count
            FROM chat_messages 
            WHERE role = 'assistant' AND image_file_path IS NOT NULL
            UNION ALL
            SELECT 'video' as media_type, COUNT(*) as count
            FROM chat_messages 
            WHERE role = 'assistant' AND video_file_path IS NOT NULL
        ''')
        media_type_stats = dict(cursor.fetchall())
        
        cursor.close()
        
        return jsonify({
            "status": "success",
            "github_upload_statistics": {
                'total_sent_media': total_sent_media,
                'media_type_statistics': media_type_stats,
                'note': 'Using existing database columns - all files are automatically uploaded to GitHub'
            }
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
        
@app.route('/admin/search', methods=['GET'])
def search_chat_messages():
    """Search for messages"""
    try:
        search_term = request.args.get('q', '')
        phone_number = request.args.get('phone_number')
        
        if not search_term:
            return jsonify({"status": "error", "message": "Search term is required"}), 400
        
        results = chatbot.history_manager.search_messages(search_term, phone_number)
        
        return jsonify({
            "status": "success",
            "search_term": search_term,
            "results": results,
            "total_results": len(results)
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/media/<path:filename>')
def serve_media(filename):
    full_path = os.path.join('media', filename)
    if not os.path.isfile(full_path):
        return f"File not found: {full_path}", 404
    return send_from_directory('media', filename)
    
# Media file serving endpoints
@app.route('/api/audio/<path:file_path>')
def serve_audio_file(file_path):
    """Serve audio files for frontend playback"""
    try:
        full_path = chatbot.history_manager.audio_storage_path / file_path
        
        if not full_path.exists():
            return jsonify({"error": "Audio file not found"}), 404
        
        # Get mime type
        mime_type = mimetypes.guess_type(str(full_path))[0] or 'audio/ogg'
        
        return send_file(
            str(full_path),
            mimetype=mime_type,
            as_attachment=False,
            download_name=full_path.name
        )
        
    except Exception as e:
        logging.error(f"Error serving audio file: {e}")
        return jsonify({"error": "Error serving audio file"}), 500

@app.route('/api/image/<path:file_path>')
def serve_image_file(file_path):
    """Serve image files"""
    try:
        full_path = chatbot.history_manager.image_storage_path / file_path
        
        if not full_path.exists():
            return jsonify({"error": "Image file not found"}), 404
        
        mime_type = mimetypes.guess_type(str(full_path))[0] or 'image/jpeg'
        
        return send_file(
            str(full_path),
            mimetype=mime_type,
            as_attachment=False,
            download_name=full_path.name
        )
        
    except Exception as e:
        logging.error(f"Error serving image file: {e}")
        return jsonify({"error": "Error serving image file"}), 500

@app.route('/api/video/<path:file_path>')
def serve_video_file(file_path):
    """Serve video files"""
    try:
        full_path = chatbot.history_manager.video_storage_path / file_path
        
        if not full_path.exists():
            return jsonify({"error": "Video file not found"}), 404
        
        mime_type = mimetypes.guess_type(str(full_path))[0] or 'video/mp4'
        
        return send_file(
            str(full_path),
            mimetype=mime_type,
            as_attachment=False,
            download_name=full_path.name
        )
        
    except Exception as e:
        logging.error(f"Error serving video file: {e}")
        return jsonify({"error": "Error serving video file"}), 500

@app.route('/api/document/<path:file_path>')
def serve_document_file(file_path):
    """Serve document files"""
    try:
        full_path = chatbot.history_manager.document_storage_path / file_path
        
        if not full_path.exists():
            return jsonify({"error": "Document file not found"}), 404
        
        mime_type = mimetypes.guess_type(str(full_path))[0] or 'application/octet-stream'
        
        return send_file(
            str(full_path),
            mimetype=mime_type,
            as_attachment=True,  # Documents should be downloaded
            download_name=full_path.name
        )
        
    except Exception as e:
        logging.error(f"Error serving document file: {e}")
        return jsonify({"error": "Error serving document file"}), 500

@app.route('/api/media/<media_type>/<filename>')
def serve_media_file(media_type, filename):
    """Serve media files (images/videos) for WhatsApp"""
    try:
        file_path = chatbot.media_handler.get_media_file_path(filename, media_type)
        
        if not file_path.exists():
            return jsonify({"error": "Media file not found"}), 404
        
        # Get mime type
        mime_type = mimetypes.guess_type(str(file_path))[0]
        if not mime_type:
            if media_type == 'image':
                mime_type = 'image/jpeg'  # Default for images
            elif media_type == 'video':
                mime_type = 'video/mp4'   # Default for videos
            else:
                mime_type = 'application/octet-stream'
        
        return send_file(
            str(file_path),
            mimetype=mime_type,
            as_attachment=False,
            download_name=filename
        )
        
    except Exception as e:
        logging.error(f"Error serving media file: {e}")
        return jsonify({"error": "Error serving media file"}), 500

@app.route('/api/private/send_media/<path:file_path>')
def serve_private_send_media_file(file_path):
    """UPDATED: Serve files with new organized structure"""
    try:
        # Handle both old and new path structures
        if file_path.startswith('external/'):
            # Old structure - try to serve if exists
            full_path = Path("public/send_media") / file_path
        else:
            # New organized structure: media_type/year/month/day/filename
            full_path = Path("public/send_media") / file_path
        
        if not full_path.exists():
            logging.error(f"Organized media file not found: {full_path}")
            return jsonify({
                "error": "Media file not found", 
                "requested_path": file_path,
                "structure": "organized_year_month_day"
            }), 404
        
        # Get mime type
        mime_type = mimetypes.guess_type(str(full_path))[0] or 'application/octet-stream'
        
        logging.info(f"Serving organized media file: {file_path}")
        
        return send_file(
            str(full_path),
            mimetype=mime_type,
            as_attachment=False,
            download_name=full_path.name
        )
        
    except Exception as e:
        logging.error(f"Error serving organized media file: {e}")
        return jsonify({"error": "Error serving media file"}), 500

@app.route('/admin/followup-keep', methods=['GET'])
def get_followup_keep_data():
    """Get all phone numbers in chat_followup_keep table"""
    try:
        if not chatbot.history_manager.connection:
            return jsonify({
                "status": "error", 
                "message": "Database connection not available"
            }), 500
        
        cursor = chatbot.history_manager.connection.cursor(dictionary=True)
        
        # Get pagination parameters
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)
        offset = (page - 1) * per_page
        
        # Get total count
        cursor.execute('SELECT COUNT(*) as total FROM chat_followup_keep')
        total_count = cursor.fetchone()['total']
        
        # Get paginated data
        cursor.execute('''
            SELECT id, phone_number, created_datetime
            FROM chat_followup_keep
            ORDER BY created_datetime DESC
            LIMIT %s OFFSET %s
        ''', (per_page, offset))
        
        followup_data = cursor.fetchall()
        cursor.close()
        
        return jsonify({
            "status": "success",
            "followup_keep_data": followup_data,
            "pagination": {
                "total": total_count,
                "page": page,
                "per_page": per_page,
                "total_pages": (total_count + per_page - 1) // per_page
            }
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/followup-keep/stats', methods=['GET'])
def get_followup_keep_stats():
    """Get statistics for chat_followup_keep table"""
    try:
        if not chatbot.history_manager.connection:
            return jsonify({
                "status": "error", 
                "message": "Database connection not available"
            }), 500
        
        cursor = chatbot.history_manager.connection.cursor(dictionary=True)
        
        # Total unique phone numbers
        cursor.execute('SELECT COUNT(*) as total FROM chat_followup_keep')
        total_unique = cursor.fetchone()['total']
        
        # New entries today
        cursor.execute('''
            SELECT COUNT(*) as today_count 
            FROM chat_followup_keep 
            WHERE DATE(created_datetime) = DATE(CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00'))
        ''')
        today_count = cursor.fetchone()['today_count']
        
        # New entries this week
        cursor.execute('''
            SELECT COUNT(*) as today_count 
            FROM chat_followup_keep 
            WHERE YEARWEEK(created_datetime, 1) = YEARWEEK(CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00'), 1)
        ''')
        week_count = cursor.fetchone()['week_count']
        
        # New entries this month
        cursor.execute('''
            SELECT COUNT(*) as month_count 
            FROM chat_followup_keep 
            WHERE YEAR(created_datetime) = YEAR(CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00'))
            AND MONTH(created_datetime) = MONTH(CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00'))
        ''')
        month_count = cursor.fetchone()['month_count']
        
        cursor.close()
        
        return jsonify({
            "status": "success",
            "followup_keep_statistics": {
                "total_unique_phone_numbers": total_unique,
                "new_today": today_count,
                "new_this_week": week_count,
                "new_this_month": month_count
            }
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
        
@app.route('/admin/private-repo-media-messages', methods=['GET'])
def get_private_repo_media_messages_endpoint():
    """UPDATED: Handle organized file structure in GitHub URLs"""
    try:
        phone_number = request.args.get('phone_number')
        media_type = request.args.get('media_type')
        limit = request.args.get('limit', 50, type=int)
        
        if not chatbot.history_manager.connection:
            return jsonify({"status": "error", "message": "Database connection not available"}), 500
        
        cursor = chatbot.history_manager.connection.cursor(dictionary=True)
        
        # Build query for private repo uploads
        where_conditions = ["role = 'assistant'"]
        params = []
        
        if phone_number:
            where_conditions.append("phone_number = %s")
            params.append(phone_number)
        
        if media_type == 'image':
            where_conditions.append("image_file_path IS NOT NULL")
            where_conditions.append("(message LIKE '%Private Repo%' OR message LIKE '%Special Image%')")
        elif media_type == 'video':
            where_conditions.append("video_file_path IS NOT NULL") 
            where_conditions.append("(message LIKE '%Private Repo%' OR message LIKE '%Special Video%')")
        elif media_type:
            return jsonify({"status": "error", "message": "Invalid media_type"}), 400
        else:
            where_conditions.append("(image_file_path IS NOT NULL OR video_file_path IS NOT NULL)")
            where_conditions.append("(message LIKE '%Private Repo%' OR message LIKE '%Special Image%' OR message LIKE '%Special Video%')")
        
        where_clause = " AND ".join(where_conditions)
        params.append(limit)
        
        query = f'''
            SELECT message_id, phone_number, message, message_type, timestamp,
                   image_file_path, image_file_name, image_mime_type, image_file_size, image_hash,
                   video_file_path, video_file_name, video_mime_type, video_file_size, video_hash
            FROM chat_messages 
            WHERE {where_clause}
            ORDER BY timestamp DESC 
            LIMIT %s
        '''
        
        cursor.execute(query, params)
        messages = cursor.fetchall()
        cursor.close()
        
        # Add URLs with organized structure support
        for message in messages:
            if message['image_file_path']:
                # Local API URL
                message['media_url'] = f"/api/private/send_media/{message['image_file_path']}"
                
                # UPDATED: GitHub URL with organized structure
                if message['image_file_path'].startswith('image/'):
                    # New organized structure: image/2025/08/26/filename.jpg
                    message['github_url'] = f"https://github.com/Welsh510/reactjs-appointmentwhatsapp/blob/master/public/send_media/{message['image_file_path']}"
                    message['path_structure'] = 'organized'
                else:
                    # Old structure
                    message['github_url'] = f"https://github.com/Welsh510/reactjs-appointmentwhatsapp/blob/master/public/send_media/{message['image_file_path']}"
                    message['path_structure'] = 'legacy'
                
                message['media_type'] = 'image'
                message['filename'] = message['image_file_name']
                
            elif message['video_file_path']:
                # Local API URL
                message['media_url'] = f"/api/private/send_media/{message['video_file_path']}"
                
                # UPDATED: GitHub URL with organized structure
                if message['video_file_path'].startswith('video/'):
                    # New organized structure: video/2025/08/26/filename.mp4
                    message['github_url'] = f"https://github.com/Welsh510/reactjs-appointmentwhatsapp/blob/master/public/send_media/{message['video_file_path']}"
                    message['path_structure'] = 'organized'
                else:
                    # Old structure
                    message['github_url'] = f"https://github.com/Welsh510/reactjs-appointmentwhatsapp/blob/master/public/send_media/{message['video_file_path']}"
                    message['path_structure'] = 'legacy'
                
                message['media_type'] = 'video'
                message['filename'] = message['video_file_name']
        
        return jsonify({
            "status": "success",
            "private_repo_media_messages": messages,
            "total": len(messages),
            "repository": "reactjs-appointmentwhatsapp (private)",
            "branch": "master",
            "file_structure": "organized_by_date (media_type/YYYY/MM/DD/filename.ext)",
            "note": "New uploads use organized year/month/day folder structure"
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/private-repo-upload-stats', methods=['GET'])
def get_private_repo_upload_stats():
    """FIXED: Update statistics to show master branch info"""
    try:
        if not chatbot.history_manager.connection:
            return jsonify({"status": "error", "message": "Database connection not available"}), 500
        
        cursor = chatbot.history_manager.connection.cursor()
        
        # Count all special reply media (both successful uploads and URL-served)
        cursor.execute('''
            SELECT COUNT(*) FROM chat_messages 
            WHERE role = 'assistant' 
            AND (message LIKE '%Private Repo%' OR message LIKE '%Special Image%' OR message LIKE '%Special Video%')
            AND (image_file_path IS NOT NULL OR video_file_path IS NOT NULL)
        ''')
        total_special_replies = cursor.fetchone()[0]
        
        # Count successful uploads vs URL-served
        cursor.execute('''
            SELECT 
                SUM(CASE WHEN message LIKE '%Private Repo%' THEN 1 ELSE 0 END) as uploaded_to_repo,
                SUM(CASE WHEN message LIKE '%Special Image%' OR message LIKE '%Special Video%' THEN 1 ELSE 0 END) as served_from_url
            FROM chat_messages 
            WHERE role = 'assistant' 
            AND (image_file_path IS NOT NULL OR video_file_path IS NOT NULL)
        ''')
        upload_stats = cursor.fetchone()
        
        cursor.close()
        
        return jsonify({
            "status": "success",
            "private_repo_statistics": {
                'total_special_reply_media': total_special_replies,
                'uploaded_to_github': upload_stats[0] or 0,
                'served_from_url': upload_stats[1] or 0,
                'repository': 'reactjs-appointmentwhatsapp',
                'branch': 'master',  # FIXED: Show master branch
                'upload_path': 'public/send_media/',
                'github_url': 'https://github.com/Welsh510/reactjs-appointmentwhatsapp/tree/master',
                'note': 'Files uploaded to master branch. Some may be served from external URLs if local copy not found.'
            }
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/sync-private-repo', methods=['POST'])
def sync_private_repo():
    """FIXED: Manual sync with master branch support"""
    try:
        if not chatbot.history_manager.special_reply_github_manager.github_token:
            return jsonify({
                "status": "error", 
                "message": "GitHub token not configured"
            }), 400
        
        # FIXED: Use master branch instead of main
        default_branch = "master"
        
        try:
            # Try to push to master branch
            result = subprocess.run([
                "git", "push", "private", f"HEAD:{default_branch}"
            ], cwd=".", capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0:
                return jsonify({
                    "status": "success",
                    "message": f"Successfully synced with private repository ({default_branch} branch)",
                    "branch": default_branch,
                    "output": result.stdout
                })
            else:
                # Try alternative push method for master branch
                alt_result = subprocess.run([
                    "git", "push", "private", f"HEAD:refs/heads/{default_branch}"
                ], cwd=".", capture_output=True, text=True, timeout=60)
                
                if alt_result.returncode == 0:
                    return jsonify({
                        "status": "success",
                        "message": f"Successfully synced with alternative method ({default_branch} branch)",
                        "branch": default_branch,
                        "output": alt_result.stdout
                    })
                else:
                    return jsonify({
                        "status": "error",
                        "message": f"Both push methods failed to {default_branch} branch",
                        "error_output": result.stderr,
                        "alt_error_output": alt_result.stderr
                    }), 500
                
        except subprocess.TimeoutExpired:
            return jsonify({
                "status": "error",
                "message": "Sync timeout - operation took too long"
            }), 408
            
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/variation-cache-stats', methods=['GET'])
def get_variation_cache_stats():
    """Get statistics about the text variation cache"""
    try:
        cache = chatbot.text_variation_generator.variation_cache
        return jsonify({
            "status": "success",
            "cache_statistics": {
                "total_cached_variations": len(cache),
                "cache_max_size": chatbot.text_variation_generator.cache_max_size,
                "cache_usage_percentage": round(len(cache) / chatbot.text_variation_generator.cache_max_size * 100, 2),
                "sample_cache_keys": list(cache.keys())[:5] if cache else []
            }
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
        
@app.route('/admin/special-reply-usage/<phone_number>', methods=['GET'])
def get_user_special_reply_usage(phone_number):
    """Get special reply usage statistics for a specific user"""
    try:
        stats = chatbot.history_manager.get_user_special_reply_stats(phone_number)
        return jsonify({
            "status": "success",
            "phone_number": phone_number,
            "special_reply_usage": stats,
            "total_triggered": len(stats)
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/reset-cooldown', methods=['POST'])
def reset_user_cooldown():
    """Reset cooldown for a specific user and special reply"""
    try:
        data = request.get_json()
        phone_number = data.get('phone_number')
        special_reply_id = data.get('special_reply_id')
        
        if not phone_number or not special_reply_id:
            return jsonify({
                "status": "error", 
                "message": "phone_number and special_reply_id are required"
            }), 400
        
        if not chatbot.history_manager.connection:
            return jsonify({
                "status": "error", 
                "message": "Database connection not available"
            }), 500
        
        cursor = chatbot.history_manager.connection.cursor()
        
        # MODIFIED: Delete using user_id instead of phone_number
        cursor.execute('''
            DELETE FROM special_reply_usage 
            WHERE user_id = %s AND special_reply_id = %s
        ''', (phone_number, special_reply_id))
        
        rows_affected = cursor.rowcount
        chatbot.history_manager.connection.commit()
        cursor.close()
        
        if rows_affected > 0:
            return jsonify({
                "status": "success",
                "message": f"Cooldown reset for user {phone_number}, special reply {special_reply_id}"
            })
        else:
            return jsonify({
                "status": "warning",
                "message": f"No cooldown record found for user {phone_number}, special reply {special_reply_id}"
            })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/admin/cooldown-stats', methods=['GET'])
def get_cooldown_stats():
    """Get overall cooldown statistics"""
    try:
        if not chatbot.history_manager.connection:
            return jsonify({
                "status": "error", 
                "message": "Database connection not available"
            }), 500
        
        cursor = chatbot.history_manager.connection.cursor(dictionary=True)
        
        # MODIFIED: Get active cooldowns using user_id instead of phone_number
        cursor.execute('''
            SELECT 
                COUNT(*) as total_active_cooldowns,
                COUNT(DISTINCT user_id) as users_with_cooldowns,
                COUNT(DISTINCT special_reply_id) as special_replies_with_cooldowns
            FROM special_reply_usage sru
            JOIN special_reply sr ON sru.special_reply_id = sr.PKKEY
            WHERE TIMESTAMPDIFF(HOUR, sru.last_triggered_at, CONVERT_TZ(UTC_TIMESTAMP(), '+00:00', '+08:00')) < sr.COOLDOWN_HOURS
        ''')
        
        cooldown_stats = cursor.fetchone()
        
        # MODIFIED: Get top triggered special replies using user_id instead of phone_number
        cursor.execute('''
            SELECT 
                sru.special_reply_id,
                sr.KEYWORDS,
                COUNT(*) as total_triggers,
                COUNT(DISTINCT sru.user_id) as unique_users
            FROM special_reply_usage sru
            JOIN special_reply sr ON sru.special_reply_id = sr.PKKEY
            GROUP BY sru.special_reply_id, sr.KEYWORDS
            ORDER BY total_triggers DESC
            LIMIT 10
        ''')
        
        top_triggered = cursor.fetchall()
        cursor.close()
        
        return jsonify({
            "status": "success",
            "cooldown_statistics": cooldown_stats,
            "top_triggered_special_replies": top_triggered
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
        
# Additional endpoints for media and administration
@app.route('/admin/media-messages', methods=['GET'])
def get_media_messages_endpoint():
    """API endpoint to get all media messages"""
    try:
        phone_number = request.args.get('phone_number')
        media_type = request.args.get('media_type')  # 'image', 'video', 'document', 'audio'
        limit = request.args.get('limit', 50, type=int)
        
        messages = chatbot.history_manager.get_media_messages(phone_number, media_type, limit)
        
        return jsonify({
            "status": "success",
            "media_messages": messages,
            "total": len(messages),
            "media_type": media_type or "all"
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/health')
def health_check():
    return jsonify({"status": "ok"}), 200

@app.route('/')
def home():
    return jsonify({"message": "Chatbot API is running."})
    
@app.route('/debug-mysql', methods=['GET'])
def debug_mysql():
    """Debug MySQL connection issues"""
    debug_info = {
        "mysql_url_exists": bool(os.getenv("MYSQL_URL")),
        "mysql_host": os.getenv("MYSQL_HOST", "Not set"),
        "mysql_user": os.getenv("MYSQL_USER", "Not set"),
        "mysql_database": os.getenv("MYSQL_DATABASE", "Not set"),
        "mysql_port": os.getenv("MYSQL_PORT", "Not set"),
        "mysql_password_exists": bool(os.getenv("MYSQL_PASSWORD")),
        "connection_status": "Not attempted"
    }
    
    try:
        # Try to create a connection
        test_connection = chatbot.history_manager.create_connection()
        if test_connection and test_connection.is_connected():
            debug_info["connection_status"] = "SUCCESS"
            test_connection.close()
        else:
            debug_info["connection_status"] = "FAILED - Connection not active"
    except Exception as e:
        debug_info["connection_status"] = f"FAILED - Error: {str(e)}"
    
    return jsonify(debug_info)

def validate_environment():
    """Validate required environment variables"""
    required_vars = [
        "OPENAI_API_KEY",
        "META_WHATSAPP_TOKEN", 
        "META_WHATSAPP_PHONE_ID",
        "META_VERIFY_TOKEN"
    ]
    
    missing_vars = [var for var in required_vars if not os.getenv(var)]
    
    if missing_vars:
        logger.error(f"Missing required environment variables: {', '.join(missing_vars)}")
        return False
    
    return True
    
if __name__ == "__main__":
    print("📊 Enhanced WhatsApp Bot with Language Detection Starting...")
    print(f"🕐 Current Malaysia Time: {get_malaysia_timestamp()}")  # ADD this line
    
    # Validate environment
    if not validate_environment():
        print("❌ Please check your .env file and ensure all required variables are set.")
        exit(1)

    print("🚀 Starting server...")
    
    # Check if chatbot initialized properly
    try:
        if chatbot.history_manager.connection:
            print("✅ Database connected with Malaysia timezone")  # Modified message
        else:
            print("⚠️ Database not connected - some features may not work")
    except Exception as e:
        print(f"⚠️ Database connection check failed: {e}")

    # Get port from Railway environment - Railway automatically provides this
    port = int(os.getenv('PORT', 8000))
    print(f"🌐 Starting server on port {port}")
    
    # IMPORTANT: Set debug=False for production and ensure proper host binding
    app.run(
        host='0.0.0.0',  # Bind to all interfaces (required for Railway)
        port=port,       # Use Railway's PORT
        debug=False,     # MUST be False in production
        threaded=True    # Enable threading for better performance
    )